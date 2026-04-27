"""Generate realistic source documents to pair with the real-world
templates downloaded into ``tests/real_world/``.

These sources are NOT copies of the templates — they carry realistic
content a Brazilian university department would put into a POP, in
free-form prose, so the LLM mapper has to do real work segmenting and
filling.

Output:

- ``tests/real_world/source_unifap.docx`` — ``POP — Solicitação de
  Compras de Material de Consumo`` with sections written as
  free-flowing paragraphs.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


def build_source_unifap(path: Path) -> None:
    doc = Document()

    sec = doc.sections[0]
    hdr = sec.header
    hdr.paragraphs[0].text = "Universidade Federal do Amapá — DIPLAN"
    hdr.add_paragraph("POP-DIPLAN-014 v1.2 — abril/2024")

    doc.add_paragraph("POP — Solicitação de Compras de Material de Consumo")

    doc.add_paragraph("Descrição")
    doc.add_paragraph(
        "Este procedimento descreve as etapas para solicitação, aprovação "
        "e acompanhamento da aquisição de materiais de consumo (papel, "
        "toner, material de escritório) pelas unidades acadêmicas e "
        "administrativas da UNIFAP, integrando-se ao calendário anual de "
        "compras consolidado pela Divisão de Planejamento de Aquisições."
    )

    doc.add_paragraph("Objetivos")
    doc.add_paragraph(
        "Padronizar o fluxo de solicitação para reduzir retrabalho, "
        "eliminar pedidos fora do calendário oficial e garantir "
        "rastreabilidade ponta-a-ponta — do requerente até a entrega."
    )

    doc.add_paragraph("Público-Alvo")
    doc.add_paragraph(
        "Servidores técnico-administrativos lotados em qualquer setor "
        "responsável por requisição, bem como chefes de divisão que "
        "homologam pedidos."
    )

    doc.add_paragraph("Pré-requisitos")
    doc.add_paragraph(
        "Acesso ao SIPAC, perfil 'Requisitor' ativo, e relatório de "
        "consumo dos últimos 12 meses do setor solicitante."
    )

    doc.add_paragraph("Responsáveis")
    doc.add_paragraph(
        "Chefe da Divisão de Planejamento de Aquisições — define o "
        "calendário e consolida as demandas. "
        "Chefe da Divisão de Materiais — orça e elabora o termo de "
        "referência. "
        "Procurador Chefe — analisa juridicamente. "
        "Ordenador de Despesa — autoriza."
    )

    doc.add_paragraph("Lista de Contatos")
    doc.add_paragraph("Maria Lopes — DIPLAN — (96) 3213-1010 — diplan@unifap.br")
    doc.add_paragraph("João Pedro — DIMAT — (96) 3213-1020 — dimat@unifap.br")
    doc.add_paragraph("Beatriz Rocha — Procuradoria — (96) 3213-1030 — proc@unifap.br")

    doc.add_paragraph("Atividades")
    doc.add_paragraph(
        "1. Setor identifica necessidade de compra com base em consumo histórico e estoque atual."
    )
    doc.add_paragraph(
        "2. Requerente acessa o SIPAC, abre requisição de material de consumo, anexa justificativa."
    )
    doc.add_paragraph("3. Chefe da unidade homologa a requisição no sistema.")
    doc.add_paragraph("4. DIPLAN recebe a requisição, valida calendário e disponibilidade orçamentária.")
    doc.add_paragraph("5. DIMAT realiza pesquisa de preços e elabora termo de referência.")
    doc.add_paragraph("6. Procuradoria analisa juridicamente o processo.")
    doc.add_paragraph("7. Ordenador autoriza a abertura do certame.")
    doc.add_paragraph("8. Pregoeiro conduz o procedimento licitatório.")

    doc.add_paragraph("Definições")
    doc.add_paragraph(
        "SIPAC: Sistema Integrado de Patrimônio, Administração e Contratos. "
        "Termo de Referência: documento técnico que descreve o objeto a "
        "ser contratado. "
        "Pregão eletrônico: modalidade licitatória conduzida pela "
        "internet."
    )

    doc.add_paragraph("Material de Suporte")
    doc.add_paragraph(
        "Manual SIPAC: https://sipac.unifap.br/manual. "
        "Calendário de compras 2024 disponível em "
        "https://proad.unifap.br/calendario."
    )

    doc.add_paragraph("Referências")
    doc.add_paragraph("Lei 14.133/2021 (Nova Lei de Licitações).")
    doc.add_paragraph("IN SEGES/ME 65/2021.")
    doc.add_paragraph("Resolução CONSAD/UNIFAP 12/2023.")

    doc.add_paragraph("Participantes na elaboração do documento")
    doc.add_paragraph("Maria Lopes (DIPLAN), João Pedro (DIMAT), Beatriz Rocha (Procuradoria).")

    doc.add_paragraph("Histórico de Revisões")
    rev = doc.add_table(rows=4, cols=4)
    rev.rows[0].cells[0].text = "Versão"
    rev.rows[0].cells[1].text = "Data"
    rev.rows[0].cells[2].text = "Autor"
    rev.rows[0].cells[3].text = "Alterações"
    rev.rows[1].cells[0].text = "1.0"
    rev.rows[1].cells[1].text = "15/03/2023"
    rev.rows[1].cells[2].text = "Maria Lopes"
    rev.rows[1].cells[3].text = "Versão inicial"
    rev.rows[2].cells[0].text = "1.1"
    rev.rows[2].cells[1].text = "10/09/2023"
    rev.rows[2].cells[2].text = "João Pedro"
    rev.rows[2].cells[3].text = "Adequação Lei 14.133"
    rev.rows[3].cells[0].text = "1.2"
    rev.rows[3].cells[1].text = "22/04/2024"
    rev.rows[3].cells[2].text = "Beatriz Rocha"
    rev.rows[3].cells[3].text = "Inclusão de fluxo procuradoria"

    doc.save(str(path))


def build_source_corentoc(path: Path) -> None:
    """Source for the nursing council POP template (Corentocantins)."""
    doc = Document()
    sec = doc.sections[0]
    hdr = sec.header
    hdr.paragraphs[0].text = "Hospital Geral de Palmas — Serviço de Enfermagem"

    doc.add_paragraph("POP do Serviço de Enfermagem — Administração de Medicação Endovenosa")
    doc.add_paragraph("Versão 02 — vigência a partir de 01/05/2024")

    doc.add_paragraph("Objetivo")
    doc.add_paragraph(
        "Padronizar a técnica de administração de medicação endovenosa "
        "por enfermeiros e técnicos de enfermagem do Serviço de "
        "Enfermagem do Hospital Geral de Palmas, garantindo segurança "
        "do paciente e conformidade com as boas práticas de assistência."
    )

    doc.add_paragraph("Aplicação")
    doc.add_paragraph(
        "Aplica-se a todas as unidades de internação, pronto-socorro e unidade de terapia intensiva."
    )

    doc.add_paragraph("Responsáveis")
    doc.add_paragraph(
        "Enfermeiro responsável pelo plantão — supervisiona a execução. "
        "Técnico de enfermagem — executa a administração conforme "
        "prescrição médica. "
        "Coordenador do Serviço de Enfermagem — homologa o procedimento "
        "e responde por sua atualização."
    )

    doc.add_paragraph("Materiais Necessários")
    doc.add_paragraph(
        "Bandeja, luvas de procedimento, álcool 70%, dispositivo "
        "intravenoso adequado, etiquetas para identificação, "
        "prescrição médica conferida em duas vias."
    )

    doc.add_paragraph("Procedimento")
    doc.add_paragraph("1. Higienizar as mãos antes do procedimento.")
    doc.add_paragraph(
        "2. Confirmar os 9 certos: paciente, medicamento, dose, via, "
        "horário, registro, orientação, forma e resposta."
    )
    doc.add_paragraph("3. Apresentar-se ao paciente, explicar o procedimento.")
    doc.add_paragraph("4. Posicionar o paciente confortavelmente.")
    doc.add_paragraph("5. Realizar a punção venosa observando técnica asséptica.")
    doc.add_paragraph("6. Administrar o medicamento na velocidade prescrita.")
    doc.add_paragraph("7. Observar reações imediatas.")
    doc.add_paragraph("8. Registrar o procedimento no prontuário.")
    doc.add_paragraph("9. Higienizar as mãos após o procedimento.")

    doc.add_paragraph("Cuidados de Enfermagem")
    doc.add_paragraph(
        "Vigiar sinais vitais a cada 15 minutos durante a primeira hora. "
        "Comunicar a equipe médica sobre qualquer reação adversa."
    )

    doc.add_paragraph("Referências")
    doc.add_paragraph("Resolução COFEN 564/2017.")
    doc.add_paragraph("Manual de Procedimentos de Enfermagem do MS, 2022.")
    doc.add_paragraph("ANVISA RDC 36/2013.")

    doc.add_paragraph("Histórico de Revisões")
    rev = doc.add_table(rows=3, cols=4)
    rev.rows[0].cells[0].text = "Versão"
    rev.rows[0].cells[1].text = "Data"
    rev.rows[0].cells[2].text = "Responsável"
    rev.rows[0].cells[3].text = "Alterações"
    rev.rows[1].cells[0].text = "01"
    rev.rows[1].cells[1].text = "10/01/2022"
    rev.rows[1].cells[2].text = "Enf. Carla Mendes"
    rev.rows[1].cells[3].text = "Emissão inicial"
    rev.rows[2].cells[0].text = "02"
    rev.rows[2].cells[1].text = "01/05/2024"
    rev.rows[2].cells[2].text = "Enf. Patricia Silva"
    rev.rows[2].cells[3].text = "Atualização conforme RDC 36/2013"

    doc.save(str(path))


def main() -> None:
    base = Path("tests/real_world")
    base.mkdir(parents=True, exist_ok=True)
    build_source_unifap(base / "source_unifap.docx")
    build_source_corentoc(base / "source_corentoc.docx")
    print(f"wrote {base}/source_unifap.docx + source_corentoc.docx")


if __name__ == "__main__":
    main()
