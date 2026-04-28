"""Generate a synthetic non-Engeman vendor pair for cross-vendor testing.

Vendor B differs from the Engeman pair on every dimension:

- Language: English (not PT-BR).
- Placeholder shapes: ``{{DOC_CODE}}``, ``[Title]``, ``Author:``,
  ``Reviewer:``, ``Issue Date:``, ``Document Reference:`` — different
  from Engeman's ``XXXX`` / ``(TITULO)`` / ``Elaborado:`` /
  ``Aprovado:`` / ``Data:`` / ``TITULO``.
- Section taxonomy: ``PURPOSE``, ``SCOPE``, ``REFERENCES``,
  ``DEFINITIONS``, ``PROCEDURE``, ``ROLES AND RESPONSIBILITIES``,
  ``REVISION HISTORY`` (English; not in the rules engine synonym table).
- Tables: ``# | Date | Description`` (revision history) and
  ``Activity | Owner`` (single Owner column, not duplicate).

Run:
    python scripts/build_vendor_b_fixtures.py

Generates ``tests/vendor_b/template.docx`` + ``tests/vendor_b/source.docx``.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


def build_template(path: Path) -> None:
    doc = Document()

    # Real Word section header — registered as a relationship so the
    # output keeps it after python-docx writes.
    sec = doc.sections[0]
    hdr = sec.header
    hdr.paragraphs[0].text = "Document Reference: {{DOC_CODE}}"
    hdr.add_paragraph("Revision: 0")
    hdr.add_paragraph("[Title]")

    # Title page
    p = doc.add_paragraph("{{DOC_CODE}}")
    p.runs[0].font.bold = True
    doc.add_paragraph("Author:")
    doc.add_paragraph("Reviewer:")
    doc.add_paragraph("Issue Date:")
    doc.add_paragraph("[Title]")
    doc.add_paragraph("")

    # Sections (uppercase = headings)
    for section in (
        "PURPOSE",
        "SCOPE",
        "REFERENCES",
        "DEFINITIONS",
        "PROCEDURE",
        "ROLES AND RESPONSIBILITIES",
        "REVISION HISTORY",
    ):
        p = doc.add_paragraph(section)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(12)
        # 3 empty body slots per section
        for _ in range(3):
            doc.add_paragraph("")

    # Roles table: Activity | Owner (single Owner col, no duplication).
    table = doc.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "Activity"
    table.rows[0].cells[1].text = "Owner"
    # rows[1..3] left empty for the LLM to fill.

    doc.add_paragraph("")

    # Revision history table: # | Date | Description.
    rev = doc.add_table(rows=3, cols=3)
    rev.rows[0].cells[0].text = "#"
    rev.rows[0].cells[1].text = "Date"
    rev.rows[0].cells[2].text = "Description"

    doc.save(str(path))


def build_source(path: Path) -> None:
    doc = Document()

    # Real section header carrying the source's metadata.
    sec = doc.sections[0]
    hdr = sec.header
    hdr.paragraphs[0].text = "PROC-OPS-2024-007"
    hdr.add_paragraph("Plant 4 Scheduled Shutdown Procedure")
    hdr.add_paragraph("Version 1")
    hdr.add_paragraph("Approved by: Robert Smith")
    hdr.add_paragraph("Date: 2024-03-15")

    # Sections — different wording than the template
    doc.add_paragraph("Objective")  # synonym for PURPOSE
    doc.add_paragraph(
        "List the maintenance and inspection tasks performed by Operations "
        "Technicians at Plant 4 during scheduled shutdowns."
    )

    doc.add_paragraph("Applicability")  # synonym for SCOPE
    doc.add_paragraph("All ammonia and urea production lines at the FAFEN-NE complex.")

    doc.add_paragraph("Reference Documents")  # synonym for REFERENCES
    doc.add_paragraph("ISO 9001:2015 Quality Management")
    doc.add_paragraph("ISO 14001:2015 Environmental Management")
    doc.add_paragraph("OSHA 29 CFR 1910.119 Process Safety Management")
    doc.add_paragraph("ASME PCC-1 Pressure Boundary Bolted Flange Joints")

    doc.add_paragraph("Glossary")  # synonym for DEFINITIONS
    doc.add_paragraph("PSM: Process Safety Management.")
    doc.add_paragraph("MOC: Management of Change procedure.")
    doc.add_paragraph("LOTO: Lockout/Tagout safety protocol.")
    doc.add_paragraph("PPE: Personal Protective Equipment.")

    doc.add_paragraph("Method")  # synonym for PROCEDURE
    doc.add_paragraph("Pre-shutdown checks")
    doc.add_paragraph(
        "Verify all utilities (steam, instrument air, cooling water) are available and within spec."
    )
    doc.add_paragraph("Confirm spare parts inventory at the warehouse.")
    doc.add_paragraph("Validate that all instrumentation has been calibrated within the last 12 months.")
    doc.add_paragraph("Shutdown execution")
    doc.add_paragraph("Notify the control room 24 hours before T-zero.")
    doc.add_paragraph("Reduce throughput in 10% increments every 30 minutes.")
    doc.add_paragraph("Apply LOTO once the unit is at safe state.")

    doc.add_paragraph("Roles")  # synonym for ROLES AND RESPONSIBILITIES
    doc.add_paragraph("The Plant Manager owns the following")
    doc.add_paragraph("Approve the procedure.")
    doc.add_paragraph("Assign the on-call coordinator.")
    doc.add_paragraph("The Shift Supervisor owns the following")
    doc.add_paragraph("Brief the operations team.")
    doc.add_paragraph("Verify LOTO compliance.")
    doc.add_paragraph("Sign off on shutdown completion.")

    doc.add_paragraph("Revision History")
    # Revision history goes in a body table.
    rev = doc.add_table(rows=2, cols=4)
    rev.rows[0].cells[0].text = "Version"
    rev.rows[0].cells[1].text = "Date Issued"
    rev.rows[0].cells[2].text = "Author"
    rev.rows[0].cells[3].text = "Changes"
    rev.rows[1].cells[0].text = "1"
    rev.rows[1].cells[1].text = "2024-03-15"
    rev.rows[1].cells[2].text = "Jane Doe"
    rev.rows[1].cells[3].text = "Initial release"

    doc.save(str(path))


def main() -> None:
    out_dir = Path("tests/vendor_b")
    out_dir.mkdir(parents=True, exist_ok=True)
    template_path = out_dir / "template.docx"
    source_path = out_dir / "source.docx"
    build_template(template_path)
    build_source(source_path)
    print(f"wrote {template_path}")
    print(f"wrote {source_path}")


if __name__ == "__main__":
    main()
