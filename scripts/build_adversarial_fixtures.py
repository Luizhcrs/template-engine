"""Generate adversarial vendor pairs for cross-vendor stress testing.

These fixtures target weaknesses the Engeman + Vendor B pairs do not
exercise. The goal is to break the LLM mode — what survives is what
the lib actually generalises to.

Vendors covered:

- **Vendor C** — ABNT-style academic (Brazilian undergraduate thesis).
  Title-case headings (``Resumo``, ``Introdução``, ``Metodologia``),
  three-level nested numbered sub-sections (``2.1``, ``2.1.3``,
  ``2.1.3.1``), placeholder of shape ``§§§§``, ``<<TITULO>>`` and
  underscore-run lines. Source uses long-form prose with abstracts.

- **Vendor D** — Government bilingual form (Portuguese + English in
  the same template). Section headings repeat in both languages
  (``OBJETIVO / OBJECTIVE``). Placeholders of shape
  ``[_______________]`` (variable underscore length), ``< nome >``
  (angle brackets with spaces), dotted-leader fields
  (``Endereço: ......``). Tables with merged cells and signature
  blocks.

- **Vendor E** — Legal contract. Numbered clauses ``1.``,
  ``1.1.``, ``1.1.1.``, ``1.1.1.1.``. Parties block (``CONTRATANTE`` +
  ``CONTRATADO`` with name + CNPJ + endereço placeholders). Witness +
  signature placeholders. Anexo references. Tables with ``Item /
  Quantidade / Valor unitário`` shape. Source carries the same content
  in narrative form (no formal section headings, just paragraphs in a
  draft contract).

Each fixture pair:

- ``tests/vendor_<x>/template.docx`` — the empty form.
- ``tests/vendor_<x>/source.docx`` — content the mapper has to fit
  into the template.

Run:

    python scripts/build_adversarial_fixtures.py

Then exercise every pair with mode='llm' against a real provider:

    python scripts/run_adversarial_llm.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

# =============================================================================
# Vendor C — ABNT academic
# =============================================================================


def build_vendor_c_template(path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    hdr = sec.header
    hdr.paragraphs[0].text = "UNIVERSIDADE FEDERAL DO BRASIL"
    hdr.add_paragraph("Departamento de §§§§§")  # placeholder for area

    p = doc.add_paragraph("<<TITULO_DO_TRABALHO>>")
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(14)
    doc.add_paragraph("Autor: __________________")
    doc.add_paragraph("Orientador: __________________")
    doc.add_paragraph("Data de Defesa: __/__/____")
    doc.add_paragraph("")

    # Top-level Title-case headings (lower-case detector trips here).
    for heading in (
        "Resumo",
        "Abstract",
        "1. Introdução",
        "2. Revisão da Literatura",
        "3. Metodologia",
        "4. Resultados",
        "5. Discussão",
        "6. Conclusão",
        "Referências",
    ):
        p = doc.add_paragraph(heading)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(12)
        for _ in range(2):
            doc.add_paragraph("")

    # Approval table (jury panel).
    table = doc.add_table(rows=4, cols=3)
    table.rows[0].cells[0].text = "Função"
    table.rows[0].cells[1].text = "Nome"
    table.rows[0].cells[2].text = "Assinatura"

    doc.add_paragraph("")

    # Revision history (academic version control).
    rev = doc.add_table(rows=3, cols=3)
    rev.rows[0].cells[0].text = "Versão"
    rev.rows[0].cells[1].text = "Data"
    rev.rows[0].cells[2].text = "Modificações"

    doc.save(str(path))


def build_vendor_c_source(path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    hdr = sec.header
    hdr.paragraphs[0].text = "Universidade Federal de Minas Gerais"
    hdr.add_paragraph("Departamento de Engenharia Química")

    # Cover info
    doc.add_paragraph("Otimização de Reatores Catalíticos via Aprendizado de Máquina")
    doc.add_paragraph("Maria Silva Santos")
    doc.add_paragraph("Prof. Dr. João Carlos Pereira")
    doc.add_paragraph("15 de março de 2024")

    # Body — long-form academic prose with sub-sections at 3 levels.
    doc.add_paragraph("Resumo")
    doc.add_paragraph(
        "Este trabalho propõe uma metodologia híbrida para otimização de "
        "reatores catalíticos industriais combinando modelagem fenomenológica "
        "com algoritmos de aprendizado de máquina. Três configurações foram "
        "testadas em escala piloto, alcançando ganho de 12% em conversão."
    )

    doc.add_paragraph("Introdução")
    doc.add_paragraph("1.1 Motivação")
    doc.add_paragraph(
        "Reatores catalíticos representam mais de 60% da capacidade produtiva "
        "da indústria petroquímica brasileira."
    )
    doc.add_paragraph("1.2 Objetivos")
    doc.add_paragraph("1.2.1 Objetivo Geral")
    doc.add_paragraph("Desenvolver pipeline híbrido para otimização operacional contínua.")
    doc.add_paragraph("1.2.2 Objetivos Específicos")
    doc.add_paragraph("Caracterizar três famílias de catalisadores comerciais.")
    doc.add_paragraph("Treinar modelos preditivos com dados de planta.")
    doc.add_paragraph("Validar ganhos via experimentação controlada.")

    doc.add_paragraph("Metodologia")
    doc.add_paragraph("3.1 Aquisição de Dados")
    doc.add_paragraph("Dados foram coletados de três plantas industriais ao longo de 18 meses.")
    doc.add_paragraph("3.2 Modelagem")
    doc.add_paragraph("3.2.1 Modelo Fenomenológico")
    doc.add_paragraph("Equações de balanço de massa e energia foram resolvidas em CFD.")
    doc.add_paragraph("3.2.2 Modelo de Aprendizado")
    doc.add_paragraph("Random Forest treinado em 50k amostras com validação cruzada k=5.")

    doc.add_paragraph("Resultados")
    doc.add_paragraph("A configuração híbrida superou ambas as baselines em 9 dos 10 testes.")

    doc.add_paragraph("Conclusão")
    doc.add_paragraph(
        "A integração de modelos fenomenológicos com aprendizado de máquina "
        "viabiliza otimização contínua sem perda de interpretabilidade."
    )

    # Source has its own version table.
    rev = doc.add_table(rows=3, cols=3)
    rev.rows[0].cells[0].text = "Versão"
    rev.rows[0].cells[1].text = "Data"
    rev.rows[0].cells[2].text = "Modificações"
    rev.rows[1].cells[0].text = "1.0"
    rev.rows[1].cells[1].text = "2024-03-15"
    rev.rows[1].cells[2].text = "Versão inicial defendida"
    rev.rows[2].cells[0].text = "1.1"
    rev.rows[2].cells[1].text = "2024-04-02"
    rev.rows[2].cells[2].text = "Correções da banca"

    doc.save(str(path))


# =============================================================================
# Vendor D — Government bilingual form
# =============================================================================


def build_vendor_d_template(path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    hdr = sec.header
    hdr.paragraphs[0].text = "MINISTÉRIO DA FAZENDA / MINISTRY OF FINANCE"
    hdr.add_paragraph("Formulário Nº [______]   |   Form Nº [______]")

    doc.add_paragraph("REQUERIMENTO / APPLICATION")

    doc.add_paragraph("Nome / Name: < nome completo >")
    doc.add_paragraph("CPF / TAX ID: ___.___.___-__")
    doc.add_paragraph("Endereço / Address: ......................................")
    doc.add_paragraph("Cidade / City: __________  UF / State: __  CEP / ZIP: _____-___")
    doc.add_paragraph("")

    # Bilingual section headings — same line, slash-separated.
    for heading in (
        "1. OBJETIVO / OBJECTIVE",
        "2. JUSTIFICATIVA / JUSTIFICATION",
        "3. DOCUMENTAÇÃO ANEXA / ATTACHED DOCUMENTS",
        "4. DECLARAÇÃO / DECLARATION",
    ):
        p = doc.add_paragraph(heading)
        p.runs[0].font.bold = True
        for _ in range(3):
            doc.add_paragraph("")

    # Documents table.
    table = doc.add_table(rows=4, cols=3)
    table.rows[0].cells[0].text = "Documento / Document"
    table.rows[0].cells[1].text = "Número / Number"
    table.rows[0].cells[2].text = "Data / Date"

    doc.add_paragraph("")
    doc.add_paragraph("Local e Data / Place and Date: ____________________")
    doc.add_paragraph("Assinatura / Signature: ____________________")

    doc.save(str(path))


def build_vendor_d_source(path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    hdr = sec.header
    hdr.paragraphs[0].text = "Receita Federal — Pedido nº 2024-998-A"

    # Source mixes Portuguese narrative + tabular data, no clean section headings.
    doc.add_paragraph("Pedido de Restituição de Imposto de Renda — Pessoa Física")
    doc.add_paragraph("Solicitante: Carlos Henrique Almeida")
    doc.add_paragraph("CPF do solicitante: 123.456.789-09")
    doc.add_paragraph(
        "Endereço residencial: Rua das Flores, 123, Apto 4B, Bairro Jardim, "
        "Cidade Belo Horizonte, Estado MG, CEP 30100-100"
    )
    doc.add_paragraph(
        "Objetivo do requerimento: solicitar restituição do imposto de renda "
        "retido na fonte referente ao exercício 2023, conforme declaração "
        "anexa, no valor total de R$ 4.872,00."
    )
    doc.add_paragraph(
        "Justificativa: o requerente trabalhou sob regime CLT durante todo "
        "o ano-calendário 2023; após reconciliação dos rendimentos, foi "
        "apurado saldo a restituir."
    )
    doc.add_paragraph(
        "Documentos apresentados: declaração de imposto de renda exercício 2023, "
        "comprovantes de rendimentos da empresa Acme Ltda., extratos bancários "
        "do Banco Itaú dos meses janeiro a dezembro de 2023, e cópia da "
        "carteira de identidade RG 12.345.678."
    )
    doc.add_paragraph(
        "Declaração: declaro, sob as penas da lei, que as informações "
        "prestadas neste requerimento são verdadeiras."
    )
    doc.add_paragraph("Local: Belo Horizonte / MG")
    doc.add_paragraph("Data: 15 de abril de 2024")

    # Documents already-tabularised in the source.
    rev = doc.add_table(rows=4, cols=3)
    rev.rows[0].cells[0].text = "Documento"
    rev.rows[0].cells[1].text = "Número"
    rev.rows[0].cells[2].text = "Data"
    rev.rows[1].cells[0].text = "Declaração de IRPF"
    rev.rows[1].cells[1].text = "12345-2023"
    rev.rows[1].cells[2].text = "30/04/2023"
    rev.rows[2].cells[0].text = "Comprovante de rendimentos"
    rev.rows[2].cells[1].text = "ACME-001"
    rev.rows[2].cells[2].text = "31/12/2023"
    rev.rows[3].cells[0].text = "Extrato bancário"
    rev.rows[3].cells[1].text = "ITAU-2023"
    rev.rows[3].cells[2].text = "31/12/2023"

    doc.save(str(path))


# =============================================================================
# Vendor E — Legal contract
# =============================================================================


def build_vendor_e_template(path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    hdr = sec.header
    hdr.paragraphs[0].text = "CONTRATO Nº ____/____"

    doc.add_paragraph("CONTRATO DE PRESTAÇÃO DE SERVIÇOS")

    # Parties block.
    doc.add_paragraph(
        "CONTRATANTE: <razão social>, inscrita no CNPJ sob o nº __.___.___/____-__, com sede em __________________."
    )
    doc.add_paragraph(
        "CONTRATADO: <nome ou razão social>, inscrito(a) no CPF/CNPJ sob o nº __________________, com endereço em __________________."
    )
    doc.add_paragraph("")

    # Numbered clauses 1./2./3. with sub-clauses 1.1, 1.1.1, 1.1.1.1.
    for clause in (
        "1. CLÁUSULA PRIMEIRA — DO OBJETO",
        "2. CLÁUSULA SEGUNDA — DAS OBRIGAÇÕES",
        "3. CLÁUSULA TERCEIRA — DO PREÇO E CONDIÇÕES DE PAGAMENTO",
        "4. CLÁUSULA QUARTA — DA VIGÊNCIA",
        "5. CLÁUSULA QUINTA — DA RESCISÃO",
        "6. CLÁUSULA SEXTA — DO FORO",
    ):
        p = doc.add_paragraph(clause)
        p.runs[0].font.bold = True
        for _ in range(3):
            doc.add_paragraph("")

    # Items table.
    table = doc.add_table(rows=4, cols=4)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Descrição"
    table.rows[0].cells[2].text = "Quantidade"
    table.rows[0].cells[3].text = "Valor unitário"

    doc.add_paragraph("")
    doc.add_paragraph("Local e Data: ____________________")
    doc.add_paragraph("CONTRATANTE: ____________________")
    doc.add_paragraph("CONTRATADO: ____________________")
    doc.add_paragraph("Testemunha 1 (CPF): ____________________")
    doc.add_paragraph("Testemunha 2 (CPF): ____________________")

    doc.save(str(path))


def build_vendor_e_source(path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    hdr = sec.header
    hdr.paragraphs[0].text = "Minuta de contrato — versão para revisão"

    # Source is narrative, no formal section headings — pure draft text.
    doc.add_paragraph(
        "Está sendo celebrado um contrato de prestação de serviços entre as "
        "partes a seguir qualificadas: a empresa Tecnologia Brasil Ltda., "
        "pessoa jurídica de direito privado, inscrita no CNPJ sob o nº "
        "12.345.678/0001-90, com sede na Avenida Paulista, 1000, São Paulo/SP, "
        "doravante denominada CONTRATANTE; e o profissional autônomo Mariana "
        "Costa, inscrita no CPF sob o nº 987.654.321-00, com endereço na Rua "
        "dos Pinheiros, 456, São Paulo/SP, doravante denominada CONTRATADO."
    )
    doc.add_paragraph(
        "O objeto deste contrato é a prestação de serviços de consultoria em "
        "engenharia de dados, incluindo a arquitetura, implementação e "
        "manutenção de pipelines analíticos sobre a stack atual da CONTRATANTE."
    )
    doc.add_paragraph(
        "São obrigações do CONTRATADO: (i) executar os serviços com diligência "
        "técnica; (ii) entregar os artefatos contratados nos prazos pactuados; "
        "(iii) preservar sigilo sobre informações da CONTRATANTE."
    )
    doc.add_paragraph(
        "São obrigações do CONTRATANTE: (i) efetuar os pagamentos nos prazos "
        "estabelecidos; (ii) fornecer os recursos e acessos necessários à "
        "execução dos serviços; (iii) designar um interlocutor técnico."
    )
    doc.add_paragraph(
        "Pelo serviço prestado, o CONTRATANTE pagará ao CONTRATADO o valor "
        "total de R$ 60.000,00 (sessenta mil reais), em parcelas mensais de "
        "R$ 10.000,00 a partir do mês seguinte à assinatura, mediante emissão "
        "de nota fiscal de serviço."
    )
    doc.add_paragraph(
        "O presente contrato vigorará pelo prazo de 6 (seis) meses, contados "
        "da data de sua assinatura, prorrogável por períodos sucessivos "
        "mediante termo aditivo."
    )
    doc.add_paragraph(
        "Qualquer das partes poderá rescindir este contrato mediante "
        "comunicação prévia de 30 (trinta) dias, sem ônus, observada a "
        "obrigação de pagamento dos serviços efetivamente prestados."
    )
    doc.add_paragraph(
        "Fica eleito o foro da Comarca de São Paulo/SP para dirimir quaisquer "
        "litígios decorrentes deste contrato."
    )

    # Items already in tabular form in the source draft.
    items = doc.add_table(rows=4, cols=4)
    items.rows[0].cells[0].text = "Item"
    items.rows[0].cells[1].text = "Descrição"
    items.rows[0].cells[2].text = "Quantidade"
    items.rows[0].cells[3].text = "Valor unitário"
    items.rows[1].cells[0].text = "1"
    items.rows[1].cells[1].text = "Arquitetura de pipeline"
    items.rows[1].cells[2].text = "1"
    items.rows[1].cells[3].text = "R$ 20.000,00"
    items.rows[2].cells[0].text = "2"
    items.rows[2].cells[1].text = "Implementação dos jobs"
    items.rows[2].cells[2].text = "1"
    items.rows[2].cells[3].text = "R$ 30.000,00"
    items.rows[3].cells[0].text = "3"
    items.rows[3].cells[1].text = "Documentação técnica"
    items.rows[3].cells[2].text = "1"
    items.rows[3].cells[3].text = "R$ 10.000,00"

    doc.add_paragraph("Local e Data: São Paulo, 27 de abril de 2026")
    doc.add_paragraph("Pela CONTRATANTE: João Silva (CPF 111.222.333-44)")
    doc.add_paragraph("Pelo CONTRATADO: Mariana Costa")
    doc.add_paragraph("Testemunha 1: Pedro Almeida (CPF 555.666.777-88)")
    doc.add_paragraph("Testemunha 2: Ana Beatriz Lima (CPF 999.888.777-66)")

    doc.save(str(path))


# =============================================================================
# Driver
# =============================================================================


def main() -> None:
    out = Path("tests")
    for tag, fn_t, fn_s in (
        ("vendor_c", build_vendor_c_template, build_vendor_c_source),
        ("vendor_d", build_vendor_d_template, build_vendor_d_source),
        ("vendor_e", build_vendor_e_template, build_vendor_e_source),
    ):
        d = out / tag
        d.mkdir(parents=True, exist_ok=True)
        fn_t(d / "template.docx")
        fn_s(d / "source.docx")
        print(f"wrote {d}/template.docx + source.docx")


if __name__ == "__main__":
    main()
