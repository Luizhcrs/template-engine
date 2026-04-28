"""Unit tests for engine.section_mapper.typed_fill."""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx import Document

from engine.section_mapper.schemas.builtins import CONTACT_LIST_SCHEMA
from engine.section_mapper.typed_fill import (
    TypedFillRequest,
    apply_typed_fills,
    validate_value_for_column,
)

if TYPE_CHECKING:
    from pathlib import Path

    from engine.section_mapper.schemas.types import TableSchema


def _build_template(tmp: Path, schema: TableSchema) -> Path:
    """Tiny docx with one table whose row 0 carries the schema's
    column names and one empty body row to fill."""
    doc = Document()
    table = doc.add_table(rows=2, cols=len(schema.columns))
    for ci, col in enumerate(schema.columns):
        table.rows[0].cells[ci].text = col.name
    path = tmp / "template.docx"
    doc.save(str(path))
    return path


# --- validation -------------------------------------------------------------


def test_validate_value_for_column_accepts_valid_phone() -> None:
    from engine.section_mapper.schemas.types import ColumnType

    assert validate_value_for_column("(96) 3213-1010", ColumnType.PHONE) is True


def test_validate_value_for_column_rejects_invalid_phone() -> None:
    from engine.section_mapper.schemas.types import ColumnType

    assert validate_value_for_column("not a phone", ColumnType.PHONE) is False


def test_validate_value_for_column_accepts_valid_email() -> None:
    from engine.section_mapper.schemas.types import ColumnType

    assert validate_value_for_column("user@unifap.br", ColumnType.EMAIL) is True


def test_validate_value_for_column_rejects_invalid_email() -> None:
    from engine.section_mapper.schemas.types import ColumnType

    assert validate_value_for_column("not-an-email", ColumnType.EMAIL) is False


def test_validate_value_for_column_accepts_valid_date() -> None:
    from engine.section_mapper.schemas.types import ColumnType

    assert validate_value_for_column("15/03/2023", ColumnType.DATE) is True


def test_validate_value_for_column_rejects_invalid_date() -> None:
    from engine.section_mapper.schemas.types import ColumnType

    assert validate_value_for_column("2023-03-15", ColumnType.DATE) is False


def test_validate_value_for_column_free_accepts_anything() -> None:
    from engine.section_mapper.schemas.types import ColumnType

    assert validate_value_for_column("anything goes", ColumnType.FREE) is True


def test_validate_value_for_column_rejects_date_in_name_column() -> None:
    """UNIFAP regression: signature_box's Nome column was receiving
    ``15/03/2023`` because NAME validation used to accept any
    non-empty string. Real names don't look like dates."""
    from engine.section_mapper.schemas.types import ColumnType

    assert validate_value_for_column("15/03/2023", ColumnType.NAME) is False
    assert validate_value_for_column("(96) 3213-1010", ColumnType.NAME) is False
    assert validate_value_for_column("user@unifap.br", ColumnType.NAME) is False
    assert validate_value_for_column("1.0", ColumnType.NAME) is False
    # Real names still pass.
    assert validate_value_for_column("Maria Lopes", ColumnType.NAME) is True
    assert validate_value_for_column("João d'Avila", ColumnType.NAME) is True


def test_validate_value_for_column_rejects_phone_in_sector_column() -> None:
    from engine.section_mapper.schemas.types import ColumnType

    assert validate_value_for_column("(96) 3213-1010", ColumnType.SECTOR) is False
    assert validate_value_for_column("DIPLAN", ColumnType.SECTOR) is True


# --- apply_typed_fills ------------------------------------------------------


def test_apply_typed_fills_writes_phone_in_phone_column(tmp_path: Path) -> None:
    """Happy path: schema-driven fill writes a phone number into the
    Telefone cell. The output cell text matches the requested value."""
    template = _build_template(tmp_path, CONTACT_LIST_SCHEMA)
    output = tmp_path / "out.docx"

    request = TypedFillRequest(
        table_index=0,
        schema=CONTACT_LIST_SCHEMA,
        cell_fills={
            (1, 2): "(96) 3213-1010",  # Telefone column
            (1, 3): "user@unifap.br",  # e-mail column
        },
    )

    n = apply_typed_fills(template, output, [request])
    assert n == 2

    doc = Document(str(output))
    cells = doc.tables[0].rows[1].cells
    assert cells[2].text.strip() == "(96) 3213-1010"
    assert cells[3].text.strip() == "user@unifap.br"


def test_apply_typed_fills_rejects_value_violating_column_type(tmp_path: Path) -> None:
    """A name routed into the Telefone column is rejected — the cell
    stays empty rather than write garbage."""
    template = _build_template(tmp_path, CONTACT_LIST_SCHEMA)
    output = tmp_path / "out.docx"

    request = TypedFillRequest(
        table_index=0,
        schema=CONTACT_LIST_SCHEMA,
        cell_fills={(1, 2): "Maria Lopes"},  # name in phone col
    )

    n = apply_typed_fills(template, output, [request])
    assert n == 0  # Nothing written

    doc = Document(str(output))
    assert doc.tables[0].rows[1].cells[2].text.strip() == ""


def test_apply_typed_fills_accepts_free_text_for_name_columns(tmp_path: Path) -> None:
    """NAME columns accept any non-empty string (we don't validate
    against a "name regex" because real names contain accents,
    hyphens, particles, etc)."""
    template = _build_template(tmp_path, CONTACT_LIST_SCHEMA)
    output = tmp_path / "out.docx"

    request = TypedFillRequest(
        table_index=0,
        schema=CONTACT_LIST_SCHEMA,
        cell_fills={(1, 1): "Maria d'Avila Lopes-Silva"},  # Nome column
    )

    n = apply_typed_fills(template, output, [request])
    assert n == 1

    doc = Document(str(output))
    assert doc.tables[0].rows[1].cells[1].text.strip() == "Maria d'Avila Lopes-Silva"


def test_apply_typed_fills_handles_multiple_tables(tmp_path: Path) -> None:
    """Each request targets a specific table_index; multiple requests
    in one call write to multiple tables in the same docx."""
    doc = Document()
    # Table 0: contact list
    t0 = doc.add_table(rows=2, cols=4)
    for ci, col in enumerate(CONTACT_LIST_SCHEMA.columns):
        t0.rows[0].cells[ci].text = col.name
    # Table 1: another contact list
    t1 = doc.add_table(rows=2, cols=4)
    for ci, col in enumerate(CONTACT_LIST_SCHEMA.columns):
        t1.rows[0].cells[ci].text = col.name
    template = tmp_path / "template.docx"
    doc.save(str(template))

    output = tmp_path / "out.docx"
    requests = [
        TypedFillRequest(
            table_index=0,
            schema=CONTACT_LIST_SCHEMA,
            cell_fills={(1, 1): "Maria Lopes"},
        ),
        TypedFillRequest(
            table_index=1,
            schema=CONTACT_LIST_SCHEMA,
            cell_fills={(1, 1): "João Pedro"},
        ),
    ]

    n = apply_typed_fills(template, output, requests)
    assert n == 2

    out_doc = Document(str(output))
    assert out_doc.tables[0].rows[1].cells[1].text.strip() == "Maria Lopes"
    assert out_doc.tables[1].rows[1].cells[1].text.strip() == "João Pedro"
