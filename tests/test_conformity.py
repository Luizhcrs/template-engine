"""Tests for engine.conformity."""

from __future__ import annotations

from pathlib import Path  # noqa: TC003 — runtime needed by fixtures

import pytest
from docx import Document

from engine.conformity import (
    ConformityReport,
    DimensionResult,
    Failure,
    check_conformity,
    check_design,
    check_structural,
    check_technical,
    check_text,
    check_text_pre_extracted,
    check_visual,
    find_orphan_placeholders,
    validate_br_date,
    validate_cep,
    validate_cpf,
    validate_email,
    validate_iso_date,
    validate_phone_br,
    validate_uf,
)
from engine.conformity.report import DimensionResult as DR
from engine.conformity.structural import diff_fingerprints, fingerprint
from engine.conformity.text import _score_from_severity
from engine.hybrid_mapper import MappingResult
from engine.schema_inference import FieldSchema

# ===== fixtures =====


def _write_docx(path: Path, lines: list[str]) -> None:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(str(path))


def _write_docx_with_heading(path: Path, body: list[tuple[str, str]]) -> None:
    """Write a docx where each entry is (style, text)."""
    doc = Document()
    for style, text in body:
        doc.add_paragraph(text, style=style)
    doc.save(str(path))


@pytest.fixture
def docs(tmp_path: Path) -> tuple[Path, Path]:
    template = tmp_path / "template.docx"
    candidate = tmp_path / "candidate.docx"
    _write_docx(
        template,
        [
            "LAUDO TECNICO",
            "Codigo: {{CODIGO}}",
            "Data: {{DATA}}",
            "CPF: {{CPF}}",
        ],
    )
    _write_docx(
        candidate,
        [
            "LAUDO TECNICO",
            "Codigo: ABC-001",
            "Data: 2026-01-15",
            "CPF: 123.456.789-09",
        ],
    )
    return template, candidate


# ===== format validators =====


def test_validate_cpf_valid_and_invalid():
    assert validate_cpf("529.982.247-25") is True
    assert validate_cpf("12345678909") is True
    assert validate_cpf("111.111.111-11") is False  # all-same digits
    assert validate_cpf("123.456.789-00") is False
    assert validate_cpf("not a cpf") is False
    assert validate_cpf("") is False


def test_validate_cep():
    assert validate_cep("01310-100") is True
    assert validate_cep("01310100") is True
    assert validate_cep("0131") is False
    assert validate_cep("") is False


def test_validate_iso_date():
    assert validate_iso_date("2026-04-26") is True
    assert validate_iso_date("2026-13-40") is False
    assert validate_iso_date("26/04/2026") is False
    assert validate_iso_date("") is False


def test_validate_br_date():
    assert validate_br_date("26/04/2026") is True
    assert validate_br_date("31/02/2026") is False
    assert validate_br_date("2026-04-26") is False


def test_validate_email():
    assert validate_email("luiz@example.com") is True
    assert validate_email("luiz+tag@sub.example.co") is True
    assert validate_email("not an email") is False
    assert validate_email("@example.com") is False


def test_validate_phone_br():
    assert validate_phone_br("(81) 99999-9999") is True
    assert validate_phone_br("81999999999") is True
    assert validate_phone_br("8133333333") is True
    assert validate_phone_br("123") is False


def test_validate_uf():
    assert validate_uf("PE") is True
    assert validate_uf("sp") is True
    assert validate_uf(" pr ") is True
    assert validate_uf("XX") is False
    assert validate_uf("") is False


def test_find_orphan_placeholders():
    text = "Header {{CODIGO}}\nFooter [DATA]\nLine ___\nGood: filled"
    orphans = find_orphan_placeholders(text)
    assert "{{CODIGO}}" in orphans
    assert "[DATA]" in orphans
    assert "___" in orphans


# ===== text dimension =====


def test_text_score_from_severity_clamped():
    assert _score_from_severity(0, 0, 0, 5) == 1.0
    assert _score_from_severity(5, 0, 0, 5) == 0.0
    assert _score_from_severity(20, 0, 0, 5) == 0.0  # overflow clamped
    assert 0.5 < _score_from_severity(2, 0, 0, 5) < 0.7


@pytest.mark.asyncio
async def test_check_text_skipped_without_llm(docs):
    template, candidate = docs
    result = await check_text(template, candidate, llm=None)
    assert result.skipped is True
    assert result.score == 1.0


@pytest.mark.asyncio
async def test_check_text_uses_semantic_diff_via_stub_llm(docs):
    template, candidate = docs

    class _StubLLM:
        name = "stub"
        model = "s-1"

        async def generate_structured(self, prompt: str, json_schema: dict) -> dict:
            return {
                "discrepancies": [
                    {
                        "type": "missing_in_output",
                        "field_or_excerpt": "CPF",
                        "source_value": "123.456.789-09",
                        "output_value": None,
                        "severity": "critical",
                        "note": "lost",
                    }
                ]
            }

    result = await check_text(template, candidate, llm=_StubLLM())  # type: ignore[arg-type]
    assert result.dimension == "text"
    assert result.score < 1.0
    assert any(f.severity == "critical" for f in result.failures)


@pytest.mark.asyncio
async def test_check_text_pre_extracted_works():
    class _StubLLM:
        name = "stub"
        model = "s-1"

        async def generate_structured(self, prompt: str, json_schema: dict) -> dict:
            return {"discrepancies": []}

    result = await check_text_pre_extracted("a", "b", llm=_StubLLM())  # type: ignore[arg-type]
    assert result.score == 1.0
    assert result.failures == []


# ===== structural dimension =====


def test_structural_fingerprint_counts_headings(tmp_path):
    p = tmp_path / "x.docx"
    _write_docx_with_heading(
        p,
        [
            ("Heading 1", "Title"),
            ("Heading 2", "Section A"),
            ("Heading 2", "Section B"),
            ("Normal", "body text"),
        ],
    )
    fp = fingerprint(p)
    assert fp.headings_by_level[1] == 1
    assert fp.headings_by_level[2] == 2


def test_structural_fingerprint_counts_tables(tmp_path):
    p = tmp_path / "x.docx"
    doc = Document()
    doc.add_paragraph("intro")
    t = doc.add_table(rows=3, cols=4)
    _ = t  # unused, just to count
    doc.save(str(p))
    fp = fingerprint(p)
    assert fp.tables_count == 1
    assert fp.table_cells_total == 12


def test_structural_diff_fingerprints_emits_failures(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    _write_docx_with_heading(a, [("Heading 1", "A"), ("Heading 1", "B")])
    _write_docx_with_heading(b, [("Heading 1", "X")])
    fp_a = fingerprint(a)
    fp_b = fingerprint(b)
    failures = diff_fingerprints(fp_a, fp_b)
    assert any(f.field_or_excerpt == "headings_h1" for f in failures)


def test_check_structural_returns_dimension_result(docs):
    template, candidate = docs
    result = check_structural(template, candidate)
    assert isinstance(result, DR)
    assert result.dimension == "structural"
    assert 0.0 <= result.score <= 1.0


def test_check_structural_identical_docs_score_one(tmp_path):
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    _write_docx(a, ["x", "y", "z"])
    _write_docx(b, ["x", "y", "z"])
    result = check_structural(a, b)
    assert result.score == 1.0


# ===== technical dimension =====


def test_check_technical_passes_when_required_filled_and_formats_valid():
    schemas = [
        FieldSchema(
            name="CPF", placeholder_token="{{CPF}}", kind="mustache", field_type="cpf", required=True
        ),
        FieldSchema(
            name="DATA", placeholder_token="{{DATA}}", kind="mustache", field_type="iso_date", required=True
        ),
    ]
    mapping = {
        "CPF": MappingResult("CPF", "529.982.247-25", "regex", 1.0),
        "DATA": MappingResult("DATA", "2026-04-26", "regex", 1.0),
    }
    result = check_technical(schemas, mapping, "rendered: 529.982.247-25 / 2026-04-26")
    assert result.score == 1.0
    assert result.failures == []


def test_check_technical_flags_missing_required():
    schemas = [
        FieldSchema(name="X", placeholder_token="{{X}}", kind="mustache", required=True),
    ]
    mapping = {"X": MappingResult("X", None, "missing", 0.0)}
    result = check_technical(schemas, mapping, "doc text")
    assert result.score < 1.0
    assert any(f.severity == "critical" and "X" in f.field_or_excerpt for f in result.failures)


def test_check_technical_flags_invalid_cpf():
    schemas = [
        FieldSchema(
            name="CPF", placeholder_token="{{CPF}}", kind="mustache", field_type="cpf", required=True
        ),
    ]
    mapping = {"CPF": MappingResult("CPF", "111.111.111-11", "llm", 0.7)}
    result = check_technical(schemas, mapping, "doc")
    assert any(f.severity == "critical" and f.field_or_excerpt == "CPF" for f in result.failures)


def test_check_technical_flags_orphan_placeholders():
    schemas: list[FieldSchema] = []
    mapping: dict = {}
    result = check_technical(schemas, mapping, "Field: {{X}} stuck")
    assert any(f.field_or_excerpt == "orphan_placeholders" for f in result.failures)


# ===== visual dimension =====


def test_check_visual_skipped_without_pillow_pretends_skipped(monkeypatch, docs):
    template, candidate = docs
    monkeypatch.setattr("engine.conformity.visual._pillow_available", lambda: False)
    result = check_visual(template, candidate)
    assert result.skipped is True
    assert result.score == 1.0


def test_check_visual_runs_when_pillow_present(docs):
    pytest.importorskip("PIL")
    template, candidate = docs
    result = check_visual(template, candidate)
    assert result.dimension == "visual"
    assert 0.0 <= result.score <= 1.0


# ===== design dimension =====


@pytest.mark.asyncio
async def test_check_design_skipped_without_visual_llm(docs):
    template, candidate = docs
    result = await check_design(template, candidate, visual_llm=None)
    assert result.skipped is True
    assert result.score == 1.0


@pytest.mark.asyncio
async def test_check_design_uses_stub_provider(docs):
    template, candidate = docs

    class _StubVisual:
        name = "stub-visual"
        model = "v-1"

        async def compare_documents(self, t, c, prompt, schema):
            return {
                "score": 0.7,
                "issues": [
                    {
                        "field": "font_family",
                        "expected": "Arial",
                        "actual": "Times New Roman",
                        "severity": "warning",
                        "note": "font mismatch",
                    }
                ],
            }

    result = await check_design(template, candidate, visual_llm=_StubVisual())  # type: ignore[arg-type]
    assert result.score == 0.7
    assert len(result.failures) == 1
    assert result.failures[0].field_or_excerpt == "font_family"


@pytest.mark.asyncio
async def test_check_design_provider_error_emits_warning_failure(docs):
    """Provider error must NOT silently pass as score=1.0 — that would let a
    transient network blip translate to "design conforms". hardening fix #8.
    """
    template, candidate = docs

    class _BoomVisual:
        name = "boom"
        model = "b-1"

        async def compare_documents(self, t, c, prompt, schema):
            raise RuntimeError("upload failed")

    result = await check_design(template, candidate, visual_llm=_BoomVisual())  # type: ignore[arg-type]
    assert result.skipped is False
    assert result.score == 0.0
    assert any(f.field_or_excerpt == "provider_error" for f in result.failures)


# ===== aggregator =====


@pytest.mark.asyncio
async def test_check_conformity_returns_report_with_all_dimensions(docs):
    pytest.importorskip("PIL")
    template, candidate = docs

    schemas = [
        FieldSchema(name="CODIGO", placeholder_token="{{CODIGO}}", kind="mustache"),
    ]
    mapping = {"CODIGO": MappingResult("CODIGO", "ABC-001", "regex", 1.0)}

    report = await check_conformity(
        template,
        candidate,
        llm=None,
        visual_llm=None,
        schemas=schemas,
        mapping=mapping,
        dimensions=["text", "structural", "visual", "design", "technical"],
    )
    assert isinstance(report, ConformityReport)
    assert set(report.by_dimension.keys()) == {"text", "structural", "visual", "design", "technical"}
    # text and design were skipped (no llm) -> still get DimensionResult
    assert report.by_dimension["text"].skipped is True
    assert report.by_dimension["design"].skipped is True


@pytest.mark.asyncio
async def test_check_conformity_subset_dimensions(docs):
    template, candidate = docs
    report = await check_conformity(
        template,
        candidate,
        llm=None,
        dimensions=["structural"],
    )
    assert list(report.by_dimension.keys()) == ["structural"]


@pytest.mark.asyncio
async def test_check_conformity_threshold_decides_is_conformant(docs):
    template, candidate = docs
    schemas = [
        FieldSchema(name="X", placeholder_token="{{X}}", kind="mustache", required=True),
    ]
    mapping = {"X": MappingResult("X", None, "missing", 0.0)}
    report = await check_conformity(
        template,
        candidate,
        schemas=schemas,
        mapping=mapping,
        dimensions=["technical"],
        threshold=0.99,
    )
    # Critical failure pushes technical score below threshold
    assert report.is_conformant is False


@pytest.mark.asyncio
async def test_report_to_dict_serializable(docs):
    template, candidate = docs
    report = await check_conformity(
        template,
        candidate,
        dimensions=["structural"],
    )
    import json

    data = report.to_dict()
    serialized = json.dumps(data)
    parsed = json.loads(serialized)
    assert "score" in parsed
    assert "by_dimension" in parsed


# ===== sanity =====


def test_failure_is_frozen():
    f = Failure(
        dimension="text",
        field_or_excerpt="x",
        expected="a",
        actual="b",
        severity="warning",
        note="",
    )
    with pytest.raises((AttributeError, Exception)):
        f.severity = "critical"  # type: ignore[misc]


def test_dimension_result_passed_property():
    dr = DimensionResult(dimension="x", score=1.0)
    assert dr.passed is True
    dr2 = DimensionResult(dimension="x", score=0.99)
    assert dr2.passed is False
