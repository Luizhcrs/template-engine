"""Tests for engine.batch — orchestrator end-to-end."""

from __future__ import annotations

from pathlib import Path  # noqa: TC003 — used at runtime in fixture annotations

import pytest
from docx import Document

from engine.batch import (
    BatchItemResult,
    _apply_mapping_to_template,
    _classify_tier,
    _replace_tokens_in_paragraph,
    normalize_batch,
)
from engine.hybrid_mapper import MappingResult
from engine.schema_inference import FieldSchema
from engine.semantic_diff import Discrepancy

# ===== fixtures =====


def _write_template(path: Path, body_lines: list[str]) -> None:
    doc = Document()
    for line in body_lines:
        doc.add_paragraph(line)
    doc.save(str(path))


def _write_source(path: Path, body_lines: list[str]) -> None:
    _write_template(path, body_lines)


@pytest.fixture
def template_and_sources(tmp_path: Path) -> tuple[Path, Path, Path]:
    template = tmp_path / "template.docx"
    _write_template(
        template,
        [
            "LAUDO TECNICO",
            "Codigo: {{CODIGO}}",
            "Data: {{DATA}}",
            "Cliente: {{CLIENTE}}",
        ],
    )

    source_dir = tmp_path / "sources"
    source_dir.mkdir()
    _write_source(
        source_dir / "doc1.docx",
        [
            "LAUDO TECNICO",
            "Codigo: ABC-001",
            "Data: 2026-01-15",
            "Cliente: Empresa Alpha",
        ],
    )
    _write_source(
        source_dir / "doc2.docx",
        [
            "LAUDO TECNICO",
            "Codigo: ABC-002",
            "Data: 2026-02-20",
            "Cliente: Cia Beta",
        ],
    )
    output_dir = tmp_path / "out"
    return template, source_dir, output_dir


# ===== _classify_tier =====


def test_classify_high_when_all_regex_no_discrepancies():
    schemas = [FieldSchema(name="X", placeholder_token="{{X}}", kind="mustache")]
    mapping = {"X": MappingResult("X", "v", "regex", 1.0)}
    tier = _classify_tier(mapping, [], schemas=schemas)
    assert tier == "high"


def test_classify_medium_when_llm_used():
    schemas = [FieldSchema(name="X", placeholder_token="{{X}}", kind="mustache")]
    mapping = {"X": MappingResult("X", "v", "llm", 0.8)}
    tier = _classify_tier(mapping, [], schemas=schemas)
    assert tier == "medium"


def test_classify_medium_when_warning_discrepancy():
    schemas = [FieldSchema(name="X", placeholder_token="{{X}}", kind="mustache")]
    mapping = {"X": MappingResult("X", "v", "regex", 1.0)}
    discrepancies = [Discrepancy("missing_in_output", "extra", "v", None, "warning", "")]
    tier = _classify_tier(mapping, discrepancies, schemas=schemas)
    assert tier == "medium"


def test_classify_low_when_critical_discrepancy():
    schemas = [FieldSchema(name="X", placeholder_token="{{X}}", kind="mustache")]
    mapping = {"X": MappingResult("X", "v", "regex", 1.0)}
    discrepancies = [Discrepancy("missing_in_output", "CPF", "x", None, "critical", "")]
    tier = _classify_tier(mapping, discrepancies, schemas=schemas)
    assert tier == "low"


def test_classify_low_when_required_field_missing():
    schemas = [
        FieldSchema(name="X", placeholder_token="{{X}}", kind="mustache", required=True),
    ]
    mapping = {"X": MappingResult("X", None, "missing", 0.0)}
    tier = _classify_tier(mapping, [], schemas=schemas)
    assert tier == "low"


def test_classify_high_when_optional_field_missing():
    schemas = [
        FieldSchema(name="X", placeholder_token="{{X}}", kind="mustache", required=True),
        FieldSchema(name="Y", placeholder_token="{{Y}}", kind="mustache", required=False),
    ]
    mapping = {
        "X": MappingResult("X", "v", "regex", 1.0),
        "Y": MappingResult("Y", None, "missing", 0.0),
    }
    tier = _classify_tier(mapping, [], schemas=schemas)
    assert tier == "high"


# ===== normalize_batch — regex-only path =====


@pytest.mark.asyncio
async def test_normalize_batch_regex_only_no_llm(template_and_sources, tmp_path):
    template, source_dir, output_dir = template_and_sources

    gold_docs = [
        "Codigo: ABC-001\nData: 2026-01-15\nCliente: Empresa Alpha",
        "Codigo: ABC-002\nData: 2026-02-20\nCliente: Cia Beta",
        "Codigo: ABC-003\nData: 2026-03-30\nCliente: Org Gamma",
    ]
    field_examples = {
        "CODIGO": ["ABC-001", "ABC-002", "ABC-003"],
        "DATA": ["2026-01-15", "2026-02-20", "2026-03-30"],
        "CLIENTE": ["Empresa Alpha", "Cia Beta", "Org Gamma"],
    }

    report = await normalize_batch(
        template,
        source_dir,
        output_dir,
        llm=None,
        gold_docs=gold_docs,
        field_examples=field_examples,
    )

    # All docs went through regex tier alone
    assert report.by_tier["high"] >= 1
    assert report.llm_call_count == 0
    assert len(report.items) == 2

    # Output docx files exist + have substituted values
    for item in report.items:
        assert item.output_path is not None
        assert item.output_path.exists()
        out_doc = Document(str(item.output_path))
        text = "\n".join(p.text for p in out_doc.paragraphs)
        assert "{{CODIGO}}" not in text  # placeholder substituted
        assert "ABC-" in text


@pytest.mark.asyncio
async def test_normalize_batch_writes_outputs_named_after_sources(template_and_sources, tmp_path):
    template, source_dir, output_dir = template_and_sources
    report = await normalize_batch(template, source_dir, output_dir, llm=None)

    output_names = {item.output_path.name for item in report.items if item.output_path}
    assert "doc1.normalized.docx" in output_names
    assert "doc2.normalized.docx" in output_names


@pytest.mark.asyncio
async def test_normalize_batch_handles_empty_source_dir(tmp_path):
    template = tmp_path / "tpl.docx"
    _write_template(template, ["x"])
    src = tmp_path / "src"
    src.mkdir()
    out = tmp_path / "out"
    report = await normalize_batch(template, src, out, llm=None)
    assert report.items == []
    assert report.by_tier == {"high": 0, "medium": 0, "low": 0, "error": 0}


# ===== normalize_batch — LLM fallback path =====


class _StubLLM:
    name = "stub"
    model = "stub-1"

    def __init__(self, *, fallback_response: dict | None = None) -> None:
        self.fallback_response = fallback_response or {}
        self.calls = 0

    async def generate_structured(self, prompt: str, json_schema: dict) -> dict:
        self.calls += 1
        # Schema enrichment: response shape {field_type, format_hint, required}
        if "Placeholder name:" in prompt:
            return {"field_type": "freetext", "format_hint": None, "required": True}
        # Semantic diff response
        if "<<<SOURCE_START>>>" in prompt:
            return {"discrepancies": []}
        # Hybrid fallback response
        return self.fallback_response


@pytest.mark.asyncio
async def test_normalize_batch_uses_llm_for_missing_fields(template_and_sources):
    template, source_dir, output_dir = template_and_sources
    stub = _StubLLM(
        fallback_response={
            "CODIGO": {"value": "FALLBACK-CODE", "confidence": 0.6},
            "DATA": {"value": "2099-01-01", "confidence": 0.6},
            "CLIENTE": {"value": "Fallback Cia", "confidence": 0.6},
        }
    )

    # No gold docs / examples → regex tier skipped → all routes to LLM
    report = await normalize_batch(
        template,
        source_dir,
        output_dir,
        llm=stub,  # type: ignore[arg-type]
    )

    assert report.llm_call_count > 0
    # Each item should have all 3 fields filled by LLM
    for item in report.items:
        assert item.tier in {"medium", "low"}  # never "high" without regex
        assert all(r.source == "llm" for r in item.mapping.values())


# ===== BatchReport.to_dict =====


@pytest.mark.asyncio
async def test_report_to_dict_is_json_serializable(template_and_sources):
    template, source_dir, output_dir = template_and_sources
    report = await normalize_batch(template, source_dir, output_dir, llm=None)
    import json

    data = report.to_dict()
    serialized = json.dumps(data)
    parsed = json.loads(serialized)
    assert "by_tier" in parsed
    assert "items" in parsed
    assert "fields" in parsed


@pytest.mark.asyncio
async def test_report_includes_per_item_mapping_summary(template_and_sources):
    template, source_dir, output_dir = template_and_sources
    report = await normalize_batch(template, source_dir, output_dir, llm=None)

    data = report.to_dict()
    for item_dict in data["items"]:
        assert "mapping_summary" in item_dict
        assert "by_source" in item_dict["mapping_summary"]


# ===== error handling =====


@pytest.mark.asyncio
async def test_failed_doc_marked_as_error_tier_without_killing_batch(tmp_path):
    template = tmp_path / "tpl.docx"
    _write_template(template, ["{{X}}"])
    src = tmp_path / "src"
    src.mkdir()
    # Valid doc
    _write_source(src / "good.docx", ["X: ok"])
    # Invalid doc — write a file with .docx extension but invalid content
    bad = src / "bad.docx"
    bad.write_bytes(b"not a docx at all")

    out = tmp_path / "out"
    report = await normalize_batch(template, src, out, llm=None)

    tiers = [item.tier for item in report.items]
    assert "error" in tiers
    error_item = next(i for i in report.items if i.tier == "error")
    assert error_item.error is not None


# ===== fragmented runs renderer =====


def _add_fragmented_paragraph(doc, fragments: list[str]):  # type: ignore[no-untyped-def]
    """Helper: add a paragraph whose text is split across N runs.

    Word frequently splits ``{{NAME}}`` into 3 or more runs whenever the user
    edits the template (different fonts, autocomplete etc). The renderer must
    cope with this.
    """
    p = doc.add_paragraph()
    for fragment in fragments:
        run = p.add_run(fragment)
        _ = run  # keep formatting default
    return p


def test_replace_tokens_handles_token_in_single_run(tmp_path):
    """Sanity: simplest case still works after the two-pass refactor."""
    p = tmp_path / "x.docx"
    doc = Document()
    doc.add_paragraph("Codigo: {{CODIGO}}")
    doc.save(str(p))

    out = tmp_path / "out.docx"
    schemas = [FieldSchema("CODIGO", "{{CODIGO}}", "mustache")]
    mapping = {"CODIGO": MappingResult("CODIGO", "ABC-001", "regex", 1.0)}
    _apply_mapping_to_template(p, mapping, schemas, out)

    text = "\n".join(par.text for par in Document(str(out)).paragraphs)
    assert "{{CODIGO}}" not in text
    assert "ABC-001" in text


def test_replace_tokens_handles_token_split_across_runs(tmp_path):
    """The placeholder ``{{CODIGO}}`` is split into 3 runs.

    Without the pass-2 paragraph-level fallback this regresses silently —
    the docx ships with the literal ``{{CODIGO}}`` still in it.
    """
    p = tmp_path / "x.docx"
    doc = Document()
    _add_fragmented_paragraph(doc, ["Codigo: {{", "CODIGO", "}}"])
    doc.save(str(p))

    out = tmp_path / "out.docx"
    schemas = [FieldSchema("CODIGO", "{{CODIGO}}", "mustache")]
    mapping = {"CODIGO": MappingResult("CODIGO", "ABC-042", "regex", 1.0)}
    _apply_mapping_to_template(p, mapping, schemas, out)

    text = "\n".join(par.text for par in Document(str(out)).paragraphs)
    assert "{{CODIGO}}" not in text
    assert "{{" not in text
    assert "}}" not in text
    assert "ABC-042" in text


def test_replace_tokens_handles_multiple_tokens_split_across_runs(tmp_path):
    p = tmp_path / "x.docx"
    doc = Document()
    _add_fragmented_paragraph(doc, ["Codigo: {{", "CODIGO", "}}, Data: ", "{{DATA}}"])
    doc.save(str(p))

    out = tmp_path / "out.docx"
    schemas = [
        FieldSchema("CODIGO", "{{CODIGO}}", "mustache"),
        FieldSchema("DATA", "{{DATA}}", "mustache"),
    ]
    mapping = {
        "CODIGO": MappingResult("CODIGO", "ABC-099", "regex", 1.0),
        "DATA": MappingResult("DATA", "2026-04-26", "regex", 1.0),
    }
    _apply_mapping_to_template(p, mapping, schemas, out)

    text = "\n".join(par.text for par in Document(str(out)).paragraphs)
    assert "{{CODIGO}}" not in text
    assert "{{DATA}}" not in text
    assert "ABC-099" in text
    assert "2026-04-26" in text


def test_replace_tokens_in_table_cells_with_fragmented_runs(tmp_path):
    """Same fragmentation can happen inside table cells."""
    p = tmp_path / "x.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    cell = table.rows[0].cells[0]
    cell.text = ""  # clear default
    para = cell.paragraphs[0]
    for fragment in ("Cliente: {{", "CLIENTE", "}}"):
        para.add_run(fragment)
    doc.save(str(p))

    out = tmp_path / "out.docx"
    schemas = [FieldSchema("CLIENTE", "{{CLIENTE}}", "mustache")]
    mapping = {"CLIENTE": MappingResult("CLIENTE", "Empresa Alpha", "regex", 1.0)}
    _apply_mapping_to_template(p, mapping, schemas, out)

    out_doc = Document(str(out))
    cell_text = "\n".join(para.text for para in out_doc.tables[0].rows[0].cells[0].paragraphs)
    assert "{{CLIENTE}}" not in cell_text
    assert "Empresa Alpha" in cell_text


def test_replace_tokens_no_op_when_no_token_present(tmp_path):
    p = tmp_path / "x.docx"
    doc = Document()
    doc.add_paragraph("Plain paragraph without placeholders.")
    doc.save(str(p))

    out = tmp_path / "out.docx"
    schemas: list[FieldSchema] = []
    _apply_mapping_to_template(p, {}, schemas, out)

    text = "\n".join(par.text for par in Document(str(out)).paragraphs)
    assert text == "Plain paragraph without placeholders."


def test_replace_tokens_in_paragraph_helper_directly():
    """Unit test on the helper without going through file IO."""
    doc = Document()
    p = doc.add_paragraph()
    for frag in ("Hello ", "{{NAME", "}}, welcome"):
        p.add_run(frag)
    _replace_tokens_in_paragraph(p, {"{{NAME}}": "Luiz"})
    assert p.text == "Hello Luiz, welcome"


# Suppress unused import warning
_ = BatchItemResult
