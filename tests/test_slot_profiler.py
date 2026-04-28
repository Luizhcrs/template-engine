"""Unit tests for engine.section_mapper.slot_profiler.

Focus on the empty-paragraph fillability rule, which is what regulates
how aggressively the profiler floods the LLM with bogus body slots.

The Corentocantins POP regression (2026-04-27): 19 consecutive empty
paragraphs sit between two all-caps title lines (page padding before
a mega-table). The previous rule "any empty after a heading is a
slot" flagged all 19 as fillable, causing the LLM to either fill them
with arbitrary header text or skip them — blowing up false-positive
slot count from ~16 to 86.
"""

from __future__ import annotations

from engine.section_mapper.slot_profiler import (
    _empty_idxs_under_headings,
    _looks_like_heading,
)


class _FakePara:
    """Minimal stand-in for ``docx.text.paragraph.Paragraph`` — only
    needs a ``.text`` attribute for these tests."""

    def __init__(self, text: str) -> None:
        self.text = text


def _build(paragraphs: list[str]) -> list[_FakePara]:
    return [_FakePara(t) for t in paragraphs]


# --- _looks_like_heading ------------------------------------------------------


def test_looks_like_heading_uppercase_multiword() -> None:
    assert _looks_like_heading("LISTA DE CONTATOS:")
    assert _looks_like_heading("LOGOMARCA DA INSTITUICAO")


def test_looks_like_heading_rejects_sentence_with_period() -> None:
    assert not _looks_like_heading("This is a normal sentence.")


def test_looks_like_heading_rejects_long_uppercase() -> None:
    long_caps = "A" * 100
    assert not _looks_like_heading(long_caps)


# --- _empty_idxs_under_headings: short runs are slots ------------------------


def test_one_empty_between_two_headings_is_fillable() -> None:
    """UNIFAP-style: heading -> empty -> heading. Single empty stays
    fillable so the LLM can drop content under the first heading."""
    paras = _build(["LISTA DE CONTATOS:", "", "LEGENDA"])
    assert _empty_idxs_under_headings(paras) == {1}


def test_two_empties_between_headings_are_fillable() -> None:
    """Cap-2 rule: a run of exactly 2 empties is still treated as a
    legitimate body slot (some templates leave a paragraph + spacer)."""
    paras = _build(["OBJETIVO", "", "", "ESCOPO"])
    assert _empty_idxs_under_headings(paras) == {1, 2}


def test_empty_between_heading_and_body_is_fillable() -> None:
    paras = _build(["OBJETIVO", "", "Conteudo do objetivo aqui"])
    assert _empty_idxs_under_headings(paras) == {1}


# --- _empty_idxs_under_headings: long runs are PADDING, not slots ------------


def test_three_empties_between_headings_are_padding() -> None:
    """A run of 3 empties between two headings is page padding, not a
    fillable slot — drop the lot."""
    paras = _build(["LOGOMARCA", "", "", "", "PROCEDIMENTO"])
    assert _empty_idxs_under_headings(paras) == set()


def test_corentocantins_19_empties_between_titles_are_padding() -> None:
    """Regression for the Corentocantins POP: between the two title
    lines there are 19 empty paragraphs of page-layout padding. None
    of them should be fillable."""
    paras = _build(["LOGOMARCA DA INSTITUICAO", *([""] * 19), "PROCEDIMENTO OPERACIONAL PADRAO"])
    assert _empty_idxs_under_headings(paras) == set()


def test_long_run_of_empties_at_end_of_doc_are_padding() -> None:
    """Empties trailing the last heading with no content after them are
    padding before the next table or end-of-doc, not slots."""
    paras = _build(["OBJETIVO GERAL", "", "", "", "", "", ""])
    assert _empty_idxs_under_headings(paras) == set()


# --- _empty_idxs_under_headings: heading must precede the empty -------------


def test_empty_after_body_is_not_fillable() -> None:
    paras = _build(["Texto qualquer", "", "Mais texto"])
    assert _empty_idxs_under_headings(paras) == set()


def test_empty_at_top_of_doc_is_not_fillable() -> None:
    paras = _build(["", "", "OBJETIVO GERAL"])
    assert _empty_idxs_under_headings(paras) == set()


# --- vMerge continuation + image cells ---------------------------------------


def test_iter_row_tcs_skips_vmerge_continuation_cells(
    tmp_path,  # type: ignore[no-untyped-def]
) -> None:
    """A vertically-merged group should produce ONE slot per logical
    cell, not one per row participating in the merge. The restart cell
    carries the content; the continuation cells are visual padding and
    must NOT be returned."""
    import shutil

    from engine.section_mapper.slot_profiler import _iter_row_tcs

    base = _make_minimal_docx(tmp_path / "base.docx")
    target = tmp_path / "vmerge.docx"
    shutil.copy(base, target)

    table_xml = """<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr>
      <w:tblGrid><w:gridCol w:w="2000"/><w:gridCol w:w="2000"/></w:tblGrid>
      <w:tr>
        <w:tc>
          <w:tcPr><w:vMerge w:val="restart"/></w:tcPr>
          <w:p><w:r><w:t>logo cell</w:t></w:r></w:p>
        </w:tc>
        <w:tc>
          <w:p><w:r><w:t>r0c1</w:t></w:r></w:p>
        </w:tc>
      </w:tr>
      <w:tr>
        <w:tc>
          <w:tcPr><w:vMerge/></w:tcPr>
          <w:p/>
        </w:tc>
        <w:tc>
          <w:p><w:r><w:t>r1c1</w:t></w:r></w:p>
        </w:tc>
      </w:tr>
    </w:tbl>"""

    _inject_table_into_docx(target, table_xml)

    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(target))
    tab = doc.tables[0]
    rows = tab._tbl.findall(qn("w:tr"))
    # Row 0 has 2 cells (restart + normal). Row 1 should yield 1 cell
    # (the vMerge continuation must be skipped).
    assert len(_iter_row_tcs(rows[0])) == 2
    assert len(_iter_row_tcs(rows[1])) == 1


def test_classify_cell_with_image_is_data(
    tmp_path,  # type: ignore[no-untyped-def]
) -> None:
    """Cells containing a drawing (logo, picture) must NOT be marked
    fillable — the LLM has nothing useful to write into them."""
    import shutil

    from engine.section_mapper.slot_profiler import profile_slots

    base = _make_minimal_docx(tmp_path / "base.docx")
    target = tmp_path / "image.docx"
    shutil.copy(base, target)

    # A tc that carries an empty paragraph but also has a <w:drawing>
    # descendant. The text is empty so the old rule would flag it as
    # fillable.
    table_xml = """<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr>
      <w:tblGrid><w:gridCol w:w="2000"/></w:tblGrid>
      <w:tr>
        <w:tc>
          <w:p>
            <w:r>
              <w:drawing/>
            </w:r>
          </w:p>
        </w:tc>
      </w:tr>
    </w:tbl>"""

    _inject_table_into_docx(target, table_xml)

    inv = profile_slots(target)
    cell_slots = [s for s in inv.slots if s.address.location == "table_cell"]
    assert len(cell_slots) == 1
    assert cell_slots[0].is_fillable is False
    assert cell_slots[0].kind == "data"


# --- _classify: placeholder shapes -------------------------------------------


def test_classify_long_xxx_placeholder_inline_in_prose() -> None:
    """UNIFAP regression: ``Chefe da Divisão XXXXX ou Diretor XXXX, etc``
    is a placeholder even though the text is longer than 40 chars."""
    from engine.section_mapper.slot_profiler import _classify

    kind, fillable = _classify("Chefe da Divisão XXXXX ou Diretor XXXX, etc")
    assert kind == "placeholder"
    assert fillable is True


def test_classify_trailing_isolated_x_is_placeholder() -> None:
    """UNIFAP regression: ``Gestor do processo X`` and ``Diretor do
    Departamento X`` end with an isolated ``X`` token used as
    placeholder marker."""
    from engine.section_mapper.slot_profiler import _classify

    kind, _ = _classify("Gestor do processo X")
    assert kind == "placeholder"
    kind, _ = _classify("Diretor do Departamento X")
    assert kind == "placeholder"
    kind, _ = _classify("Chefe da Divisão X")
    assert kind == "placeholder"


def test_classify_trailing_isolated_x_does_not_match_normal_words() -> None:
    """``Mr X`` / ``raio X`` are too short to be POP-style placeholders,
    AND a trailing X surrounded by long prose isn't a marker either —
    only treat it as placeholder for short-ish role/title strings."""
    from engine.section_mapper.slot_profiler import _classify

    # Plain prose that happens to mention an isolated X — keep as data.
    kind, _ = _classify("O paciente fez raio X.")
    assert kind == "data"


# --- sdt-wrapped cells (Word content controls) ------------------------------


def test_profile_slots_finds_cells_inside_sdt_content_controls(
    tmp_path,  # type: ignore[no-untyped-def]
) -> None:
    """UNIFAP regression: enterprise templates wrap date / dropdown
    cells in ``<w:sdt><w:sdtContent><w:tc>``. python-docx's
    ``row.cells`` skips them, so the profiler used to be blind to
    template-default text like ``15/10/2014`` carried inside those
    cells. Walking ``<w:tc>`` descendants of the row brings them
    back."""
    import shutil

    from engine.section_mapper.slot_profiler import profile_slots

    base = _make_minimal_docx(tmp_path / "base.docx")
    target = tmp_path / "with_sdt.docx"
    shutil.copy(base, target)

    sdt_table_xml = """<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr>
      <w:tblGrid><w:gridCol w:w="2000"/><w:gridCol w:w="2000"/></w:tblGrid>
      <w:tr>
        <w:tc>
          <w:tcPr><w:tcW w:w="2000" w:type="dxa"/></w:tcPr>
          <w:p><w:r><w:t>Data da Revisao</w:t></w:r></w:p>
        </w:tc>
        <w:sdt>
          <w:sdtPr><w:date><w:dateFormat w:val="dd/MM/yyyy"/></w:date></w:sdtPr>
          <w:sdtContent>
            <w:tc>
              <w:tcPr><w:tcW w:w="2000" w:type="dxa"/></w:tcPr>
              <w:p><w:r><w:t>15/10/2014</w:t></w:r></w:p>
            </w:tc>
          </w:sdtContent>
        </w:sdt>
      </w:tr>
    </w:tbl>"""

    _inject_table_into_docx(target, sdt_table_xml)

    inv = profile_slots(target)
    cell_slots = [s for s in inv.slots if s.address.location == "table_cell"]

    # Must see TWO cells in the row (the normal one + the sdt-wrapped one),
    # not just one.
    assert len(cell_slots) == 2, (
        f"expected 2 cells, got {len(cell_slots)}: {[s.current_text for s in cell_slots]}"
    )
    texts = sorted(s.current_text.strip() for s in cell_slots)
    assert texts == ["15/10/2014", "Data da Revisao"]


def _make_minimal_docx(path):  # type: ignore[no-untyped-def]
    from docx import Document

    doc = Document()
    doc.add_paragraph("placeholder")
    doc.save(str(path))
    return path


def _inject_table_into_docx(path, table_xml: str) -> None:  # type: ignore[no-untyped-def]
    """Replace document.xml's body with a body that contains only this table."""
    import shutil
    import tempfile
    import zipfile

    body_xml = f"""<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {table_xml}
    <w:p><w:r><w:t>end</w:t></w:r></w:p>
  </w:body>
</w:document>"""

    fd, tmp = tempfile.mkstemp(suffix=".docx", dir=str(path.parent))
    import os

    os.close(fd)
    tmp_path = type(path)(tmp)
    try:
        with (
            zipfile.ZipFile(str(path), "r") as zin,
            zipfile.ZipFile(str(tmp_path), "w", zipfile.ZIP_DEFLATED) as zout,
        ):
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = body_xml.encode("utf-8")
                zout.writestr(item, data)
        shutil.move(str(tmp_path), str(path))
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


# --- table cell context anchors column header --------------------------------


def test_profile_slots_table_cell_context_includes_column_header(
    tmp_path,  # type: ignore[no-untyped-def]
) -> None:
    """UNIFAP regression: two empty cells in the same row must get
    DISTINCT contexts so the LLM can tell which column it's filling."""
    from docx import Document

    from engine.section_mapper.slot_profiler import profile_slots

    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "Nº"
    table.rows[0].cells[1].text = "Nome"
    table.rows[0].cells[2].text = "Telefone"
    table.rows[0].cells[3].text = "e-mail"
    table.rows[1].cells[0].text = "1"
    # row 1, columns 1-3 left empty (fillable slots)

    template_path = tmp_path / "two_col_table.docx"
    doc.save(str(template_path))

    inv = profile_slots(template_path)
    by_id = {s.id: s for s in inv.slots}

    # Header row cells must NOT carry a column anchor (they ARE the
    # header — a self-reference would just repeat their text). They
    # still see their row siblings.
    assert not by_id["cell_t0_r0_c0"].context.startswith("column=")

    # Body row empty cells must carry distinct column anchors.
    assert by_id["cell_t0_r1_c1"].context.startswith('column="Nome"')
    assert by_id["cell_t0_r1_c2"].context.startswith('column="Telefone"')
    assert by_id["cell_t0_r1_c3"].context.startswith('column="e-mail"')

    # Row siblings still appear after the column anchor.
    assert "1" in by_id["cell_t0_r1_c1"].context
