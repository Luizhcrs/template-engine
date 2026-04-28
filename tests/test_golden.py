"""Golden fixture suite for the schema-driven table fill pipeline.

Locks expected cell-by-cell output against a known-good run. The
suite mocks out :func:`engine.section_mapper.record_extractor.extract_records`
with canned records so the test is deterministic, fast, and runs in
CI without API keys / spend.

Adding a new golden fixture: drop ``<name>_canned_records.json`` and
``<name>_expected.json`` next to the existing UNIFAP files, then add
a parametrised entry to :func:`_GOLDEN_CASES`.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import TYPE_CHECKING

import pytest

from engine.section_mapper.record_aligner import (
    TableInventory,
    TableRow,
    align_records_to_rows,
)
from engine.section_mapper.record_extractor import Record
from engine.section_mapper.schemas.builtins import (
    CONTACT_LIST_SCHEMA,
    PARTICIPANT_TABLE_SCHEMA,
    REVISION_TABLE_SCHEMA,
    SIGNATURE_BOX_SCHEMA,
)
from engine.section_mapper.slot_profiler import _is_vmerge_continuation
from engine.section_mapper.typed_fill import TypedFillRequest, apply_typed_fills

if TYPE_CHECKING:
    from engine.section_mapper.schemas.types import TableSchema


_GOLDEN_DIR = Path(__file__).parent / "golden"

_SCHEMA_BY_NAME: dict[str, TableSchema] = {
    "contact_list": CONTACT_LIST_SCHEMA,
    "participant_table": PARTICIPANT_TABLE_SCHEMA,
    "revision_table": REVISION_TABLE_SCHEMA,
    "signature_box": SIGNATURE_BOX_SCHEMA,
}


_GOLDEN_CASES = [
    {
        "name": "unifap",
        "template": Path("tests/real_world/pop_unifap.docx"),
        "canned": _GOLDEN_DIR / "unifap_canned_records.json",
        "expected": _GOLDEN_DIR / "unifap_expected.json",
    },
]


def _build_inventory_from_template(
    template_path: Path,
    table_index: int,
    schema: TableSchema,
) -> TableInventory:
    """Recreate the same TableInventory the orchestrator builds from a
    template — visual-column-aligned cell_fillable, vmerge continuation
    flagged, body rows treated as fillable for schema-driven writes."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(template_path))
    table = doc.tables[table_index]
    tr_elements = table._tbl.findall(qn("w:tr"))

    rows: list[TableRow] = []
    for ri, tr in enumerate(tr_elements):
        all_tcs = list(tr.iter(qn("w:tc")))
        first_tc = tr.find(qn("w:tc"))
        row_is_continuation = ri > 0 and first_tc is not None and _is_vmerge_continuation(first_tc)
        fillable: list[bool] = []
        is_body_row = ri > 0
        for tc in all_tcs:
            if _is_vmerge_continuation(tc):
                fillable.append(False)
                continue
            fillable.append(is_body_row)
        fillable = fillable[: len(schema.columns)]
        rows.append(TableRow(vmerge_with_above=row_is_continuation, cell_fillable=fillable))

    return TableInventory(schema=schema, rows=rows)


@pytest.mark.parametrize("case", _GOLDEN_CASES, ids=[c["name"] for c in _GOLDEN_CASES])
def test_golden_pipeline_produces_expected_cells(case: dict, tmp_path: Path) -> None:
    """Run the deterministic part of the schema-driven pipeline
    against canned records and assert the output docx matches every
    cell in the expected fixture."""
    template = case["template"]
    if not template.exists():
        pytest.skip(f"template fixture missing: {template}")

    canned = json.loads(case["canned"].read_text(encoding="utf-8"))
    expected = json.loads(case["expected"].read_text(encoding="utf-8"))

    output = tmp_path / f"{case['name']}_output.docx"

    requests: list[TypedFillRequest] = []
    for ti_str, table_data in canned["tables"].items():
        ti = int(ti_str)
        schema = _SCHEMA_BY_NAME[table_data["schema"]]
        records = [Record(r) for r in table_data["records"]]
        inv = _build_inventory_from_template(template, ti, schema)
        cell_fills = align_records_to_rows(inv, records)
        if not cell_fills:
            continue
        requests.append(
            TypedFillRequest(
                table_index=ti,
                schema=schema,
                cell_fills=cell_fills,
            )
        )

    n = apply_typed_fills(template, output, requests)
    assert n > 0, "schema-driven pipeline wrote zero cells"

    # Read back every cell asserted in the expected map.
    from docx import Document
    from docx.oxml.ns import qn

    out_doc = Document(str(output))

    failures: list[str] = []
    for table_key, cell_map in expected.items():
        if not table_key.startswith("table_"):
            continue
        ti = int(table_key.split("_")[1])
        if ti >= len(out_doc.tables):
            failures.append(f"{table_key}: missing in output")
            continue
        rows = out_doc.tables[ti]._tbl.findall(qn("w:tr"))
        for cell_key, expected_text in cell_map.items():
            parts = cell_key.split("_")
            ri = int(parts[1])
            ci = int(parts[3])
            if ri >= len(rows):
                failures.append(f"{table_key}/{cell_key}: row missing")
                continue
            tcs = list(rows[ri].iter(qn("w:tc")))
            if ci >= len(tcs):
                failures.append(f"{table_key}/{cell_key}: col missing")
                continue
            wp = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            actual = "".join(t.text or "" for t in tcs[ci].iter(f"{wp}t")).strip()
            if actual != expected_text:
                failures.append(f"{table_key}/{cell_key}: expected={expected_text!r} actual={actual!r}")

    assert not failures, "Golden mismatch:\n" + "\n".join(failures)
