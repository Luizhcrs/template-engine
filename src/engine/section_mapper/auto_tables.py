"""Auto-detect empty tables in a template and synthesize sensible defaults.

Industrial templates ship a few canonical empty tables that 90% of users
will fill the same way:

- ``Histórico de Revisões`` (``Rev. | Data | Alteração``) — first row
  defaults to ``"00"`` / today / ``"Emissão inicial"``. When the source
  ``.docx`` has its own revision history (e.g. ``Versão | Data | Autor |
  Alterações``), :func:`detect_default_specs_with_source` extracts those
  rows and adds a final ``"Migração para o novo modelo padrão"`` entry.
- ``Atribuições e Responsabilidades`` (``Atividades | Responsabilidade``)
  — populated from the source's ``RESPONSABILIDADES /  AUTORIDADES``
  paragraphs when available (``Compete à gerência``: gerência column gets
  ``X``; ``Compete aos supervisores``: supervisores column gets ``X``).

The orchestrator wires this in when the caller doesn't pass ``table_specs``,
so a brand-new template-source pair gets sensible output without manual
configuration.
"""

from __future__ import annotations

import re
import unicodedata
from datetime import UTC, date, datetime
from typing import TYPE_CHECKING

from engine.section_mapper.table_filler import TableSpec

if TYPE_CHECKING:
    from pathlib import Path


def _normalize(text: str) -> str:
    nkfd = unicodedata.normalize("NFKD", text)
    no_accent = "".join(c for c in nkfd if not unicodedata.combining(c))
    return re.sub(r"[^A-Z0-9 ]+", " ", no_accent.upper()).strip()


def _today_iso() -> str:
    return datetime.now(UTC).date().isoformat()


_HISTORICO_HEADERS: frozenset[str] = frozenset({"REV", "DATA", "ALTERACAO"})
_RESPONSABILIDADE_HEADERS: frozenset[str] = frozenset({"ATIVIDADES", "RESPONSABILIDADE"})

_VERSION_COL_TOKENS: frozenset[str] = frozenset({"VERSAO", "REV"})
_DATE_COL_TOKENS: frozenset[str] = frozenset({"DATA"})
_CHANGE_COL_TOKENS: frozenset[str] = frozenset({"ALTERACAO", "ALTERACOES"})
_AUTHOR_COL_TOKENS: frozenset[str] = frozenset({"AUTOR", "REVISOR"})


def detect_default_specs(template_path: Path) -> list[TableSpec]:
    """Walk template tables and synthesize a default :class:`TableSpec` per
    canonical empty table found. When *source_path* is provided, defaults
    are upgraded to source-derived rows where the source has matching data.
    """
    return _detect_specs_impl(template_path, source_path=None)


def detect_default_specs_with_source(
    template_path: Path,
    source_path: Path,
) -> list[TableSpec]:
    """Like :func:`detect_default_specs`, but populated from *source_path*.

    Histórico table is filled from the source's revision history (renumbered
    starting at ``00``) plus a trailing ``"Migração para o novo modelo
    padrão"`` row. Responsabilidade table is filled from source paragraphs
    under ``Compete à gerência`` / ``Compete aos supervisores`` (or
    equivalents).
    """
    return _detect_specs_impl(template_path, source_path=source_path)


def _detect_specs_impl(
    template_path: Path,
    *,
    source_path: Path | None,
) -> list[TableSpec]:
    from docx import Document

    doc = Document(str(template_path))
    specs: list[TableSpec] = []

    source_history = _extract_source_history(source_path) if source_path else []
    source_resp = _extract_source_responsibilities(source_path) if source_path else None

    for table in doc.tables:
        if not table.rows:
            continue
        header = {_normalize(c.text) for c in table.rows[0].cells if c.text.strip()}
        if not header:
            continue

        if header >= _HISTORICO_HEADERS:
            specs.append(_historico_spec(table, source_history))
            continue

        # Responsabilidade table: header set ⊇ {ATIVIDADES, RESPONSABILIDADE}
        if {"ATIVIDADES", "RESPONSABILIDADE"} <= header and source_resp:
            specs.append(_responsabilidade_spec(table, source_resp))

    return specs


def _historico_spec(
    table: object,
    source_history: list[dict[str, str]],
) -> TableSpec:
    headers = [c.text.strip() for c in table.rows[0].cells]  # type: ignore[attr-defined]

    if not source_history:
        return TableSpec(
            headers=headers,
            rows=[
                {
                    "Rev.": "00",
                    "Rev": "00",
                    "Data": _today_iso(),
                    "Alteração": "Emissão inicial",
                    "Alteracao": "Emissão inicial",
                }
            ],
        )

    rows: list[dict[str, str]] = []
    for i, h in enumerate(source_history):
        rows.append(
            {
                "Rev.": f"{i:02d}",
                "Rev": f"{i:02d}",
                "Data": h.get("date", _today_iso()),
                "Alteração": h.get("change", ""),
                "Alteracao": h.get("change", ""),
            }
        )
    next_idx = len(source_history)
    rows.append(
        {
            "Rev.": f"{next_idx:02d}",
            "Rev": f"{next_idx:02d}",
            "Data": _today_iso(),
            "Alteração": "Migração para o novo modelo padrão.",
            "Alteracao": "Migração para o novo modelo padrão.",
        }
    )
    return TableSpec(headers=headers, rows=rows)


def _responsabilidade_spec(
    table: object,
    source_resp: dict[str, list[str]],
) -> TableSpec:
    headers = [c.text.strip() for c in table.rows[0].cells]  # type: ignore[attr-defined]

    # The template's column-2 and column-3 sub-headers (in row 2) say
    # "Gerente Setorial" and "Supervisores" — use the same labels.
    gerente_label = "Gerente Setorial"
    supervisores_label = "Supervisores"

    rows: list[dict[str, str]] = []
    for activity in source_resp.get("gerencia", []):
        rows.append(
            {
                "Atividades": activity,
                gerente_label: "X",
                supervisores_label: "",
            }
        )
    for activity in source_resp.get("supervisores", []):
        rows.append(
            {
                "Atividades": activity,
                gerente_label: "",
                supervisores_label: "X",
            }
        )

    subheaders = [""] * len(headers)
    if len(headers) >= 3:
        subheaders[1] = gerente_label
        subheaders[2] = supervisores_label
    return TableSpec(headers=headers, rows=rows, subheaders=subheaders)


def _extract_source_history(source_path: Path) -> list[dict[str, str]]:
    """Find a revision-history table in the source ``.docx``.

    Returns rows shaped as ``{"date": "...", "change": "..."}``. Author
    is concatenated into change when present.
    """
    if source_path.suffix.lower() != ".docx":
        return []
    try:
        from docx import Document

        doc = Document(str(source_path))
    except Exception:
        return []

    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [_normalize(c.text) for c in table.rows[0].cells]
        col_map = _classify_history_columns(header_cells)
        if col_map is None:
            continue

        rows: list[dict[str, str]] = []
        for r in table.rows[1:]:
            cells = [c.text.strip() for c in r.cells]
            if not any(cells):
                continue
            date_val = cells[col_map["date"]] if "date" in col_map else ""
            change_val = cells[col_map["change"]] if "change" in col_map else ""
            author_val = cells[col_map["author"]] if "author" in col_map else ""
            full_change = change_val
            if author_val:
                full_change = f"{change_val} ({author_val})" if change_val else author_val
            rows.append({"date": date_val, "change": full_change})
        if rows:
            return rows

    return []


def _classify_history_columns(header_cells: list[str]) -> dict[str, int] | None:
    out: dict[str, int] = {}
    for i, h in enumerate(header_cells):
        if any(tok in h for tok in _VERSION_COL_TOKENS):
            out.setdefault("version", i)
        if any(tok in h for tok in _DATE_COL_TOKENS):
            out.setdefault("date", i)
        if any(tok in h for tok in _CHANGE_COL_TOKENS):
            out.setdefault("change", i)
        if any(tok in h for tok in _AUTHOR_COL_TOKENS):
            out.setdefault("author", i)
    if "date" in out and "change" in out:
        return out
    return None


_GERENCIA_RE = re.compile(r"COMPETE\s+A?\s*GER", re.IGNORECASE)
_SUPERVISORES_RE = re.compile(r"COMPETE\s+AOS?\s+SUPERV", re.IGNORECASE)


def _extract_source_responsibilities(
    source_path: Path,
) -> dict[str, list[str]] | None:
    """Walk source ``.docx`` paragraphs; under ``Compete à gerência`` / ``...
    aos supervisores`` collect each child paragraph as an activity.

    Uses ``<w:numPr>`` ilvl to detect bucket boundaries: a sub-heading
    paragraph at ilvl=1 (or any numbered ilvl <= current bucket's) opens
    the next bucket; an ilvl=0 paragraph leaves the responsibilities
    section entirely.
    """
    if source_path.suffix.lower() != ".docx":
        return None
    try:
        from docx import Document

        from engine.section_mapper.numbering import extract_num_pr
    except Exception:
        return None
    try:
        doc = Document(str(source_path))
    except Exception:
        return None

    out: dict[str, list[str]] = {"gerencia": [], "supervisores": []}
    bucket: str | None = None
    bucket_ilvl: int | None = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        norm = _normalize(text)

        np = extract_num_pr(para._p.xml)
        ilvl = np[1] if np is not None else None

        # ilvl=0 = top-level section change. Always closes any active bucket.
        if ilvl == 0:
            bucket = None
            bucket_ilvl = None
            continue

        # Sub-heading at the bucket's level (or shallower) opens / closes.
        if ilvl is not None and (bucket_ilvl is None or ilvl <= bucket_ilvl):
            if _GERENCIA_RE.search(norm):
                bucket = "gerencia"
                bucket_ilvl = ilvl
                continue
            if _SUPERVISORES_RE.search(norm):
                bucket = "supervisores"
                bucket_ilvl = ilvl
                continue
            # A different ilvl=1 subsection inside the same parent closes
            # the bucket without opening a new one.
            bucket = None
            bucket_ilvl = None
            continue

        if bucket is not None:
            out[bucket].append(text)

    if not out["gerencia"] and not out["supervisores"]:
        return None
    return out


def merge_specs(
    auto: list[TableSpec],
    user: list[TableSpec] | None,
) -> list[TableSpec]:
    """User-supplied specs win over auto-detected ones with the same headers."""
    if not user:
        return auto

    user_keys = [{_normalize(h) for h in sp.headers} for sp in user]
    out = list(user)
    for auto_spec in auto:
        auto_key = {_normalize(h) for h in auto_spec.headers}
        if auto_key in user_keys:
            continue
        out.append(auto_spec)
    return out


__all__ = [
    "detect_default_specs",
    "detect_default_specs_with_source",
    "merge_specs",
]


# Suppress unused import warning when called from the package
_ = date
