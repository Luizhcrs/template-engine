"""Type-validated, deterministic per-cell write.

Given a list of :class:`TypedFillRequest` (one per table covered by
the schema-driven path), open the template, write each
``(row, col) -> value`` pair into the matching docx table, validate
the value against the column's :class:`ColumnType` first, and reject
anything that doesn't match.

This is the sink of the schema-driven pipeline. Profiler + extractor
+ aligner produce typed cell fills; this module commits them.
"""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass, field
from typing import TYPE_CHECKING

import structlog

from engine.section_mapper.schemas.types import ColumnType

if TYPE_CHECKING:
    from pathlib import Path

    from engine.section_mapper.schemas.types import TableSchema


log = structlog.get_logger(__name__)


_VALIDATORS: dict[ColumnType, re.Pattern[str]] = {
    ColumnType.PHONE: re.compile(r"^\(?\d{2}\)?\s*\d{4,5}-?\d{4}$"),
    ColumnType.EMAIL: re.compile(r"^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$"),
    ColumnType.DATE: re.compile(r"^\d{2}/\d{2}/\d{4}$"),
    ColumnType.VERSION: re.compile(r"^\d+\.\d+(?:\.\d+)?$"),
    ColumnType.NUMBER: re.compile(r"^\d+$"),
}

# Negative patterns — values that should NEVER end up in a free-form
# string column (NAME / SECTOR / ROLE). Reject dates, phone numbers,
# emails and pure version strings landing where a person's name or a
# department name belongs.
_FREEFORM_REJECT_RE = re.compile(
    r"^(?:"
    r"\d{2}/\d{2}/\d{4}"  # DD/MM/YYYY
    r"|\d{4}-\d{2}-\d{2}"  # ISO date
    r"|\(?\d{2}\)?\s*\d{4,5}-?\d{4}"  # phone
    r"|\d+\.\d+(?:\.\d+)?"  # version (1.0, 1.0.0)
    r"|[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}"  # email
    r")$"
)


def validate_value_for_column(value: str, column_type: ColumnType) -> bool:
    """Return True when *value* matches the regex / format expected
    for *column_type*. Free-form types (``NAME``, ``SECTOR``,
    ``ROLE``, ``FREE``) accept any non-empty string but reject values
    that obviously belong to another column type (a date in a Nome
    column, a phone in a Setor column, etc).
    """
    if column_type in {ColumnType.NAME, ColumnType.SECTOR, ColumnType.ROLE}:
        if not value:
            return False
        return not _FREEFORM_REJECT_RE.match(value)
    if column_type == ColumnType.FREE:
        return bool(value)
    pattern = _VALIDATORS.get(column_type)
    if pattern is None:
        return True
    return bool(pattern.match(value))


@dataclass(frozen=True)
class TypedFillRequest:
    """One table's worth of typed cell fills.

    Attributes:
        table_index: 0-based index into ``Document.tables``.
        schema: the :class:`TableSchema` matched against this table —
            drives per-cell type validation.
        cell_fills: ``(row, col) -> value`` for every cell to write.
            Coordinates use python-docx's ``row.cells`` indexing for
            simple, non-merged tables. Merged-cell handling lives in
            :mod:`engine.section_mapper.slot_renderer` and is reused
            indirectly for tables with vmerge / sdt-wrapped cells.
    """

    table_index: int
    schema: TableSchema
    cell_fills: dict[tuple[int, int], str] = field(default_factory=dict)


def apply_typed_fills(
    template_path: Path,
    output_path: Path,
    requests: list[TypedFillRequest],
) -> int:
    """Open *template_path*, apply every :class:`TypedFillRequest` to
    its addressed table, save to *output_path*. Returns the count of
    cells actually written (after type validation drops invalid
    values)."""
    from docx import Document
    from docx.oxml.ns import qn

    from engine.section_mapper.slot_renderer import _set_tc_text

    if template_path != output_path:
        shutil.copy(str(template_path), str(output_path))
    doc = Document(str(output_path))

    written = 0
    for request in requests:
        if request.table_index < 0 or request.table_index >= len(doc.tables):
            log.info(
                "section_mapper.typed_fill.table_index_out_of_range",
                index=request.table_index,
                total=len(doc.tables),
            )
            continue
        table = doc.tables[request.table_index]
        rows = list(table._tbl.findall(qn("w:tr")))
        for (ri, ci), value in request.cell_fills.items():
            if ri < 0 or ri >= len(rows):
                continue
            if ci < 0 or ci >= len(request.schema.columns):
                continue
            col = request.schema.columns[ci]
            if not validate_value_for_column(value, col.type):
                log.info(
                    "section_mapper.typed_fill.value_rejected",
                    table=request.table_index,
                    row=ri,
                    col=ci,
                    column_name=col.name,
                    column_type=col.type.value,
                    value=value[:60],
                )
                continue
            target_tc = _column_aligned_tc(rows[ri], ci)
            if target_tc is None:
                continue
            _set_tc_text(target_tc, value)
            written += 1

    doc.save(str(output_path))
    log.info(
        "section_mapper.typed_fill.applied",
        requests=len(requests),
        cells_written=written,
    )
    return written


def _column_aligned_tc(tr, target_col: int):  # type: ignore[no-untyped-def]
    """Walk *tr*'s direct ``<w:tc>`` children in document order and
    return the tc that occupies VISUAL column *target_col*.

    Vertical-merge continuation cells (``<w:vMerge/>`` without
    ``w:val="restart"``) DO occupy a visual column — they're part of
    the merged group above — but writing to them is a no-op in Word.
    Return None for those so the caller skips the write rather than
    silently shifting subsequent columns.
    """
    from docx.oxml.ns import qn

    from engine.section_mapper.slot_profiler import _is_vmerge_continuation

    direct = list(tr.findall(qn("w:tc")))
    if direct and target_col < len(direct):
        tc = direct[target_col]
        return None if _is_vmerge_continuation(tc) else tc
    # Fall back to descendants (sdt-wrapped tcs live deeper).
    wp = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    descendants = list(tr.iter(f"{wp}tc"))
    if target_col < len(descendants):
        tc = descendants[target_col]
        return None if _is_vmerge_continuation(tc) else tc
    return None


__all__ = [
    "TypedFillRequest",
    "apply_typed_fills",
    "validate_value_for_column",
]
