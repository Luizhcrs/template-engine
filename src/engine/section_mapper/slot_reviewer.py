"""Slot pipeline — closed-loop self-review pass.

After the first :func:`engine.section_mapper.slot_filler.fill_slots`
round and :func:`engine.section_mapper.slot_renderer.apply_slot_fills`
write the output docx, the LLM has never seen what it actually
produced. Bugs slip through silently:

- a name landing in the Telefone column instead of the e-mail column,
- ``Nº`` cells repeating ``1, 1, 2, 3`` instead of ``1, 2, 3, 4``,
- a ``Gestor do processo X`` placeholder left in place,
- duplicate rows where the LLM copied a template-default example.

This module renders the just-produced docx back to PNG pages, shows
them to the LLM alongside the original source content + the slot
inventory, and asks for corrections. Returns ``dict[slot_id,
corrected_text]`` to be re-applied. Slots the LLM does NOT mention are
considered fine and left alone.

One extra LLM call per doc. Skipped silently if the multimodal
renderer dependencies (``docx2pdf``, ``pymupdf``) are missing.
"""

from __future__ import annotations

import json
import re
from typing import TYPE_CHECKING

import structlog

if TYPE_CHECKING:
    from pathlib import Path

    from engine.llm.base import LLMProvider
    from engine.section_mapper.slots import SlotInventory
    from engine.section_mapper.source_profiler import SourceStructure


log = structlog.get_logger(__name__)


_PROMPT = """\
You are reviewing a docx that has already been filled by an earlier
pass. Your job is to spot mistakes the first pass made and propose
specific corrections.

DEFAULT IS NO CHANGE. Only propose a correction when you are
HIGHLY CONFIDENT the current text is wrong AND you can justify the
new text with a specific quote from SOURCE or a specific column
mismatch in the rendered PNG. If you are unsure, OMIT the slot.
Rewriting an already-correct cell is worse than leaving it alone.

You receive:

1. RENDERED PAGES — PNG renders of the OUTPUT docx (after first-pass
   fill). Read them as the user would.
2. SOURCE — the source document profile that drove the fill.
3. SLOTS — the slot inventory (each entry has ``id``, ``address``,
   ``current_text`` (post-fill text in the output), ``kind``,
   ``context``).

Each slot entry may carry an ``expected_column`` field. If present,
the ``current_text`` MUST semantically match that column header:

  expected_column="Telefone"   -> phone number ((96) 1234-5678)
  expected_column="e-mail"     -> e-mail address (foo@bar.br)
  expected_column="Data"       -> date (DD/MM/YYYY)
  expected_column="Nome"       -> person name
  expected_column="Setor"      -> department / unit name
  expected_column="Função"     -> role / job title
  expected_column="Versão"     -> version number
  expected_column="Nº" / "#"   -> distinct incrementing integer

If ``current_text`` does not match its ``expected_column``, propose a
correction. Look at the surrounding rows in the rendered PNG to find
the right value (a name shoved into a Telefone column has typically
displaced the real phone number — recover it from the page).

Other failure modes:

- Row-index columns (``Nº`` / ``#`` / ``Item``) carrying repeated
  values: must be distinct ascending integers (``1, 2, 3, 4`` not
  ``1, 1, 2, 3``). Re-number the whole column if you spot a repeat.
- Placeholder cells (``Gestor do processo X``,
  ``Diretor do Departamento X``, ``Chefe da Divisão XXXXX``) left
  unchanged when the source carries a real role / value.
- Duplicate rows where the first pass copied a template-default
  example (``Fulano de Tal | Reitoria | Secretário da Reitoria``)
  into an empty fill row instead of using real source data.
- Visible template tokens still showing in the output:
  ``{{X}}``, ``[FOO]``, ``<<X>>``, ``XXXXX``, ``XX/XX/2022``,
  ``Fulano (Titular)``, ``Ciclano (Substituto)``, ``Sicrano``, etc.
- Cells that received a value when they should have stayed empty
  (the LLM filled a logo cell with header text, etc).
- Same string repeated across two adjacent cells of the same row
  (LLM duplicated content instead of producing distinct values).

Output JSON ONLY, with one key:

```
{
  "corrections": [
    {"slot_id": "<exact slot id>", "new_text": "<corrected text>"}
  ]
}
```

Skip slots that are already correct — do NOT echo them. ``new_text``
must be the FULL replacement string, not a diff. Keep heading
prefixes (``1. OBJETIVO:``, ``Elaboração   Data:``) when the slot
kind is ``heading_with_hint`` / ``label_value``.

If everything looks correct, return ``{"corrections": []}``.

SOURCE structure:
{source_json}

SLOTS (post-fill):
{slots_json}

Output JSON only. No prose. No markdown.
"""


_CONTROL_CHAR_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
_COLUMN_HEADER_RE = re.compile(r'column="([^"]+)"')


def _clean(s: str) -> str:
    return _CONTROL_CHAR_RE.sub("", s)


def _extract_column_header(context: str) -> str:
    """Pull the ``column="<name>"`` anchor out of a slot's context
    string, if present. Returns ``""`` for body paragraphs / header
    cells / anything without a column anchor.
    """
    m = _COLUMN_HEADER_RE.search(context)
    return m.group(1) if m else ""


def _build_schema(inventory: SlotInventory) -> dict:
    fillable_ids = sorted({s.id for s in inventory.fillable()}) or ["__none__"]
    return {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "corrections": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "slot_id": {"type": "string", "enum": fillable_ids},
                        "new_text": {"type": "string"},
                    },
                    "required": ["slot_id", "new_text"],
                },
            }
        },
        "required": ["corrections"],
    }


async def review_slots(
    output_path: Path,
    inventory: SlotInventory,
    source: SourceStructure,
    *,
    llm: LLMProvider,
    output_image_urls: list[str] | None = None,
) -> dict[str, str]:
    """Issue ONE multimodal LLM call against the post-fill docx and
    return ``{slot_id: corrected_text}`` for the slots the LLM wants
    to fix. Slots the LLM does not mention are considered fine.

    *output_image_urls* may be passed pre-rendered (so the orchestrator
    can render once and reuse). If omitted, ``review_slots`` renders
    *output_path* on its own. Returns ``{}`` when no images can be
    produced (graceful no-op when docx2pdf/pymupdf are missing).
    """
    images = output_image_urls
    if images is None:
        try:
            from engine.section_mapper.template_renderer import render_pages

            pages = render_pages(output_path, max_pages=4)
            images = [p.as_data_url() for p in pages]
        except Exception as exc:
            log.info("section_mapper.slot_reviewer.render_skipped", error=str(exc))
            return {}

    if not images:
        log.info("section_mapper.slot_reviewer.no_images")
        return {}

    fillable = inventory.fillable()
    if not fillable:
        return {}

    # Read the post-fill cell texts so the LLM sees what's currently
    # in the output, not the template defaults.
    post_fill_texts = _read_current_texts(output_path, inventory)
    slots_payload = []
    for s in fillable:
        entry: dict[str, object] = {
            "id": s.id,
            "kind": s.kind,
            "context": s.context,
            "current_text": post_fill_texts.get(s.id, s.current_text),
        }
        # Lift the column name out of the context so the LLM does not
        # have to parse the ``column="..."`` prefix itself. Reviewer
        # was missing wrong-column mistakes because the anchor was
        # buried in the context string.
        col = _extract_column_header(s.context)
        if col:
            entry["expected_column"] = col
        slots_payload.append(entry)

    slots_json = json.dumps(slots_payload, ensure_ascii=False)
    source_json = json.dumps(source.to_dict(), ensure_ascii=False)

    prompt = _PROMPT.replace("{slots_json}", slots_json[:80000]).replace("{source_json}", source_json[:80000])

    schema = _build_schema(inventory)

    try:
        response = await llm.generate_structured(  # type: ignore[call-arg]
            prompt, schema, image_urls=images
        )
    except Exception as exc:
        log.warning("section_mapper.slot_reviewer.llm_failed", error=str(exc))
        return {}

    return _parse_response(response)


def _read_current_texts(
    output_path: Path,
    inventory: SlotInventory,
) -> dict[str, str]:
    """Re-read each fillable slot's text from the post-fill docx so
    the LLM reviews REAL output, not the template original.
    """
    from docx import Document
    from docx.oxml.ns import qn

    from engine.section_mapper.slot_profiler import _iter_row_tcs, _tc_text

    out: dict[str, str] = {}
    try:
        doc = Document(str(output_path))
    except Exception:
        return out

    fillable = inventory.fillable()
    for slot in fillable:
        loc = slot.address.location
        if loc == "table_cell":
            ti = slot.address.table_index
            ri = slot.address.row
            ci = slot.address.col
            if ti is None or ri is None or ci is None:
                continue
            if ti < 0 or ti >= len(doc.tables):
                continue
            rows = doc.tables[ti]._tbl.findall(qn("w:tr"))
            if ri < 0 or ri >= len(rows):
                continue
            tcs = _iter_row_tcs(rows[ri])
            if ci < 0 or ci >= len(tcs):
                continue
            out[slot.id] = _tc_text(tcs[ci]).strip()
        elif loc == "body_para":
            idx = slot.address.paragraph_idx
            if idx is None or idx < 0 or idx >= len(doc.paragraphs):
                continue
            out[slot.id] = doc.paragraphs[idx].text
    return out


# Common BR-PT template-default tokens. A reviewer correction
# proposing any of these is the LLM going BACKWARDS — replacing real
# data with a generic placeholder. Reject the correction silently.
_BANNED_NEW_TEXT_RE = re.compile(
    r"\b(?:Fulano(?:\s+de\s+Tal)?|Ciclano|Sicrano|Beltrano|XXXXX|XX/XX|x\.xx\.xxx\.xx)\b",
    re.IGNORECASE,
)


def _parse_response(response: object) -> dict[str, str]:
    if not isinstance(response, dict):
        return {}
    raw = response.get("corrections") or []
    if not isinstance(raw, list):
        return {}
    out: dict[str, str] = {}
    for entry in raw:
        if not isinstance(entry, dict):
            continue
        sid = entry.get("slot_id")
        text = entry.get("new_text")
        if not (isinstance(sid, str) and isinstance(text, str)):
            continue
        cleaned = _clean(text).strip()
        if not cleaned:
            continue
        if _BANNED_NEW_TEXT_RE.search(cleaned):
            log.info(
                "section_mapper.slot_reviewer.banned_correction_dropped",
                slot_id=sid,
                proposed=cleaned[:60],
            )
            continue
        out[sid] = cleaned
    return out


__all__ = ["review_slots"]
