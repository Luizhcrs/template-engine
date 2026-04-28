"""Generic template profiler — extracts structure from any ``.docx``
template without vendor-specific heuristics.

The rules engine pipeline relies on hardcoded rules (Engeman placeholder names,
Brazilian-PT headings, canonical Histórico / Responsabilidade tables).
That works for one vendor's templates; it does not generalise.

This module produces a vendor-agnostic :class:`TemplateStructure` that
the LLM-driven mapper (:mod:`engine.section_mapper.auto_mapper`) consumes
to build a substitution plan for ANY template + source pair.

Generic detection rules:

- **Placeholders** are detected by shape, not by name:

  - Repeated character runs: ``XXXX``, ``0000``, ``XX/XX/XX``.
  - Parenthesised labels: ``(TITULO)``, ``(NOME)``.
  - Square-bracket tokens: ``[FOO]``, ``[X]`` (also covers batch orchestrator
    ``{{X}}`` and ``___`` runs).
  - Empty-suffix labels: ``Label:`` followed by whitespace inside the
    same paragraph (``Elaborado:``, ``Aprovado:``, ``Data:``).
  - Revision-like literals: ``Rev. 00``, ``Versão 0``, ``Vers. 1``.

- **Empty tables** are any table with at least one fully blank row after
  the header row.

- **Headings** reuse :func:`engine.section_mapper.parser.parse_docx`
  unchanged.

The struct is JSON-serialisable so callers can stuff it into an LLM
prompt or persist a profile for later re-use.
"""

from __future__ import annotations

import re
from dataclasses import asdict, dataclass, field
from typing import TYPE_CHECKING

from engine.section_mapper.parser import DocxSection, parse_docx

if TYPE_CHECKING:
    from pathlib import Path


# --- placeholder shape regexes ------------------------------------------------


_REPEATED_CHARS_RE = re.compile(r"(?<![A-Za-z0-9])([X0]{3,}|[X0]{2,}(?:[/\-.][X0]{2,})+)(?![A-Za-z0-9])")
_PARENS_LABEL_RE = re.compile(r"\(([A-ZÁÉÍÓÚÂÊÔÃÕÇ_][A-ZÁÉÍÓÚÂÊÔÃÕÇ_ ]{1,40})\)")
_BRACKET_LABEL_RE = re.compile(r"\[([A-Za-zÀ-ÿ_][\w.-]{0,40})\]")
_CURLY_TOKEN_RE = re.compile(r"\{\{\s*([A-Za-z_][\w.-]*)\s*\}\}")
_DOUBLE_ANGLE_RE = re.compile(r"<<\s*([A-Za-z_][\w. -]*)\s*>>")
_ANGLE_LABEL_RE = re.compile(r"<\s*([a-zà-ÿ][a-zà-ÿ ]{1,30})\s*>")  # lowercase only to skip XML-like tags
_UNDERSCORE_RUN_RE = re.compile(r"_{3,}")
_DOT_LEADER_RE = re.compile(r"\.{6,}")
_SYMBOL_RUN_RE = re.compile(r"([§¶†‡])\1{2,}")
_LABEL_EMPTY_RE = re.compile(r"^([A-ZÁÉÍÓÚÂÊÔÃÕÇ][A-Za-zÀ-ÿ ]{1,30}):\s*$")
_LABEL_WITH_LEADER_RE = re.compile(
    r"^([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ /]{1,40}):\s*(?:\.{3,}|_{3,})\s*$",
)
_REVISION_RE = re.compile(r"\b(Rev\.?|Vers[aã]o|Ver\.)\s*0+\b", re.IGNORECASE)


@dataclass(frozen=True)
class TemplatePlaceholder:
    """A single detected placeholder.

    Attributes:
        text: literal placeholder text as it appears in the doc
            (``"XXXX"``, ``"(TITULO)"``, ``"Elaborado:"``, ...).
        kind: ``"repeated"`` / ``"parens"`` / ``"brackets"`` / ``"curly"``
            / ``"underscore"`` / ``"label_empty"`` / ``"revision"``.
        location: ``"header"`` / ``"footer"`` / ``"body"``.
        paragraph_idx: index in the doc's paragraph list (body only;
            ``-1`` for header / footer placeholders).
        context: surrounding text (full paragraph text) so the mapper
            understands the field's purpose.
    """

    text: str
    kind: str
    location: str
    paragraph_idx: int
    context: str

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass(frozen=True)
class TemplateEmptyTable:
    """A table in the template with empty data rows.

    Attributes:
        index: table index in ``doc.tables`` order.
        primary_headers: row 0 cell texts.
        sub_headers: row 1 cell texts (often empty; may carry
            sub-headers like ``["", "Gerente Setorial", "Supervisores"]``
            when row 0 has duplicate primary headers).
        empty_row_count: how many empty rows after the header(s).
        location: ``"body"`` always (table fill happens in body only).
    """

    index: int
    primary_headers: list[str]
    sub_headers: list[str]
    empty_row_count: int
    location: str = "body"

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass(frozen=True)
class TemplateCell:
    """A single cell in any template table.

    Captures cell-level layout so the LLM can fill ``mega-table``
    layouts (Corentocantins-style POPs where the entire document is
    one big table with embedded headings + body slots).

    Attributes:
        table_index: index in ``doc.tables`` order.
        row: 0-based row index inside the table.
        col: 0-based column index.
        text: current cell text (template default / placeholder /
            empty).
        is_fillable: heuristic flag — ``True`` when the cell looks
            like a slot the user is expected to fill (empty, contains
            imperative help text, contains ``XX`` / ``___`` masks,
            ends in ``:`` with no value, etc).
    """

    table_index: int
    row: int
    col: int
    text: str
    is_fillable: bool

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass(frozen=True)
class TemplateHeading:
    """A heading paragraph from the template.

    Attributes:
        name: canonical heading name (uppercase, no accents).
        raw_heading: original heading text.
        number: dotted section number or ``None`` (for unnumbered).
        level: heading depth.
        paragraph_idx: index in doc.paragraphs.
    """

    name: str
    raw_heading: str
    number: str | None
    level: int
    paragraph_idx: int

    @classmethod
    def from_section(cls, s: DocxSection) -> TemplateHeading:
        return cls(
            name=s.name,
            raw_heading=s.raw_heading,
            number=s.number,
            level=s.level,
            paragraph_idx=s.heading_paragraph_idx,
        )

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass(frozen=True)
class TemplateStructure:
    """Vendor-agnostic profile of a template.

    Attributes:
        template_path: source path (string for serialisation).
        headings: list of detected heading paragraphs.
        placeholders: list of detected placeholders (header/footer/body).
        empty_tables: list of empty tables awaiting data.
        cells: every cell of every body table — captures cell-level
            layout for mega-tables where heading + body live in
            adjacent cells (Corentocantins-style POPs).
    """

    template_path: str
    headings: list[TemplateHeading]
    placeholders: list[TemplatePlaceholder]
    empty_tables: list[TemplateEmptyTable]
    cells: list[TemplateCell] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "template_path": self.template_path,
            "headings": [h.to_dict() for h in self.headings],
            "placeholders": [p.to_dict() for p in self.placeholders],
            "empty_tables": [t.to_dict() for t in self.empty_tables],
            "cells": [c.to_dict() for c in self.cells],
        }


# --- detection ----------------------------------------------------------------


def profile_template(template_path: Path) -> TemplateStructure:
    """Walk *template_path* and return a :class:`TemplateStructure`."""
    from docx import Document

    doc = Document(str(template_path))

    headings = [TemplateHeading.from_section(s) for s in parse_docx(template_path)]
    placeholders = _detect_placeholders(doc, template_path)
    empty_tables = _detect_empty_tables(doc)
    cells = _detect_cells(doc)

    return TemplateStructure(
        template_path=str(template_path),
        headings=headings,
        placeholders=placeholders,
        empty_tables=empty_tables,
        cells=cells,
    )


def _detect_placeholders(doc, template_path: Path) -> list[TemplatePlaceholder]:  # type: ignore[no-untyped-def]
    """Detect placeholders in body, headers, footers."""
    out: list[TemplatePlaceholder] = []
    seen: set[tuple[str, str, int]] = set()  # (text, location, paragraph_idx)

    # body paragraphs
    for idx, para in enumerate(doc.paragraphs):
        for ph in _scan_paragraph_text(para.text, location="body", paragraph_idx=idx):
            key = (ph.text, ph.location, ph.paragraph_idx)
            if key in seen:
                continue
            seen.add(key)
            out.append(ph)

    # body tables (placeholders inside cells, like "(TITULO)" inside a header banner)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for ph in _scan_paragraph_text(cell.text, location="body", paragraph_idx=-1):
                    key = (ph.text, ph.location, -1)
                    if key in seen:
                        continue
                    seen.add(key)
                    out.append(ph)

    # headers + footers — scan raw XML so we also catch text inside text
    # boxes (<w:txbxContent>) and other non-paragraph containers that
    # python-docx's .paragraphs / .tables iterators skip.
    for hdr_kind, xml_files in (
        ("header", _list_part_xml(template_path, "header")),
        ("footer", _list_part_xml(template_path, "footer")),
    ):
        for xml in xml_files:
            for ph in _scan_xml_paragraphs(xml, location=hdr_kind):
                key = (ph.text, ph.location, ph.paragraph_idx)
                if key in seen:
                    continue
                seen.add(key)
                out.append(ph)

    return out


def _list_part_xml(docx_path: Path, prefix: str) -> list[str]:
    """Read every ``word/{prefix}*.xml`` from the docx zip."""
    import zipfile

    out: list[str] = []
    pattern = re.compile(rf"^word/{prefix}\d*\.xml$")
    try:
        with zipfile.ZipFile(str(docx_path)) as z:
            for name in z.namelist():
                if pattern.match(name):
                    out.append(z.read(name).decode("utf-8"))
    except (OSError, zipfile.BadZipFile):
        return []
    return out


def _scan_xml_paragraphs(xml: str, *, location: str) -> list[TemplatePlaceholder]:
    """Extract paragraph-level texts from a docx part XML and scan each
    for placeholders. Walks every ``<w:p>`` (including those nested
    inside text boxes, table cells, etc) regardless of container.
    """
    out: list[TemplatePlaceholder] = []
    para_pattern = re.compile(r"<w:p\b[^>]*>(.*?)</w:p>", re.DOTALL)
    text_pattern = re.compile(r"<w:t[^>]*>([^<]*)</w:t>")
    for p_match in para_pattern.finditer(xml):
        body = p_match.group(1)
        text = "".join(text_pattern.findall(body))
        if not text.strip():
            continue
        out.extend(_scan_paragraph_text(text, location=location, paragraph_idx=-1))
    return out


def _scan_container(container, *, location: str) -> list[TemplatePlaceholder]:  # type: ignore[no-untyped-def]
    out: list[TemplatePlaceholder] = []
    for para in container.paragraphs:
        out.extend(_scan_paragraph_text(para.text, location=location, paragraph_idx=-1))
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                out.extend(_scan_paragraph_text(cell.text, location=location, paragraph_idx=-1))
    return out


def _scan_paragraph_text(
    text: str,
    *,
    location: str,
    paragraph_idx: int,
) -> list[TemplatePlaceholder]:
    """Apply every placeholder-shape regex to *text* and return matches."""
    if not text or not text.strip():
        return []

    out: list[TemplatePlaceholder] = []

    for m in _REPEATED_CHARS_RE.finditer(text):
        out.append(
            TemplatePlaceholder(
                text=m.group(1),
                kind="repeated",
                location=location,
                paragraph_idx=paragraph_idx,
                context=text.strip(),
            )
        )
    for m in _PARENS_LABEL_RE.finditer(text):
        out.append(
            TemplatePlaceholder(
                text=m.group(1),
                kind="parens",
                location=location,
                paragraph_idx=paragraph_idx,
                context=text.strip(),
            )
        )
    for m in _BRACKET_LABEL_RE.finditer(text):
        out.append(
            TemplatePlaceholder(
                text=m.group(0),
                kind="brackets",
                location=location,
                paragraph_idx=paragraph_idx,
                context=text.strip(),
            )
        )
    for m in _CURLY_TOKEN_RE.finditer(text):
        out.append(
            TemplatePlaceholder(
                text=m.group(0),
                kind="curly",
                location=location,
                paragraph_idx=paragraph_idx,
                context=text.strip(),
            )
        )
    for m in _DOUBLE_ANGLE_RE.finditer(text):
        out.append(
            TemplatePlaceholder(
                text=m.group(0),
                kind="double_angle",
                location=location,
                paragraph_idx=paragraph_idx,
                context=text.strip(),
            )
        )
    for m in _ANGLE_LABEL_RE.finditer(text):
        out.append(
            TemplatePlaceholder(
                text=m.group(0),
                kind="angle",
                location=location,
                paragraph_idx=paragraph_idx,
                context=text.strip(),
            )
        )
    for m in _UNDERSCORE_RUN_RE.finditer(text):
        out.append(
            TemplatePlaceholder(
                text=m.group(0),
                kind="underscore",
                location=location,
                paragraph_idx=paragraph_idx,
                context=text.strip(),
            )
        )
    for m in _DOT_LEADER_RE.finditer(text):
        out.append(
            TemplatePlaceholder(
                text=m.group(0),
                kind="dot_leader",
                location=location,
                paragraph_idx=paragraph_idx,
                context=text.strip(),
            )
        )
    for m in _SYMBOL_RUN_RE.finditer(text):
        out.append(
            TemplatePlaceholder(
                text=m.group(0),
                kind="symbol_run",
                location=location,
                paragraph_idx=paragraph_idx,
                context=text.strip(),
            )
        )
    for m in _REVISION_RE.finditer(text):
        out.append(
            TemplatePlaceholder(
                text=m.group(0),
                kind="revision",
                location=location,
                paragraph_idx=paragraph_idx,
                context=text.strip(),
            )
        )
    label_match = _LABEL_EMPTY_RE.match(text.strip())
    if label_match:
        out.append(
            TemplatePlaceholder(
                text=label_match.group(0),
                kind="label_empty",
                location=location,
                paragraph_idx=paragraph_idx,
                context=text.strip(),
            )
        )

    return out


_IMPERATIVE_INSTRUCTION_RE = re.compile(
    r"^\s*(Descrever|Identificar|Listar|Citar|Apontar|Indicar|Informar|Mencionar|"
    r"Inserir|Detalhar|Especificar|Descreva|Cite|Liste|Aponte|"
    r"Describe|List|Identify|State|Specify|Explain)\b",
    re.IGNORECASE,
)
_XX_MASK_RE = re.compile(r"X{2,}|0{2,}|_{3,}|\.{6,}|/{2,}")
_PARENTHESIZED_HINT_RE = re.compile(r"^\s*\(.+\)\s*$")
_LABEL_NO_VALUE_RE = re.compile(r"^[\w\sÀ-ÿ/]{2,40}:\s*$")


def _is_cell_fillable(text: str) -> bool:
    """Heuristic: does this cell look like a fill-me slot?

    Empty, imperative instruction, mask shape, parenthesized hint,
    label-with-empty-value — all qualify.
    """
    stripped = text.strip()
    if not stripped:
        return True
    if _IMPERATIVE_INSTRUCTION_RE.match(stripped):
        return True
    if _XX_MASK_RE.search(stripped):
        return True
    if _PARENTHESIZED_HINT_RE.match(stripped) and len(stripped) <= 80:
        return True
    if _LABEL_NO_VALUE_RE.match(stripped):
        return True
    # Default-placeholder names commonly used by templates.
    return stripped in {
        "Fulano",
        "Fulano de Tal",
        "Fulano (Titular)",
        "Ciclano",
        "Ciclano (Substituto)",
        "Beltrano",
        "Beltrano de Tal",
    }


def _detect_cells(doc) -> list[TemplateCell]:  # type: ignore[no-untyped-def]
    """Return every cell of every body table, with fillability flag."""
    out: list[TemplateCell] = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip()
                out.append(
                    TemplateCell(
                        table_index=ti,
                        row=ri,
                        col=ci,
                        text=text,
                        is_fillable=_is_cell_fillable(text),
                    )
                )
    return out


def _detect_empty_tables(doc) -> list[TemplateEmptyTable]:  # type: ignore[no-untyped-def]
    """Find every table with at least one fully-empty data row."""
    out: list[TemplateEmptyTable] = []
    for i, table in enumerate(doc.tables):
        if not table.rows:
            continue
        primary = [c.text.strip() for c in table.rows[0].cells]
        sub: list[str] = []
        if len(table.rows) >= 2:
            sub = [c.text.strip() for c in table.rows[1].cells]

        # Count empty data rows. Skip first (header) + sub-header (when
        # present and clearly different from data row).
        sub_is_subheader = bool(sub) and any(s for s in sub) and sub != primary
        data_rows_start = 2 if sub_is_subheader else 1

        empty_count = 0
        for r in table.rows[data_rows_start:]:
            cells = [c.text.strip() for c in r.cells]
            if not any(cells):
                empty_count += 1

        if empty_count >= 1:
            out.append(
                TemplateEmptyTable(
                    index=i,
                    primary_headers=primary,
                    sub_headers=sub if sub_is_subheader else [""] * len(primary),
                    empty_row_count=empty_count,
                )
            )
    return out


__all__ = [
    "TemplateCell",
    "TemplateEmptyTable",
    "TemplateHeading",
    "TemplatePlaceholder",
    "TemplateStructure",
    "profile_template",
]


# Suppress unused-import warning for asdict re-export.
_ = field
