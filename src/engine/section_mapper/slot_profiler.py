"""Slot pipeline — flat profiler that emits :class:`Slot` records covering
every fillable place in a template.

Walks:

1. Body paragraphs.
2. Body table cells.
3. Header / footer paragraphs (via the docx zip's ``word/header*.xml``
   and ``word/footer*.xml`` parts — python-docx skips paragraphs nested
   inside text boxes).

Each emitted :class:`Slot` carries a stable ID, the address inside the
docx tree, the current text, a kind classification, and a fillability
flag.

The profiler does NOT decide what content goes where — that's the
LLM's job. It just produces a complete inventory.
"""

from __future__ import annotations

import re
import zipfile
from typing import TYPE_CHECKING

from engine.section_mapper.slots import Slot, SlotAddress, SlotInventory

if TYPE_CHECKING:
    from pathlib import Path

# --- fillability heuristics ---------------------------------------------------


_PLACEHOLDER_RE = re.compile(
    r"""(
        \{\{[^}]+\}\}                  # {{X}}
        | \[[A-Za-zÀ-ÿ_][^\]]{0,40}\]  # [FOO]
        | <<[^>]+>>                    # <<X>>
        | \([A-ZÁÉÍÓÚÂÊÔÃÕÇ_][A-ZÁÉÍÓÚÂÊÔÃÕÇ_ ]{2,40}\)  # (LABEL)
    )""",
    re.VERBOSE,
)

_LEADER_RE = re.compile(r"_{3,}|\.{6,}|/{2,}")
_XX_MASK_RE = re.compile(r"X{3,}|0{4,}")
# Matches strings ending with an isolated trailing ``X`` token used as
# a fill marker in BR-PT POP templates: ``Gestor do processo X``,
# ``Chefe da Divisão X``, ``Diretor do Departamento X``. The string
# must START with a known role token AND end with the bare ``X``,
# otherwise prose like ``O paciente fez raio X`` would match.
_ROLE_TRAILING_X_RE = re.compile(
    r"^(?:Diretor[a]?|Chefe|Gestor[a]?|Gerente|Coordenador[a]?|Supervisor[a]?"
    r"|Reitor[a]?|Pró-Reitor[a]?|Pró-Reitoria|Vice-Reitor[a]?|Secretári[oa]"
    r"|Procurador[a]?|Superintendente|Departamento|Divisão|Unidade|Setor)"
    r"[\w\s\-/áéíóúâêôãõçÁÉÍÓÚÂÊÔÃÕÇ]{0,55}\sX\s*$",
    re.IGNORECASE,
)
_LABEL_NO_VALUE_RE = re.compile(r"^[\w\sÀ-ÿ/]{2,40}:\s*$")
_LABEL_WITH_LEADER_RE = re.compile(r"^[\w\sÀ-ÿ/]{1,40}:\s*(?:_{3,}|\.{3,}|/+)\s*$")
_IMPERATIVE_RE = re.compile(
    r"^\s*(Descrever|Identificar|Listar|Citar|Apontar|Indicar|Informar|Mencionar|"
    r"Inserir|Detalhar|Especificar|Descreva|Cite|Liste|Aponte|Provê|Prover|"
    r"Describe|List|Identify|State|Specify|Explain)\b",
    re.IGNORECASE,
)
_PARENS_HINT_RE = re.compile(r"\([^)]{5,}\)")
_NUMBERED_HEADING_HINT_RE = re.compile(r"^\s*\d+(?:\.\d+)*[.:]?\s+[A-ZÁÉÍÓÚÂÊÔÃÕÇ]")


def _classify(text: str) -> tuple[str, bool]:
    """Return ``(kind, is_fillable)`` for a paragraph / cell text.

    Kinds:

    - ``"empty"`` — literally empty.
    - ``"placeholder"`` — contains a delimited placeholder token.
    - ``"label_value"`` — label-with-leader / label-no-value.
    - ``"instruction"`` — imperative-instruction help text.
    - ``"heading_with_hint"`` — numbered heading + parenthesised hint
      in the same cell (common in Corentocantins POPs: ``1. OBJETIVO:
      (Descrição clara...)``). Renderer must keep the heading prefix
      and replace only the hint.
    - ``"heading"`` — numbered heading with no fillable hint.
    - ``"data"`` — already filled data (preserve).

    Fillable kinds: empty / placeholder / label_value / instruction /
    heading_with_hint.
    """
    stripped = text.strip()
    if not stripped:
        return ("empty", True)

    has_numbered_heading = bool(_NUMBERED_HEADING_HINT_RE.match(stripped))
    has_parens_hint = bool(_PARENS_HINT_RE.search(stripped))

    if has_numbered_heading and has_parens_hint:
        return ("heading_with_hint", True)
    if _PLACEHOLDER_RE.search(stripped):
        return ("placeholder", True)
    if _LABEL_WITH_LEADER_RE.match(stripped) or _LABEL_NO_VALUE_RE.match(stripped):
        return ("label_value", True)
    if _IMPERATIVE_RE.match(stripped):
        return ("instruction", True)
    if _LEADER_RE.search(stripped) and len(stripped) <= 80:
        return ("placeholder", True)
    if _XX_MASK_RE.search(stripped):
        return ("placeholder", True)
    if _ROLE_TRAILING_X_RE.match(stripped):
        return ("placeholder", True)
    if has_numbered_heading:
        return ("heading", False)  # heading without fillable hint — preserve
    return ("data", False)


# --- profilers ----------------------------------------------------------------


def profile_slots(template_path: Path) -> SlotInventory:
    """Walk *template_path* and emit a complete :class:`SlotInventory`.

    Order: body paragraphs → body table cells → header parts → footer
    parts. The LLM gets this inventory in document order so it can
    reason about flow.
    """
    from docx import Document

    doc = Document(str(template_path))

    out: list[Slot] = []

    # Body paragraphs. Only flag empty paragraphs as fillable when they
    # sit DIRECTLY UNDER a heading paragraph — otherwise the LLM treats
    # every empty paragraph (including spacer / separator paragraphs
    # above tables) as a slot and fills them with arbitrary content.
    fillable_empty_idxs = _empty_idxs_under_headings(doc.paragraphs)
    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        kind, fillable = _classify(text)
        if kind == "empty" and idx not in fillable_empty_idxs:
            fillable = False
        out.append(
            Slot(
                id=f"body_para_{idx}",
                address=SlotAddress(location="body_para", paragraph_idx=idx),
                current_text=text,
                kind=kind,
                context=_neighbouring_paragraph(doc.paragraphs, idx),
                is_fillable=fillable,
            )
        )

    # Body table cells. python-docx's ``row.cells`` returns N entries
    # even when columns are MERGED (same underlying ``<w:tc>`` repeats).
    # Dedupe by the underlying XML element identity so the LLM sees one
    # logical slot per merged group instead of 8 of the same cell.
    for ti, table in enumerate(doc.tables):
        header_texts = _row_tc_texts(_iter_row_tcs(table.rows[0]._tr)) if table.rows else []
        for ri, row in enumerate(table.rows):
            row_tcs = _iter_row_tcs(row._tr)
            row_texts = _row_tc_texts(row_tcs)
            for ci, tc in enumerate(row_tcs):
                text = _tc_text(tc)
                if _tc_has_drawing(tc):
                    kind, fillable = "data", False
                else:
                    kind, fillable = _classify(text)
                out.append(
                    Slot(
                        id=f"cell_t{ti}_r{ri}_c{ci}",
                        address=SlotAddress(
                            location="table_cell",
                            table_index=ti,
                            row=ri,
                            col=ci,
                        ),
                        current_text=text,
                        kind=kind,
                        context=_cell_context_xml(row_texts, ci, header_texts, ri),
                        is_fillable=fillable,
                    )
                )

    # Header / footer parts via raw XML
    for part_name, slot_kind in (("header", "header_para"), ("footer", "footer_para")):
        out.extend(_profile_part_slots(template_path, part_name, slot_kind))

    return SlotInventory(template_path=str(template_path), slots=out)


_MAX_EMPTY_RUN_AS_SLOT = 2


def _empty_idxs_under_headings(paragraphs: list) -> set[int]:  # type: ignore[type-arg]
    """Return the indices of empty paragraphs that DIRECTLY follow a
    heading paragraph AND belong to a short run.

    Rule: when a heading is found, look at the run of consecutive empty
    paragraphs that follow it. If the run length is ≤
    :data:`_MAX_EMPTY_RUN_AS_SLOT`, every empty in the run is a fillable
    body slot. If the run is longer, the empties are page-layout
    padding (typical before a mega-table or page break) and NONE of
    them are flagged.

    Why a cap: the Corentocantins POP regression. Its template puts 19
    empty paragraphs between two title lines purely for page padding —
    the previous "any empty after a heading is a slot" rule flagged all
    19, causing the LLM to either fill them with header text or skip
    them, exploding false-positive slot count from ~16 to 86.
    """
    out: set[int] = set()
    n = len(paragraphs)
    i = 0
    while i < n:
        text = paragraphs[i].text.strip()
        if not text or not _looks_like_heading(text):
            i += 1
            continue
        run_start = i + 1
        j = run_start
        while j < n and not paragraphs[j].text.strip():
            j += 1
        run_len = j - run_start
        if 0 < run_len <= _MAX_EMPTY_RUN_AS_SLOT:
            out.update(range(run_start, j))
        i = max(j, i + 1)
    return out


_HEADING_LIKE_RE = re.compile(
    r"^\s*(?:\d+(?:\.\d+)*[.:]?\s+\S|[A-ZÁÉÍÓÚÂÊÔÃÕÇ][A-ZÁÉÍÓÚÂÊÔÃÕÇ\s\-—–:]{3,80})\s*$",
)


def _looks_like_heading(text: str) -> bool:
    """Heuristic: numbered (``1. OBJETIVO``) or all-caps multi-word
    (``OBJETIVO``, ``Descrição``-style Title-case headings handled by
    the bold-aware detector elsewhere).
    """
    if not _HEADING_LIKE_RE.match(text):
        return False
    if text.endswith("."):
        # Body sentences ending in period are NOT headings.
        return False
    return len(text) <= 80


def _neighbouring_paragraph(paragraphs: list, idx: int) -> str:  # type: ignore[type-arg]
    """Return the nearest non-empty paragraph above *idx* as context.

    Helps the LLM understand which heading the slot lives under.
    """
    for j in range(idx - 1, -1, -1):
        text = paragraphs[j].text.strip()
        if text:
            return text[:120]
    return ""


_WP_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_W_TC = f"{_WP_NS}tc"
_W_T = f"{_WP_NS}t"
_W_TCPR = f"{_WP_NS}tcPr"
_W_VMERGE = f"{_WP_NS}vMerge"
_W_VAL = f"{_WP_NS}val"
_W_DRAWING = f"{_WP_NS}drawing"


def _iter_row_tcs(tr) -> list:  # type: ignore[no-untyped-def, type-arg]
    """Return every ``<w:tc>`` descendant of *tr* in document order,
    excluding vertical-merge CONTINUATION cells.

    Why descendants and not direct children: enterprise templates wrap
    date / dropdown / plain-text cells inside
    ``<w:sdt><w:sdtContent><w:tc>`` (Word "Content Controls"), which
    python-docx's ``row.cells`` skips.

    Why skip continuation cells: a vertically-merged group like a logo
    image stretching across 3 rows shows up as one tc with
    ``<w:vMerge w:val="restart"/>`` plus two tcs with ``<w:vMerge/>``
    (no val = "continue"). Treating each one as its own slot would
    flood the LLM with phantom empty cells.

    Output is order-preserving and deduplicated on lxml element
    identity so a horizontally-merged tc that appears once in the XML
    stays once in the slot inventory.
    """
    seen: set[int] = set()
    out = []
    for tc in tr.iter(_W_TC):
        if id(tc) in seen:
            continue
        seen.add(id(tc))
        if _is_vmerge_continuation(tc):
            continue
        out.append(tc)
    return out


def _is_vmerge_continuation(tc) -> bool:  # type: ignore[no-untyped-def]
    tcPr = tc.find(_W_TCPR)
    if tcPr is None:
        return False
    vmerge = tcPr.find(_W_VMERGE)
    if vmerge is None:
        return False
    val = vmerge.get(_W_VAL)
    # No val OR explicit "continue" both mean continuation; "restart"
    # means this is the top of the merged group and IS its own slot.
    return val != "restart"


def _tc_text(tc) -> str:  # type: ignore[no-untyped-def]
    """Concatenate every ``<w:t>`` text inside *tc* — same shape as
    python-docx's ``Cell.text`` so ``_classify`` keeps working.
    """
    return "".join((t.text or "") for t in tc.iter(_W_T))


def _tc_has_drawing(tc) -> bool:  # type: ignore[no-untyped-def]
    """True if the cell contains a ``<w:drawing>`` (logo, picture).
    Such cells are visual content and must NOT be flagged fillable.
    """
    return next(tc.iter(_W_DRAWING), None) is not None


def _row_tc_texts(tcs) -> list[str]:  # type: ignore[no-untyped-def, type-arg]
    return [_tc_text(tc).strip()[:60] for tc in tcs]


def _cell_context_xml(
    row_texts: list[str],
    current_col: int,
    header_texts: list[str],
    current_row: int,
) -> str:
    """Build the per-cell context string sent to the LLM.

    Layout: ``column="<header>" | siblings: <row siblings joined>``.
    Header omitted for row 0 (the header row itself). Row siblings
    preserve the historical mega-table behaviour (LLM sees the heading
    cell while deciding the body slot in the same row).
    """
    parts: list[str] = []
    if current_row != 0 and current_col < len(header_texts):
        col_header = header_texts[current_col]
        if col_header:
            parts.append(f'column="{col_header}"')
    for ci, text in enumerate(row_texts):
        if ci == current_col:
            continue
        if text:
            parts.append(text)
    return " | ".join(parts)[:200]


def _profile_part_slots(
    template_path: Path,
    prefix: str,
    slot_location: str,
) -> list[Slot]:
    """Walk every ``word/{prefix}*.xml`` part and emit one slot per
    paragraph that carries text. We use raw XML so paragraphs inside
    text boxes (``<w:txbxContent>``) get profiled too.
    """
    out: list[Slot] = []
    pattern = re.compile(rf"^word/{prefix}\d*\.xml$")
    para_pattern = re.compile(r"<w:p\b[^>]*>(.*?)</w:p>", re.DOTALL)
    text_pattern = re.compile(r"<w:t[^>]*>([^<]*)</w:t>")

    try:
        with zipfile.ZipFile(str(template_path)) as z:
            for name in z.namelist():
                if not pattern.match(name):
                    continue
                xml = z.read(name).decode("utf-8")
                for p_idx, p_match in enumerate(para_pattern.finditer(xml)):
                    body = p_match.group(1)
                    text = "".join(text_pattern.findall(body))
                    if not text.strip():
                        continue
                    kind, fillable = _classify(text)
                    short_part = name.replace("/", "_").replace(".xml", "")
                    out.append(
                        Slot(
                            id=f"{slot_location}_{short_part}_p{p_idx}",
                            address=SlotAddress(
                                location=slot_location,
                                part_name=name,
                                paragraph_idx=p_idx,
                            ),
                            current_text=text,
                            kind=kind,
                            context="",
                            is_fillable=fillable,
                        )
                    )
    except (OSError, zipfile.BadZipFile):
        return []

    return out


__all__ = [
    "profile_slots",
]
