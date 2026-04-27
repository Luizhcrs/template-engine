"""Heading parser — detects sections in docx + plain text.

Strategy: a heading is a paragraph that either:

1. has a Word ``Heading <N>`` style; or
2. matches the numbered-heading pattern (``"1. OBJETIVO"``,
   ``"3.2. Etapas..."``, etc) — the dominant convention in industrial
   procedures (Engeman, NR-12, NR-13, ISO 9001).

Each :class:`Section` owns the paragraphs (or the trailing run of source
text) between its own heading and the next heading at the same-or-higher
level.

The parser handles two inputs:

- ``parse_docx(path)`` — walks ``python-docx`` paragraphs, returns
  :class:`DocxSection` with paragraph **indices** so callers can edit the
  doc in place.
- ``parse_text(text)`` — walks plain text (extracted from PDF/docx),
  returns :class:`TextSection` with the captured **content string** for
  each heading.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from pathlib import Path


_NUMBERED_HEADING_RE = re.compile(
    r"^\s*(\d+(?:\.\d+)*)\.?\s+([A-ZÁÉÍÓÚÂÊÔÃÕÇ][A-Za-zÀ-ÿ\s\-/]{2,80}?)\s*$",
)
# All-caps unnumbered heading (template style): "OBJETIVO", "APLICAÇÃO",
# "NORMAS E DOCUMENTOS DE REFERÊNCIA". Allows letters / spaces / hyphens.
# Length 3-80; rejects single-word lowercase or sentence-cased text.
_ALLCAPS_HEADING_RE = re.compile(
    r"^\s*([A-ZÁÉÍÓÚÂÊÔÃÕÇ][A-ZÁÉÍÓÚÂÊÔÃÕÇ\s\-/]{2,80})\s*$",
)
_HEADING_STYLE_RE = re.compile(r"^Heading\s*(\d+)\s*$", re.IGNORECASE)


def _normalize_heading(text: str) -> str:
    """Uppercase + strip accents/punctuation for matching keys."""
    import unicodedata

    nkfd = unicodedata.normalize("NFKD", text)
    no_accent = "".join(c for c in nkfd if not unicodedata.combining(c))
    cleaned = re.sub(r"[^A-Za-z0-9 ]+", " ", no_accent).strip().upper()
    return re.sub(r"\s+", " ", cleaned)


@dataclass(frozen=True)
class TextSection:
    """A section parsed from plain text.

    Attributes:
        name: canonical heading text (uppercase, no accents).
        raw_heading: original heading line as it appears in the doc.
        number: dotted section number (``"3.2"``) or ``None``.
        level: depth (1 for top-level, 2 for ``3.2``, ...).
        content: text between this heading and the next same-or-higher.
    """

    name: str
    raw_heading: str
    number: str | None
    level: int
    content: str

    @property
    def is_empty(self) -> bool:
        return not self.content.strip()


@dataclass(frozen=True)
class DocxSection:
    """A section parsed from a python-docx Document.

    Attributes:
        name: canonical heading text.
        raw_heading: original heading line.
        number: dotted section number or None.
        level: depth.
        heading_paragraph_idx: index of the heading paragraph in
            ``doc.paragraphs``.
        content_paragraph_idxs: indices of the body paragraphs (between
            this heading and the next).
    """

    name: str
    raw_heading: str
    number: str | None
    level: int
    heading_paragraph_idx: int
    content_paragraph_idxs: list[int] = field(default_factory=list)


def _detect_heading(paragraph_text: str, style_name: str | None) -> tuple[str, str | None, int] | None:
    """Return (canonical_name, number, level) or None when not a heading."""
    text = paragraph_text.strip()
    if not text:
        return None

    # Word style first
    if style_name:
        m = _HEADING_STYLE_RE.match(style_name)
        if m:
            return _normalize_heading(text), None, int(m.group(1))

    # Numbered heading pattern
    m = _NUMBERED_HEADING_RE.match(text)
    if m:
        number = m.group(1)
        title = m.group(2)
        level = number.count(".") + 1
        return _normalize_heading(title), number, level

    # All-caps unnumbered heading — template style ("OBJETIVO", "APLICACAO").
    # Heuristic: 3-80 chars, ALL upper, no inner punctuation that screams
    # body-text. We refuse lines that contain a colon (label syntax) or
    # end with a digit (form fields) to lower false-positive rate.
    m = _ALLCAPS_HEADING_RE.match(text)
    if m and ":" not in text and not text.rstrip().endswith(tuple("0123456789")):
        return _normalize_heading(text), None, 1

    return None


def parse_text(text: str) -> list[TextSection]:
    """Parse plain text into a flat list of :class:`TextSection`.

    Algorithm:

    1. Walk lines.
    2. When a line matches the heading pattern, close the previous
       section's content and open a new one.
    3. Anything between two headings is the previous section's content.
    """
    lines = text.splitlines()
    sections: list[TextSection] = []
    current: dict | None = None
    buffer: list[str] = []

    def _flush() -> None:
        nonlocal current, buffer
        if current is None:
            return
        sections.append(
            TextSection(
                name=current["name"],
                raw_heading=current["raw_heading"],
                number=current["number"],
                level=current["level"],
                content="\n".join(buffer).strip(),
            )
        )
        buffer = []

    for line in lines:
        head = _detect_heading(line, style_name=None)
        if head is None:
            if current is not None:
                buffer.append(line)
            continue

        canonical, number, level = head
        _flush()
        current = {
            "name": canonical,
            "raw_heading": line.strip(),
            "number": number,
            "level": level,
        }

    _flush()
    return sections


def parse_docx(path: Path) -> list[DocxSection]:
    """Parse a ``.docx`` into :class:`DocxSection` records keyed by heading."""
    from docx import Document

    doc = Document(str(path))
    sections: list[DocxSection] = []
    current: dict | None = None

    def _flush() -> None:
        nonlocal current
        if current is None:
            return
        sections.append(
            DocxSection(
                name=current["name"],
                raw_heading=current["raw_heading"],
                number=current["number"],
                level=current["level"],
                heading_paragraph_idx=current["heading_idx"],
                content_paragraph_idxs=list(current["content_idxs"]),
            )
        )

    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else None
        head = _detect_heading(para.text, style_name)
        if head is None:
            if current is not None:
                current["content_idxs"].append(idx)
            continue

        canonical, number, level = head
        _flush()
        current = {
            "name": canonical,
            "raw_heading": para.text.strip(),
            "number": number,
            "level": level,
            "heading_idx": idx,
            "content_idxs": [],
        }

    _flush()
    return sections


def heading_index(sections: list) -> dict[str, int]:
    """Map canonical heading name -> first-occurrence index in *sections*.

    Accepts either ``list[TextSection]`` or ``list[DocxSection]``; both expose
    a ``.name`` attribute.
    """
    out: dict[str, int] = {}
    for i, s in enumerate(sections):
        if s.name not in out:
            out[s.name] = i
    return out
