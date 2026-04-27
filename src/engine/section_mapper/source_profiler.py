"""Generic source profiler — extracts a vendor-agnostic structure from
any source ``.docx`` (or PDF / txt via the existing extractor) for the
LLM-driven mapper to consume.

The returned :class:`SourceStructure` bundles everything the mapper
needs to build a complete substitution plan:

- Heading list (with auto-numbering markers resolved when the source
  is a ``.docx`` carrying ``<w:numPr>``).
- Per-section content lines (already prefixed with rendered list /
  letter / bullet markers).
- Document tables (raw cell texts, header row identified).
- Header text (raw, both glued and spaced flavors so the mapper can
  pull dotted document codes AND spaced multi-word titles).
- Sample free-text paragraphs (so the mapper can detect document-level
  metadata that the source author put in the body rather than the
  header).

JSON-serialisable. Designed to fit one mid-sized LLM prompt without
needing chunking for typical industrial-procedure documents.
"""

from __future__ import annotations

import re
import zipfile
from dataclasses import asdict, dataclass
from typing import TYPE_CHECKING

from engine.extractor import extract
from engine.section_mapper.parser import TextSection, parse_docx_source, parse_text

if TYPE_CHECKING:
    from pathlib import Path


@dataclass(frozen=True)
class SourceTable:
    """A table in the source document, flattened to cell texts."""

    index: int
    headers: list[str]
    rows: list[list[str]]

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass(frozen=True)
class SourceStructure:
    """Vendor-agnostic profile of a source document."""

    source_path: str
    sections: list[TextSection]
    tables: list[SourceTable]
    header_text_glued: str
    header_text_spaced: str
    body_paragraphs: list[str]  # flat full-text fallback for LLM

    def to_dict(self) -> dict:
        return {
            "source_path": self.source_path,
            "sections": [
                {
                    "name": s.name,
                    "raw_heading": s.raw_heading,
                    "number": s.number,
                    "level": s.level,
                    "content": s.content,
                }
                for s in self.sections
            ],
            "tables": [t.to_dict() for t in self.tables],
            "header_text_glued": self.header_text_glued,
            "header_text_spaced": self.header_text_spaced,
            "body_paragraphs": list(self.body_paragraphs),
        }


def profile_source(source_path: Path) -> SourceStructure:
    """Walk *source_path* and return a :class:`SourceStructure`."""
    sections = _profile_sections(source_path)
    tables = _profile_tables(source_path)
    glued, spaced = _profile_header_text(source_path)
    body_paragraphs = _profile_body_paragraphs(source_path)
    return SourceStructure(
        source_path=str(source_path),
        sections=sections,
        tables=tables,
        header_text_glued=glued,
        header_text_spaced=spaced,
        body_paragraphs=body_paragraphs,
    )


def _profile_body_paragraphs(source_path: Path) -> list[str]:
    """Flat list of every non-empty body paragraph.

    Used as a fallback when heading detection fails (English / Title-case
    sources, free-form documents). The LLM mapper segments this when
    ``sections`` is empty.
    """
    if source_path.suffix.lower() != ".docx":
        return []
    try:
        from docx import Document

        doc = Document(str(source_path))
    except Exception:
        return []
    return [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]


def _profile_sections(source_path: Path) -> list[TextSection]:
    """Use the existing parsers; route by extension."""
    if source_path.suffix.lower() == ".docx":
        return parse_docx_source(source_path)
    text = extract(source_path).text
    return parse_text(text)


def _profile_tables(source_path: Path) -> list[SourceTable]:
    """Read every table in a ``.docx`` source. PDFs / txt return empty."""
    if source_path.suffix.lower() != ".docx":
        return []
    try:
        from docx import Document

        doc = Document(str(source_path))
    except Exception:
        return []

    out: list[SourceTable] = []
    for i, t in enumerate(doc.tables):
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        if not rows:
            continue
        headers = rows[0]
        data_rows = [r for r in rows[1:] if any(c for c in r)]
        out.append(SourceTable(index=i, headers=headers, rows=data_rows))
    return out


def _profile_header_text(source_path: Path) -> tuple[str, str]:
    """Concatenate every ``<w:t>`` in every header file.

    Returns ``(glued, spaced)``: glued is runs joined without spacing
    (so dotted codes like ``IT.PRO.URE.387.0005`` reassemble across run
    fragmentation); spaced is runs joined with single spaces (so titles
    like ``PARTIDA DA ÁREA DE SÍNTESE`` followed by ``Ver.:`` don't
    merge into ``SÍNTESEVer``).
    """
    if source_path.suffix.lower() != ".docx":
        return "", ""

    glued_parts: list[str] = []
    spaced_parts: list[str] = []
    pattern = re.compile(r"^word/header\d*\.xml$")
    try:
        with zipfile.ZipFile(str(source_path)) as z:
            for name in z.namelist():
                if not pattern.match(name):
                    continue
                xml = z.read(name).decode("utf-8")
                texts = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)
                glued_parts.append("".join(texts))
                spaced_parts.append(" ".join(t for t in texts if t.strip()))
    except (OSError, zipfile.BadZipFile):
        return "", ""

    glued = " ".join(p.strip() for p in glued_parts if p.strip())
    spaced = " ".join(p.strip() for p in spaced_parts if p.strip())
    return glued, spaced


__all__ = [
    "SourceStructure",
    "SourceTable",
    "profile_source",
]
