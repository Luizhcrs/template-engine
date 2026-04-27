"""Apply a :class:`MappingPlan` to a template, producing the output docx.

Wraps the existing renderer / table_filler / header_filler primitives:

- ``plan.section_content`` → ``render_section_content``
- ``plan.table_data`` → ``fill_tables`` with synthesized
  :class:`TableSpec`s (one per template-table index in the plan).
- ``plan.header_substitutions`` → run-preserving substitution in every
  ``word/header*.xml`` (same engine as ``header_filler`` but driven by
  the plan instead of the hardcoded placeholder map).

Output is the same format the rules-mode pipeline produces, so callers
can swap modes without changing downstream consumers.
"""

from __future__ import annotations

import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import TYPE_CHECKING

from engine.section_mapper.parser import parse_docx
from engine.section_mapper.renderer import render_section_content
from engine.section_mapper.table_filler import TableSpec, fill_tables

if TYPE_CHECKING:
    from engine.section_mapper.auto_mapper import MappingPlan
    from engine.section_mapper.template_profiler import TemplateStructure


_HEADER_FILE_RE = re.compile(r"^word/header\d*\.xml$")
_BODY_FILE_RE = re.compile(r"^word/document\.xml$")
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def apply_mapping_plan(
    template_path: Path,
    output_path: Path,
    *,
    plan: MappingPlan,
    template: TemplateStructure,
) -> int:
    """Materialize *plan* into ``output_path``.

    Returns the number of tables filled.
    """
    target_sections = parse_docx(template_path)

    render_section_content(
        template_path,
        output_path,
        docx_sections=target_sections,
        content_by_target=plan.section_content,
    )

    specs = _build_table_specs(plan, template)
    filled = 0
    if specs:
        filled = fill_tables(template_path, output_path, specs)

    if plan.header_substitutions:
        _apply_header_substitutions(output_path, plan.header_substitutions)
        # Body placeholders (e.g. ``{{DOC_CODE}}``, ``[Title]`` on the
        # cover page) get the same substitution map applied to
        # ``word/document.xml`` — but only for placeholders whose text
        # is distinctive enough to avoid substring collision in body
        # paragraphs that contain similar runs (CNPJ masks like
        # ``__.___.___/____-__``, dotted-leader fields like
        # ``...........``).
        body_safe_subs = _filter_body_safe_subs(plan.header_substitutions)
        if body_safe_subs:
            _apply_body_substitutions(output_path, body_safe_subs)

    if plan.paragraph_rewrites:
        _apply_paragraph_rewrites(output_path, plan.paragraph_rewrites)

    if plan.cell_fills:
        _apply_cell_fills(output_path, plan.cell_fills)

    return filled


def _apply_cell_fills(docx_path: Path, fills: list) -> None:  # type: ignore[type-arg]
    """Replace text inside specific table cells addressed by
    (table_index, row, col).

    Used for mega-table layouts where the entire document is one
    table and each fillable spot is a cell coordinate.
    """
    if not fills:
        return

    from docx import Document

    doc = Document(str(docx_path))

    for fill in fills:
        if fill.table_index < 0 or fill.table_index >= len(doc.tables):
            continue
        table = doc.tables[fill.table_index]
        if fill.row < 0 or fill.row >= len(table.rows):
            continue
        row = table.rows[fill.row]
        if fill.col < 0 or fill.col >= len(row.cells):
            continue
        cell = row.cells[fill.col]
        # Replace cell text preserving the first paragraph's run formatting.
        if cell.paragraphs:
            para = cell.paragraphs[0]
            t_elements = para._p.findall(f".//{_W_NS}t")
            if t_elements:
                t_elements[0].text = fill.new_text
                for t in t_elements[1:]:
                    t.text = ""
            else:
                para.add_run(fill.new_text)
            # Clear any subsequent paragraphs in the cell.
            for extra_para in cell.paragraphs[1:]:
                for t in extra_para._p.findall(f".//{_W_NS}t"):
                    t.text = ""
        else:
            cell.text = fill.new_text

    doc.save(str(docx_path))


_DISTINCT_PLACEHOLDER_RE = re.compile(
    r"""(
        \{\{[^}]+\}\}     # curly token
        | \[[^\]]+\]      # bracket label
        | <<[^>]+>>       # double-angle
        | \([A-ZÁÉÍÓÚÂÊÔÃÕÇ_][A-ZÁÉÍÓÚÂÊÔÃÕÇ_ ]{2,}\)  # parens UPPERCASE label
    )""",
    re.VERBOSE,
)


def _filter_body_safe_subs(subs: dict[str, str]) -> dict[str, str]:
    """Drop substitutions whose placeholder text is not distinct enough
    to apply safely as a body-XML substring replace.

    Distinctive placeholders carry a delimiter pair (``{{...}}``,
    ``[...]``, ``<<...>>``, ``(LABEL)``) so the match is unambiguous in
    free-flowing paragraph text. Generic shapes (``__``, ``___``,
    ``XXXX``, ``____/____``, label-with-empty-suffix like
    ``Author:``) are header-only — they collide with surrounding text
    in body paragraphs (CPF / CNPJ masks, dotted leaders, etc.) and
    must not be substring-replaced there.
    """
    safe: dict[str, str] = {}
    for placeholder, replacement in subs.items():
        if _DISTINCT_PLACEHOLDER_RE.search(placeholder):
            safe[placeholder] = replacement
    return safe


def _build_table_specs(
    plan: MappingPlan,
    template: TemplateStructure,
) -> list[TableSpec]:
    """Translate the plan's ``table_data`` into ``TableSpec``s anchored
    to the template's empty tables (matched by ``template_table_index``).
    """
    by_index = {t.index: t for t in template.empty_tables}
    out: list[TableSpec] = []
    for entry in plan.table_data:
        tpl = by_index.get(entry.template_table_index)
        if tpl is None:
            continue
        if not entry.rows:
            continue
        out.append(
            TableSpec(
                headers=tpl.primary_headers,
                rows=entry.rows,
                subheaders=entry.sub_headers or None,
            )
        )
    return out


def _apply_header_substitutions(docx_path: Path, subs: dict[str, str]) -> None:
    """Inline-substitute placeholders in every ``word/header*.xml``.

    Placeholder match is per-``<w:t>`` (atomic node). Templates almost
    always store each header field as its own run group, so this works
    well for ``XXXX``, ``Rev. 00``, ``Elaborado:``, etc. For literals
    that span runs (rare), the LLM picks an alternate placeholder text
    that does sit in one ``<w:t>``.
    """
    if not subs:
        return

    with tempfile.NamedTemporaryFile(
        suffix=".docx",
        delete=False,
        dir=str(docx_path.parent),
    ) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with (
            zipfile.ZipFile(str(docx_path), "r") as zin,
            zipfile.ZipFile(str(tmp_path), "w", zipfile.ZIP_DEFLATED) as zout,
        ):
            for item in zin.infolist():
                data = zin.read(item.filename)
                if _HEADER_FILE_RE.match(item.filename):
                    text = data.decode("utf-8")
                    text = _replace_in_runs(text, subs)
                    data = text.encode("utf-8")
                zout.writestr(item, data)
        shutil.move(str(tmp_path), str(docx_path))
    except Exception:
        if tmp_path.exists():
            tmp_path.unlink()
        raise


def _apply_paragraph_rewrites(docx_path: Path, rewrites: list) -> None:  # type: ignore[type-arg]
    """Replace whole paragraphs whose rendered text matches
    ``rewrite.match_text`` with ``rewrite.replacement_text``. Walks
    body + headers + footers. Run formatting of the first non-empty
    run is preserved; trailing runs in the paragraph are cleared.
    """
    if not rewrites:
        return

    from docx import Document

    doc = Document(str(docx_path))

    rewrite_map = {r.match_text.strip(): r.replacement_text for r in rewrites}

    def _try_rewrite_para(para) -> None:  # type: ignore[no-untyped-def]
        text = para.text.strip()
        if text in rewrite_map:
            replacement = rewrite_map[text]
            t_elements = para._p.findall(f".//{_W_NS}t")
            if t_elements:
                t_elements[0].text = replacement
                for t in t_elements[1:]:
                    t.text = ""
            else:
                para.add_run(replacement)

    for para in doc.paragraphs:
        _try_rewrite_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _try_rewrite_para(para)

    for section in doc.sections:
        for container in (section.header, section.footer):
            for para in container.paragraphs:
                _try_rewrite_para(para)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _try_rewrite_para(para)

    doc.save(str(docx_path))


def _apply_body_substitutions(docx_path: Path, subs: dict[str, str]) -> None:
    """Substitute placeholders inside body paragraph text via python-docx
    (so XML-escaped chars like ``&lt;&lt;TITULO&gt;&gt;`` are matched
    against their unescaped form ``<<TITULO>>``).

    Walks every paragraph text node in body, body-tables, and cell
    sub-paragraphs. Header paragraphs are left to ``_apply_header_substitutions``.
    """
    if not subs:
        return

    from docx import Document

    doc = Document(str(docx_path))

    def _replace_in_para(para) -> None:  # type: ignore[no-untyped-def]
        text = para.text
        if not text:
            return
        new_text = text
        replaced = False
        for placeholder, replacement in subs.items():
            if placeholder and placeholder in new_text:
                new_text = new_text.replace(placeholder, replacement)
                replaced = True
        if not replaced or new_text == text:
            return
        # Preserve first run's formatting; clear trailing runs.
        t_elements = para._p.findall(f".//{_W_NS}t")
        if t_elements:
            t_elements[0].text = new_text
            for t in t_elements[1:]:
                t.text = ""
        else:
            para.add_run(new_text)

    for para in doc.paragraphs:
        _replace_in_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para)

    doc.save(str(docx_path))


def _replace_in_runs(xml: str, subs: dict[str, str]) -> str:
    """Apply each ``placeholder -> replacement`` inside every ``<w:t>``."""

    def repl(m: re.Match[str]) -> str:
        prefix = m.group(1)
        inner = m.group(2)
        suffix = m.group(3)
        for placeholder, replacement in subs.items():
            if not placeholder:
                continue
            if placeholder in inner:
                inner = inner.replace(placeholder, _xml_escape(replacement), 1)
        return f"{prefix}{inner}{suffix}"

    pattern = re.compile(r"(<w:t(?:\s[^>]*)?>)([^<]*)(</w:t>)")
    return pattern.sub(repl, xml)


def _xml_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


__all__ = ["apply_mapping_plan"]
