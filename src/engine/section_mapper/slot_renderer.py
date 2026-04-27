"""Wave N — apply slot fills to a docx in place.

Walks the inventory + dict of fills and writes each new text at the
slot's exact address. Never clones, inserts, or removes structure —
only substitutes the inner text of paragraphs / cells.

This is the atomic counterpart to :func:`engine.section_mapper.slot_filler.fill_slots`.
"""

from __future__ import annotations

import re
import shutil
import tempfile
import zipfile
from pathlib import Path

from engine.section_mapper.slots import Slot, SlotInventory

# Re-export to silence unused-import lint (we need Slot at runtime).
_ = (Slot, SlotInventory, Path)

_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_HEADER_FOOTER_RE = re.compile(r"^word/(?:header|footer)\d*\.xml$")


def apply_slot_fills(
    template_path: Path,
    output_path: Path,
    *,
    inventory: SlotInventory,
    fills: dict[str, str],
) -> int:
    """Open *template_path*, apply every ``fills[slot_id]`` to its
    addressed slot, save to *output_path*.

    Returns the count of slots actually filled.

    Body paragraphs and table cells are rewritten via python-docx so
    XML escapes (``<<TITULO>>`` stored as ``&lt;&lt;TITULO&gt;&gt;``)
    match against unescaped text. Header / footer paragraphs are
    rewritten directly inside the docx zip's ``word/headerN.xml`` /
    ``word/footerN.xml`` parts because python-docx's iterators skip
    paragraphs nested inside text boxes.
    """
    if not fills:
        # Still copy template → output so the caller sees a file.
        if template_path != output_path:
            shutil.copy(str(template_path), str(output_path))
        return 0

    slots_by_id = inventory.by_id()

    body_fills: dict[Slot, str] = {}
    cell_fills: dict[Slot, str] = {}
    header_fills: dict[Slot, str] = {}
    footer_fills: dict[Slot, str] = {}

    for slot_id, new_text in fills.items():
        slot = slots_by_id.get(slot_id)
        if slot is None:
            continue
        loc = slot.address.location
        if loc == "body_para":
            body_fills[slot] = new_text
        elif loc == "table_cell":
            cell_fills[slot] = new_text
        elif loc == "header_para":
            header_fills[slot] = new_text
        elif loc == "footer_para":
            footer_fills[slot] = new_text

    # Phase 1 — write body paragraphs + cells via python-docx, save to output.
    _write_body(template_path, output_path, body_fills, cell_fills)

    # Phase 2 — rewrite header / footer XML parts in the saved zip.
    if header_fills or footer_fills:
        _write_header_footer(output_path, header_fills, footer_fills)

    return len(body_fills) + len(cell_fills) + len(header_fills) + len(footer_fills)


def _write_body(
    template_path: Path,
    output_path: Path,
    body_fills: dict,  # type: ignore[type-arg]
    cell_fills: dict,  # type: ignore[type-arg]
) -> None:
    from docx import Document

    doc = Document(str(template_path))

    for slot, new_text in body_fills.items():
        idx = slot.address.paragraph_idx
        if idx is None or idx < 0 or idx >= len(doc.paragraphs):
            continue
        _set_paragraph_text(doc.paragraphs[idx], new_text)

    for slot, new_text in cell_fills.items():
        ti = slot.address.table_index
        ri = slot.address.row
        ci = slot.address.col
        if ti is None or ri is None or ci is None:
            continue
        if ti < 0 or ti >= len(doc.tables):
            continue
        table = doc.tables[ti]
        if ri < 0 or ri >= len(table.rows):
            continue
        row = table.rows[ri]
        if ci < 0 or ci >= len(row.cells):
            continue
        target = row.cells[ci]
        target_text = target.text.strip()
        # Mirror across merged-column siblings (same row, identical text).
        for sibling in row.cells:
            if sibling.text.strip() == target_text:
                _set_cell_text(sibling, new_text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _set_paragraph_text(paragraph, text: str) -> None:  # type: ignore[no-untyped-def]
    p_elem = paragraph._p
    text_elements = p_elem.findall(f".//{_W_NS}t")
    if text_elements:
        text_elements[0].text = text
        for t in text_elements[1:]:
            t.text = ""
        return
    paragraph.add_run(text)


def _set_cell_text(cell, new_text: str) -> None:  # type: ignore[no-untyped-def]
    if cell.paragraphs:
        para = cell.paragraphs[0]
        t_elements = para._p.findall(f".//{_W_NS}t")
        if t_elements:
            t_elements[0].text = new_text
            for t in t_elements[1:]:
                t.text = ""
        else:
            para.add_run(new_text)
        for extra in cell.paragraphs[1:]:
            for t in extra._p.findall(f".//{_W_NS}t"):
                t.text = ""
    else:
        cell.text = new_text


def _write_header_footer(
    docx_path: Path,
    header_fills: dict,  # type: ignore[type-arg]
    footer_fills: dict,  # type: ignore[type-arg]
) -> None:
    """Rewrite text inside header / footer XML parts.

    Each fill targets a specific paragraph index inside a specific
    ``word/headerN.xml`` / ``word/footerN.xml``. We rebuild the zip
    once, replacing the targeted paragraphs' inner ``<w:t>`` elements
    with the new text.
    """
    # Group by part name → list of (paragraph_idx, new_text)
    by_part: dict[str, list[tuple[int, str]]] = {}
    for slot, new_text in {**header_fills, **footer_fills}.items():
        part_name = slot.address.part_name
        idx = slot.address.paragraph_idx
        if part_name is None or idx is None:
            continue
        by_part.setdefault(part_name, []).append((idx, new_text))

    if not by_part:
        return

    with tempfile.NamedTemporaryFile(
        suffix=".docx",
        delete=False,
        dir=str(docx_path.parent),
    ) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with (
            zipfile.ZipFile(str(docx_path), "r") as zin,
            zipfile.ZipFile(str(tmp_path), "w", zipfile.ZIP_DEFLATED) as zout,
        ):
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename in by_part:
                    xml = data.decode("utf-8")
                    xml = _rewrite_paragraphs_in_part(xml, by_part[item.filename])
                    data = xml.encode("utf-8")
                zout.writestr(item, data)
        shutil.move(str(tmp_path), str(docx_path))
    except Exception:
        if tmp_path.exists():
            tmp_path.unlink()
        raise


def _rewrite_paragraphs_in_part(
    xml: str,
    fills: list[tuple[int, str]],
) -> str:
    """Rewrite the inner ``<w:t>`` text of paragraphs at the given
    indices. Walks ``<w:p>`` elements in document order.
    """
    fill_by_idx = dict(fills)
    para_pattern = re.compile(r"(<w:p\b[^>]*>)((?:.|\n)*?)(</w:p>)")
    text_pattern = re.compile(r"(<w:t[^>]*>)([^<]*)(</w:t>)")

    counter = {"i": 0}

    def repl(m: re.Match[str]) -> str:
        idx = counter["i"]
        counter["i"] += 1
        if idx not in fill_by_idx:
            return m.group(0)

        new_text = _xml_escape(fill_by_idx[idx])
        body = m.group(2)
        first_done = {"d": False}

        def t_repl(tm: re.Match[str]) -> str:
            if not first_done["d"]:
                first_done["d"] = True
                return f"{tm.group(1)}{new_text}{tm.group(3)}"
            return f"{tm.group(1)}{tm.group(3)}"

        new_body, n = text_pattern.subn(t_repl, body)
        if n == 0:
            # No <w:t> in the paragraph — skip rather than corrupt.
            return m.group(0)
        return f"{m.group(1)}{new_body}{m.group(3)}"

    return para_pattern.sub(repl, xml)


def _xml_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


__all__ = ["apply_slot_fills"]
