"""Section content renderer — inserts multi-line content into a docx template.

The vanilla approach of stuffing multi-line text into a single ``<w:t>``
yields garbage when the host paragraph has ``<w:jc w:val="both"/>``
(justified) — Word distributes words across the line width and the output
looks like newspaper columns.

This renderer does it right:

1. Find the heading paragraph in the template.
2. Locate the first empty body paragraph below it (the "anchor").
3. Drop ``<w:jc>`` from the anchor's pPr to prevent the justified-blowout.
4. Set the anchor's text to line 1 of the content.
5. For each remaining line, clone the anchor's ``<w:p>``, clear inner
   ``<w:t>`` text, set line N, insert via ``addnext`` so paragraph order
   is preserved.

Bullet points / dash-prefixed lines stay as separate paragraphs (the
template usually has list paragraphs already styled); we don't rewrite
the list level.
"""

from __future__ import annotations

import re
from copy import deepcopy
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from pathlib import Path

    from engine.section_mapper.parser import DocxSection


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# Sub-heading prefixes that should render as bold + space-before.
# ``5.2.`` and ``5.2.1.`` cover the dotted-number pattern; the regex
# requires a SPACE after the number so ``5.5%`` (units) is not matched.
_SUBHEADING_RE = re.compile(r"^\s*(\d+(?:\.\d+){1,3})\.?\s+\S")
_SUBSUBHEADING_RE = re.compile(r"^\s*(\d+(?:\.\d+){2,3})\.?\s+\S")
_NOTA_RE = re.compile(r"^\s*Nota\s*\d*[:.]\s", re.IGNORECASE)


def _detect_line_kind(line: str) -> str:
    """Return ``"subsubheading"`` / ``"subheading"`` / ``"nota"`` /
    ``"body"`` based on the rendered marker prefix.
    """
    if _SUBSUBHEADING_RE.match(line):
        return "subsubheading"
    if _SUBHEADING_RE.match(line):
        return "subheading"
    if _NOTA_RE.match(line):
        return "nota"
    return "body"


def _add_run_emphasis(  # type: ignore[no-untyped-def]
    p_elem,
    *,
    italic: bool = False,
    bold: bool = False,
    color: str | None = None,
) -> None:
    """Apply bold / italic / explicit color to every run in the paragraph.

    ``color`` is a 6-digit hex string (``"000000"`` for black). When set,
    overrides any inherited color so headings don't pick up Word's
    default heading-blue from a Ttulo style.
    """
    for r in p_elem.findall(f"{_W_NS}r"):
        rPr = r.find(f"{_W_NS}rPr")
        if rPr is None:
            rPr = r.makeelement(f"{_W_NS}rPr", {})
            r.insert(0, rPr)
        if italic and rPr.find(f"{_W_NS}i") is None:
            rPr.append(rPr.makeelement(f"{_W_NS}i", {}))
        if bold and rPr.find(f"{_W_NS}b") is None:
            rPr.append(rPr.makeelement(f"{_W_NS}b", {}))
        if color is not None:
            existing = rPr.find(f"{_W_NS}color")
            if existing is not None:
                rPr.remove(existing)
            rPr.append(rPr.makeelement(f"{_W_NS}color", {f"{_W_NS}val": color}))


def _set_paragraph_spacing(  # type: ignore[no-untyped-def]
    p_elem,
    *,
    before_twips: int,
    after_twips: int,
) -> None:
    """Set ``<w:spacing>`` directly on the paragraph (twips: 1pt = 20)."""
    pPr = p_elem.find(f"{_W_NS}pPr")
    if pPr is None:
        pPr = p_elem.makeelement(f"{_W_NS}pPr", {})
        p_elem.insert(0, pPr)
    existing = pPr.find(f"{_W_NS}spacing")
    if existing is not None:
        pPr.remove(existing)
    pPr.append(
        pPr.makeelement(
            f"{_W_NS}spacing",
            {
                f"{_W_NS}before": str(before_twips),
                f"{_W_NS}after": str(after_twips),
            },
        )
    )


def _decorate_for_kind(p_elem, kind: str) -> None:  # type: ignore[no-untyped-def]
    """Apply per-kind decoration via direct formatting only.

    No ``<w:pStyle>`` is applied — Word's default heading styles render
    blue, which is wrong for industrial-procedure documents that expect
    black bold sub-headings. Direct formatting (bold + spacing + black
    color) keeps the visual hierarchy without inheriting style colors.
    """
    if kind == "subheading":
        _add_run_emphasis(p_elem, bold=True, color="000000")
        _set_paragraph_spacing(p_elem, before_twips=240, after_twips=120)
        return
    if kind == "subsubheading":
        _add_run_emphasis(p_elem, bold=True, color="000000")
        _set_paragraph_spacing(p_elem, before_twips=180, after_twips=80)
        return
    if kind == "nota":
        _add_run_emphasis(p_elem, italic=True)


def _reset_paragraph_style(p_elem) -> None:  # type: ignore[no-untyped-def]
    """Drop inline bold/italic and any prior spacing override so a cloned
    anchor doesn't inherit decorations from the previous line.
    """
    pPr = p_elem.find(f"{_W_NS}pPr")
    if pPr is not None:
        # Drop our own paragraph-level spacing override.
        spacing = pPr.find(f"{_W_NS}spacing")
        if spacing is not None and (spacing.get(f"{_W_NS}before") or spacing.get(f"{_W_NS}after")):
            pPr.remove(spacing)
    for r in p_elem.findall(f"{_W_NS}r"):
        rPr = r.find(f"{_W_NS}rPr")
        if rPr is None:
            continue
        for tag in (f"{_W_NS}b", f"{_W_NS}i"):
            el = rPr.find(tag)
            if el is not None:
                rPr.remove(el)


def _strip_jc(p_elem) -> None:  # type: ignore[no-untyped-def]
    """Drop ``<w:jc>`` so a multi-line block doesn't render as columns."""
    pPr = p_elem.find(f"{_W_NS}pPr")
    if pPr is None:
        return
    jc = pPr.find(f"{_W_NS}jc")
    if jc is not None:
        pPr.remove(jc)


def _split_into_lines(content: str) -> list[str]:
    """Conservative line split. Drops empty lines, keeps bullets intact."""
    raw = [ln.strip() for ln in content.splitlines()]
    return [ln for ln in raw if ln]


def _set_paragraph_text(paragraph, text: str) -> None:  # type: ignore[no-untyped-def]
    """Set the paragraph's rendered text to *text*.

    For paragraphs that already have ``<w:r>`` runs, replaces every
    ``<w:t>`` payload (preserves run formatting). For empty paragraphs
    (no runs yet, common on template body slots), creates a fresh run via
    ``python-docx``.
    """
    p_elem = paragraph._p
    text_elements = p_elem.findall(f".//{_W_NS}t")
    if text_elements:
        text_elements[0].text = text
        for t in text_elements[1:]:
            t.text = ""
        return
    # Empty paragraph: add a fresh run so a <w:t> exists.
    paragraph.add_run(text)


def render_section_content(
    template_path: Path,
    output_path: Path,
    *,
    docx_sections: list[DocxSection],
    content_by_target: dict[str, str],
) -> None:
    """Open template, fill each empty section with its mapped content, save."""
    from docx import Document

    doc = Document(str(template_path))
    paragraphs = list(doc.paragraphs)

    heading_idx = {s.name: s.heading_paragraph_idx for s in docx_sections}
    # Track each anchor + the index of the LAST inserted paragraph in the
    # source paragraphs list so we can prune leftover empties later.
    last_filled_idx_by_target: dict[str, int] = {}

    for target_name, content in content_by_target.items():
        if not content.strip():
            continue
        idx = heading_idx.get(target_name)
        if idx is None:
            continue

        next_heading_idx = _next_heading_idx(idx, docx_sections)
        anchor = None
        anchor_idx: int | None = None
        for j in range(idx + 1, next_heading_idx if next_heading_idx else len(paragraphs)):
            if not paragraphs[j].text.strip():
                anchor = paragraphs[j]
                anchor_idx = j
                break
        if anchor is None or anchor_idx is None:
            continue

        lines = _split_into_lines(content)
        if not lines:
            continue

        _strip_jc(anchor._p)
        _set_paragraph_text(anchor, lines[0])
        _decorate_for_kind(anchor._p, _detect_line_kind(lines[0]))

        cursor = anchor._p
        for line in lines[1:]:
            new_p = deepcopy(anchor._p)
            _strip_jc(new_p)
            # Reset any heading style/emphasis copied from the previous
            # cloned anchor — each line decides its own style.
            _reset_paragraph_style(new_p)
            ts = new_p.findall(f".//{_W_NS}t")
            if ts:
                ts[0].text = line
                for t in ts[1:]:
                    t.text = ""
            _decorate_for_kind(new_p, _detect_line_kind(line))
            cursor.addnext(new_p)
            cursor = new_p

        last_filled_idx_by_target[target_name] = anchor_idx

    # Prune empty body paragraphs left over between the last filled anchor
    # and the next heading. The template ships with multiple blank slots
    # per section (so the user can paste arbitrary content); keeping the
    # ones we did not use produces large vertical gaps in Word.
    _prune_unused_body_slots(
        paragraphs,
        docx_sections,
        last_filled_idx_by_target,
    )
    # Collapse any run of >1 empty paragraphs in the body to a single
    # empty so end-of-document filler space and "between sections we
    # didn't fill" gaps don't stretch the output vertically.
    _collapse_empty_paragraph_runs(doc)

    doc.save(str(output_path))


def _collapse_empty_paragraph_runs(doc) -> None:  # type: ignore[no-untyped-def]
    """Walk the document body once; whenever 2+ consecutive empty
    paragraphs appear at the same nesting level, drop all but the first.

    Paragraphs inside tables are left alone — table cells are sensitive
    to paragraph count for cell sizing.
    """
    body = doc.element.body
    _collapse_in_element(body)


def _collapse_in_element(parent) -> None:  # type: ignore[no-untyped-def]
    """Walk *parent*'s direct children; collapse consecutive empty
    ``<w:p>`` runs (siblings only). Recurse into ``<w:tc>`` cells'
    inner bodies — but skip the table cell collapse to preserve cell
    layout.
    """
    consecutive_empty: list = []  # type: ignore[type-arg]
    for child in list(parent):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text_content = "".join(t.text or "" for t in child.findall(f".//{_W_NS}t"))
            if not text_content.strip():
                consecutive_empty.append(child)
                if len(consecutive_empty) > 1:
                    parent.remove(child)
                continue
            consecutive_empty = []
        else:
            consecutive_empty = []


def _prune_unused_body_slots(
    paragraphs: list,  # type: ignore[type-arg]
    docx_sections: list[DocxSection],
    last_filled_idx_by_target: dict[str, int],
) -> None:
    """Delete blank body paragraphs that follow a filled anchor up until
    the next heading. Anchors carry section formatting, so we can't drop
    EVERYTHING — only paragraphs strictly between our last inserted
    content and the next heading.
    """
    heading_idx_by_name = {s.name: s.heading_paragraph_idx for s in docx_sections}
    for target_name, anchor_idx in last_filled_idx_by_target.items():
        idx = heading_idx_by_name.get(target_name)
        if idx is None:
            continue
        next_idx = _next_heading_idx(idx, docx_sections)
        if next_idx is None:
            continue
        # Anchor + cloned-next-paragraphs live AFTER anchor_idx in the doc
        # tree (we used addnext). The paragraphs *list* is stale — the
        # original "blank slot" indices stay valid, but new clones are
        # inserted between them. Walk the live siblings instead.
        anchor_p = paragraphs[anchor_idx]._p
        sibling = anchor_p.getnext()
        while sibling is not None and sibling.tag.endswith("}p"):
            text_content = "".join(t.text or "" for t in sibling.findall(f".//{_W_NS}t"))
            if text_content.strip():
                # The next heading's text is non-empty too, so we stop.
                break
            parent = sibling.getparent()
            nxt = sibling.getnext()
            parent.remove(sibling)
            sibling = nxt


def _next_heading_idx(current_idx: int, sections: list[DocxSection]) -> int | None:
    """Return the next heading's paragraph index after *current_idx*, or None."""
    candidates = [s.heading_paragraph_idx for s in sections if s.heading_paragraph_idx > current_idx]
    return min(candidates) if candidates else None


def detect_orphan_paragraphs(output_path: Path) -> list[str]:
    """Return paragraph texts that still contain ``{{X}}`` / ``[X]`` / ``___``.

    Useful as a post-render sanity check.
    """
    from docx import Document

    doc = Document(str(output_path))
    pattern = re.compile(r"\{\{[A-Za-z_][\w.-]*\}\}|\[[A-Za-z_][\w.-]*\]|_{3,}")
    return [p.text for p in doc.paragraphs if pattern.search(p.text)]
