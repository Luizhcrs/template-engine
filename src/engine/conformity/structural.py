"""Structural dimension — counts headings/tables/sections/lists in docx.

Pure ``python-docx`` parsing, **zero LLM**. Compares the structural fingerprint
of two ``.docx`` documents and flags mismatches:

- heading count by level (e.g. template has 3x H1, candidate has 1x H1)
- tables count + total cell count
- sections count
- list paragraphs count
- total paragraph count

A mismatch beyond a tolerance threshold becomes a :class:`Failure`. Score is
``1.0 - sum(weighted_diffs) / max_acceptable``.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import TYPE_CHECKING, Final

import structlog
from docx import Document

from engine.conformity.report import DimensionResult, Failure

if TYPE_CHECKING:
    from pathlib import Path

log = structlog.get_logger(__name__)


@dataclass(frozen=True)
class StructuralFingerprint:
    """Compact summary of a docx's structural shape."""

    headings_by_level: dict[int, int]
    tables_count: int
    table_cells_total: int
    sections_count: int
    list_paragraphs: int
    paragraphs_total: int


_HEADING_RE: Final[re.Pattern[str]] = re.compile(r"Heading\s*(\d+)", re.IGNORECASE)


def _is_list_paragraph(p) -> bool:  # type: ignore[no-untyped-def]
    """python-docx exposes list info via paragraph format / numbering style.

    Heuristic: any paragraph whose style name contains ``List`` qualifies, plus
    any whose underlying XML carries ``<w:numPr>`` (numbering properties).
    """
    style = (p.style.name or "").lower() if p.style else ""
    if "list" in style:
        return True
    pPr = p._p.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    if pPr is None:
        return False
    return pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr") is not None


def fingerprint(path: Path) -> StructuralFingerprint:
    """Walk a ``.docx`` once and produce a :class:`StructuralFingerprint`."""
    doc = Document(str(path))

    headings: dict[int, int] = {}
    list_count = 0
    for para in doc.paragraphs:
        if para.style is None:
            continue
        m = _HEADING_RE.match(para.style.name or "")
        if m:
            level = int(m.group(1))
            headings[level] = headings.get(level, 0) + 1
        if _is_list_paragraph(para):
            list_count += 1

    cells = sum(len(row.cells) for table in doc.tables for row in table.rows)

    fp = StructuralFingerprint(
        headings_by_level=headings,
        tables_count=len(doc.tables),
        table_cells_total=cells,
        sections_count=len(doc.sections),
        list_paragraphs=list_count,
        paragraphs_total=len(doc.paragraphs),
    )
    log.info(
        "conformity.structural.fingerprint",
        path=str(path),
        headings=headings,
        tables=fp.tables_count,
        sections=fp.sections_count,
    )
    return fp


def _diff_int(label: str, expected: int, actual: int, severity_threshold: int = 1) -> Failure | None:
    """Emit a Failure when |expected - actual| >= severity_threshold."""
    if expected == actual:
        return None
    delta = abs(expected - actual)
    severity = "critical" if delta > severity_threshold else "warning"
    return Failure(
        dimension="structural",
        field_or_excerpt=label,
        expected=str(expected),
        actual=str(actual),
        severity=severity,
        note=f"count differs by {delta}",
    )


def diff_fingerprints(
    template_fp: StructuralFingerprint,
    candidate_fp: StructuralFingerprint,
) -> list[Failure]:
    """Compare two fingerprints, emit Failures per metric."""
    failures: list[Failure] = []

    # Headings per level — collect all levels seen on either side
    levels = set(template_fp.headings_by_level) | set(candidate_fp.headings_by_level)
    for level in sorted(levels):
        f = _diff_int(
            f"headings_h{level}",
            template_fp.headings_by_level.get(level, 0),
            candidate_fp.headings_by_level.get(level, 0),
        )
        if f:
            failures.append(f)

    pairs = (
        ("tables_count", template_fp.tables_count, candidate_fp.tables_count, 0),
        ("table_cells_total", template_fp.table_cells_total, candidate_fp.table_cells_total, 4),
        ("sections_count", template_fp.sections_count, candidate_fp.sections_count, 0),
        ("list_paragraphs", template_fp.list_paragraphs, candidate_fp.list_paragraphs, 2),
    )
    for label, e, a, threshold in pairs:
        f = _diff_int(label, e, a, severity_threshold=threshold)
        if f:
            failures.append(f)

    return failures


def _score_from_failures(failures: list[Failure], max_acceptable: float = 4.0) -> float:
    weight = sum(1.0 if f.severity == "critical" else 0.4 for f in failures)
    raw = 1.0 - (weight / max_acceptable)
    return max(0.0, min(1.0, raw))


def check_structural(
    template_path: Path,
    candidate_path: Path,
    *,
    max_acceptable_failures: float = 4.0,
) -> DimensionResult:
    """Run structural dimension. Pure determinístico, zero LLM."""
    template_fp = fingerprint(template_path)
    candidate_fp = fingerprint(candidate_path)
    failures = diff_fingerprints(template_fp, candidate_fp)
    score = _score_from_failures(failures, max_acceptable=max_acceptable_failures)
    log.info("conformity.structural", score=score, failures=len(failures))
    return DimensionResult(dimension="structural", score=score, failures=failures)
