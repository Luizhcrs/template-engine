import structlog
from docx.oxml.ns import qn

log = structlog.get_logger(__name__)


def _find_heading(doc, heading_text: str):
    for p in doc.paragraphs:
        if p.text.strip().upper() == heading_text.upper() and p.style.name.startswith("Heading"):
            return p
    log.warning("render_ops.heading_not_found", heading=heading_text)
    return None


def _remove_placeholder_after(heading_p):
    """Remove the paragraph immediately following heading_p if it is '[A DEFINIR]'."""
    nxt = heading_p._element.getnext()
    if nxt is None or not nxt.tag.endswith("}p"):
        return
    text = "".join(t.text or "" for t in nxt.iter(qn("w:t")) if t.text)
    if text.strip() == "[A DEFINIR]":
        nxt.getparent().remove(nxt)


def write_section(ctx: dict, params: dict) -> None:
    """Append a single text block under a named heading."""
    doc = ctx["doc"]
    content = ctx["content"]
    heading = params["heading"]
    source_key = params["source_key"]
    text = str(content.get(source_key, ""))

    h = _find_heading(doc, heading)
    if h is None:
        doc.add_heading(heading, level=1)
    else:
        _remove_placeholder_after(h)

    doc.add_paragraph(text)


def write_list(ctx: dict, params: dict) -> None:
    """Append bullet-style list items under a named heading.

    params: heading, source_key, marker (default '-')
    content[source_key] must be list[str].
    """
    doc = ctx["doc"]
    content = ctx["content"]
    heading = params["heading"]
    source_key = params["source_key"]
    marker = params.get("marker", "-")
    items = content.get(source_key, []) or []

    h = _find_heading(doc, heading)
    if h is None:
        doc.add_heading(heading, level=1)
    else:
        _remove_placeholder_after(h)

    for item in items:
        doc.add_paragraph(f"{marker} {item}")


def write_table(ctx: dict, params: dict) -> None:
    """Append a table under a named heading.

    params: heading, source_key, columns (list[str])
    content[source_key] must be list[dict]; column lookup is case-insensitive
    (tries exact key, then lowercase).
    """
    doc = ctx["doc"]
    content = ctx["content"]
    heading = params["heading"]
    source_key = params["source_key"]
    columns = params["columns"]
    rows = content.get(source_key, []) or []

    h = _find_heading(doc, heading)
    if h is None:
        doc.add_heading(heading, level=1)
    else:
        _remove_placeholder_after(h)

    t = doc.add_table(rows=1, cols=len(columns))
    t.style = "Table Grid"
    for i, col in enumerate(columns):
        t.rows[0].cells[i].text = col
    for row in rows:
        cells = t.add_row().cells
        for i, col in enumerate(columns):
            val = row.get(col, row.get(col.lower(), ""))
            cells[i].text = str(val)


def write_steps(ctx: dict, params: dict) -> None:
    """Append numbered steps under a named heading.

    params: heading, source_key, prefix (e.g. '5.2.'), note_prefix (e.g. 'NOTA')
    content[source_key] must be list[dict] with keys 'tipo' ('passo'|'nota'|'cabecalho') and 'texto'.
    """
    doc = ctx["doc"]
    content = ctx["content"]
    heading = params["heading"]
    source_key = params["source_key"]
    prefix = params.get("prefix", "")
    note_prefix = params.get("note_prefix", "NOTA")
    steps = content.get(source_key, []) or []

    h = _find_heading(doc, heading)
    if h is None:
        doc.add_heading(heading, level=1)
    else:
        _remove_placeholder_after(h)

    passo_num = 1
    nota_num = 1
    for step in steps:
        tipo = step.get("tipo", "passo")
        texto = str(step.get("texto", ""))
        if tipo == "passo":
            doc.add_paragraph(f"{prefix}{passo_num}. {texto}")
            passo_num += 1
        elif tipo == "nota":
            doc.add_paragraph(f"{note_prefix} {nota_num}: {texto}")
            nota_num += 1
        else:
            doc.add_paragraph(texto)
