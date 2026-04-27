"""Batch orchestrator — drive 1 template + N source docs → N normalized outputs.

End-to-end pipeline that wires Wave D modules together:

1. :mod:`engine.schema_inference` → derive field schema from the template.
2. :mod:`engine.pattern_inference` → if gold examples were supplied, build
   regex patterns once.
3. For each source doc, in parallel:

   a. :func:`engine.extractor.extract` to pull raw text.
   b. :func:`engine.hybrid_mapper.map_hybrid` regex-first then LLM fallback.
   c. Render output by direct token substitution in a copy of the template.
   d. :func:`engine.semantic_diff.diff_documents` to flag information loss.
   e. Bucket into a confidence tier (``high`` / ``medium`` / ``low``).

4. Write ``report.json`` summarizing counts, per-doc outcomes, total cost
   estimate (LLM call count), and any errors.

Designed for the "400 docs to normalize, pra ontem" use case: regex resolves the
bulk for free; LLM only touches what regex couldn't; semantic diff is the safety
net that catches information loss.
"""

from __future__ import annotations

import asyncio
import re
import shutil
from dataclasses import asdict, dataclass, field
from datetime import UTC, datetime
from typing import TYPE_CHECKING, Final

import structlog
from docx import Document

from engine.extractor import extract
from engine.hybrid_mapper import MappingResult, map_hybrid, summarize
from engine.llm.base import LLMError
from engine.pattern_inference import infer_field_patterns
from engine.schema_inference import detect_placeholders, enrich_with_llm
from engine.semantic_diff import Discrepancy, diff_documents

if TYPE_CHECKING:
    from pathlib import Path

    from engine.llm.base import LLMProvider
    from engine.pattern_inference import InferredPattern
    from engine.schema_inference import FieldSchema
    from engine.security.audit import AuditLog

log = structlog.get_logger(__name__)

_DEFAULT_MAX_CONCURRENT: Final[int] = 4
_SUPPORTED_EXTS: Final[tuple[str, ...]] = (".docx", ".pdf")


@dataclass(frozen=True)
class BatchItemResult:
    """Outcome for one source document."""

    source_path: Path
    output_path: Path | None
    tier: str  # "high" | "medium" | "low" | "error"
    mapping: dict[str, MappingResult]
    discrepancies: list[Discrepancy]
    error: str | None = None


@dataclass(frozen=True)
class BatchReport:
    """Aggregate report over a batch run."""

    template_path: Path
    output_dir: Path
    schemas: list[FieldSchema]
    items: list[BatchItemResult]
    started_at: str
    finished_at: str
    llm_call_count: int
    extras: dict = field(default_factory=dict)

    @property
    def by_tier(self) -> dict[str, int]:
        counts: dict[str, int] = {"high": 0, "medium": 0, "low": 0, "error": 0}
        for item in self.items:
            counts[item.tier] = counts.get(item.tier, 0) + 1
        return counts

    def to_dict(self) -> dict:
        """JSON-serializable view of the report."""
        return {
            "template_path": str(self.template_path),
            "output_dir": str(self.output_dir),
            "started_at": self.started_at,
            "finished_at": self.finished_at,
            "llm_call_count": self.llm_call_count,
            "by_tier": self.by_tier,
            "fields": [
                {
                    "name": s.name,
                    "kind": s.kind,
                    "field_type": s.field_type,
                    "required": s.required,
                }
                for s in self.schemas
            ],
            "items": [
                {
                    "source": str(item.source_path),
                    "output": str(item.output_path) if item.output_path else None,
                    "tier": item.tier,
                    "error": item.error,
                    "mapping_summary": summarize(item.mapping),
                    "mapping": {
                        f: {
                            "value": r.value,
                            "source": r.source,
                            "confidence": r.confidence,
                        }
                        for f, r in item.mapping.items()
                    },
                    "discrepancies": [asdict(d) for d in item.discrepancies],
                }
                for item in self.items
            ],
            "extras": self.extras,
        }


def _classify_tier(
    mapping: dict[str, MappingResult],
    discrepancies: list[Discrepancy],
    *,
    schemas: list[FieldSchema],
) -> str:
    """Pick a confidence tier based on extraction sources + diff severity.

    - ``high``: every required field came from regex AND no critical discrepancy.
    - ``medium``: any LLM-sourced field OR warning-level discrepancy, no missing
      required, no critical discrepancy.
    - ``low``: any missing required field OR any critical discrepancy.
    - ``low``: also when the template defines no schema fields (likely a
      configuration mistake that would otherwise count as ``high``).
    """
    if not schemas:
        return "low"

    required_names = {s.name for s in schemas if s.required}
    sources = {r.source for r in mapping.values()}
    has_missing_required = any(
        mapping.get(n) is None or mapping[n].source == "missing" for n in required_names
    )
    has_critical = any(d.severity == "critical" for d in discrepancies)
    has_warning = any(d.severity == "warning" for d in discrepancies)
    has_llm = "llm" in sources

    if has_missing_required or has_critical:
        return "low"
    if has_llm or has_warning:
        return "medium"
    return "high"


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _single_pass_replace(text: str, replacements: dict[str, str]) -> str:
    """Replace every token in *text* in a single regex pass.

    Avoids the order-dependent re-substitution bug where a value containing
    ``{{B}}`` would be re-matched when the loop processed ``B``.
    """
    if not replacements or not text:
        return text
    # Sort tokens longest-first so e.g. ``{{NOME_COMPLETO}}`` wins over ``{{NOME}}``.
    keys = sorted(replacements, key=len, reverse=True)
    pattern = re.compile("|".join(re.escape(k) for k in keys))
    return pattern.sub(lambda m: replacements[m.group(0)], text)


def _replace_tokens_in_paragraph(
    paragraph,  # type: ignore[no-untyped-def]
    replacements: dict[str, str],
) -> None:
    """Replace every key of ``replacements`` with its value inside *paragraph*.

    Walks every ``<w:r>`` element via XPath (not just ``paragraph.runs``),
    so it covers runs nested inside hyperlinks, smart-tags, sdt content
    controls, and similar containers.

    Two-pass strategy:

    1. Per-run replacement — when a token sits within a single run it is
       replaced in place, preserving the run's formatting.
    2. Paragraph-level fallback — when the token spans multiple runs (which
       Word produces routinely on edited templates: ``{{`` in run A,
       ``NAME`` in run B, ``}}`` in run C), pass 1 misses it. Pass 2 builds
       the concatenated paragraph text, applies remaining replacements, drops
       the result into the first text-bearing element, and clears the rest.
    """
    p_elem = paragraph._p
    text_elements = p_elem.findall(f".//{_W_NS}t")
    if not text_elements:
        return

    # Pass 1 — per-text-element replacement (preserves intra-paragraph formatting).
    for t_elem in text_elements:
        original = t_elem.text or ""
        if not original:
            continue
        new_text = _single_pass_replace(original, replacements)
        if new_text != original:
            t_elem.text = new_text

    # Pass 2 — paragraph-level fallback for tokens that survived because they
    # span multiple <w:t> elements (i.e. multiple runs).
    full_text = "".join((t.text or "") for t in text_elements)
    if not any(token in full_text for token in replacements):
        return

    merged = _single_pass_replace(full_text, replacements)
    text_elements[0].text = merged
    for t_elem in text_elements[1:]:
        t_elem.text = ""


def _apply_mapping_to_template(
    template_path: Path,
    mapping: dict[str, MappingResult],
    schemas: list[FieldSchema],
    output_path: Path,
) -> Path:
    """Direct token-substitution renderer.

    Copies the template ``.docx`` to ``output_path`` and replaces every
    placeholder token with the mapped value (or empty string when missing).
    Handles tokens fragmented across multiple ``<w:r>`` runs — a common
    failure mode of naive renderers on Word-edited templates.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(template_path, output_path)
    doc = Document(str(output_path))

    replacements = {
        s.placeholder_token: (mapping[s.name].value or "")
        if s.name in mapping and mapping[s.name].value is not None
        else ""
        for s in schemas
    }

    def _walk(container) -> None:  # type: ignore[no-untyped-def]
        for paragraph in container.paragraphs:
            _replace_tokens_in_paragraph(paragraph, replacements)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    _walk(cell)

    _walk(doc)
    for section in doc.sections:
        _walk(section.header)
        _walk(section.footer)

    doc.save(str(output_path))
    return output_path


async def _process_one(
    source_path: Path,
    template_path: Path,
    schemas: list[FieldSchema],
    inferred_patterns: dict[str, InferredPattern],
    output_dir: Path,
    *,
    llm: LLMProvider | None,
    enable_semantic_diff: bool,
    audit: AuditLog | None = None,
) -> BatchItemResult:
    """Process a single source doc end-to-end."""
    # Include the source extension in the output stem so `doc1.docx` and
    # `doc1.pdf` don't both render to `doc1.normalized.docx` (the second
    # would silently overwrite the first).
    suffix_label = source_path.suffix.lstrip(".")
    output_path = output_dir / f"{source_path.stem}.{suffix_label}.normalized.docx"
    try:
        source = extract(source_path)
        if audit is not None:
            from engine.security.audit import sha256_hex

            audit.log_event(
                "batch.item_start",
                doc_hash=sha256_hex(source.text),
                fields_touched=[s.name for s in schemas],
                extra={"source": str(source_path)},
            )

        mapping = await map_hybrid(schemas, inferred_patterns, source.text, llm=llm)
        _apply_mapping_to_template(template_path, mapping, schemas, output_path)

        if audit is not None:
            from engine.security.audit import sha256_hex

            for r in mapping.values():
                audit.log_event(
                    "hybrid_mapper.field",
                    doc_hash=sha256_hex(source.text),
                    source=r.source,
                    fields_touched=[r.field],
                    llm_provider=getattr(llm, "name", None) if r.source == "llm" else None,
                    llm_model=getattr(llm, "model", None) if r.source == "llm" else None,
                )

        discrepancies: list[Discrepancy] = []
        if enable_semantic_diff and llm is not None:
            try:
                discrepancies = await diff_documents(
                    source_path,
                    output_path,
                    llm=llm,
                    schemas=schemas,
                )
                if audit is not None:
                    audit.log_event(
                        "semantic_diff.done",
                        doc_hash=sha256_hex(source.text),
                        llm_provider=getattr(llm, "name", None),
                        llm_model=getattr(llm, "model", None),
                        extra={"discrepancy_count": len(discrepancies)},
                    )
            except (TimeoutError, LLMError, ValueError) as exc:
                log.warning(
                    "batch.semantic_diff_failed",
                    path=str(source_path),
                    error=str(exc),
                )

        tier = _classify_tier(mapping, discrepancies, schemas=schemas)
        if audit is not None:
            audit.log_event(
                "batch.item_end",
                doc_hash=sha256_hex(source.text),
                extra={"tier": tier, "output": str(output_path)},
            )
        return BatchItemResult(
            source_path=source_path,
            output_path=output_path,
            tier=tier,
            mapping=mapping,
            discrepancies=discrepancies,
        )
    except Exception as exc:
        log.warning(
            "batch.item_failed",
            path=str(source_path),
            error=str(exc),
            error_type=type(exc).__name__,
        )
        return BatchItemResult(
            source_path=source_path,
            output_path=None,
            tier="error",
            mapping={},
            discrepancies=[],
            error=str(exc),
        )


async def normalize_batch(
    template_path: Path,
    source_dir: Path,
    output_dir: Path,
    *,
    llm: LLMProvider | None = None,
    field_examples: dict[str, list[str]] | None = None,
    gold_docs: list[str] | None = None,
    enable_semantic_diff: bool = True,
    max_concurrent: int = _DEFAULT_MAX_CONCURRENT,
    local_only: bool = False,
    audit: AuditLog | None = None,
) -> BatchReport:
    """Run the full pipeline over a directory of source docs.

    Args:
        template_path: ``.docx`` template with placeholder tokens.
        source_dir: directory containing ``.docx``/``.pdf`` source files.
        output_dir: where normalized outputs and ``report.json`` are written.
        llm: optional :class:`LLMProvider` for schema enrichment, hybrid fallback,
            and semantic diff. When ``None``, runs in regex-only mode and skips diff.
        field_examples: pre-built ``{field_name: [examples]}`` for
            :func:`engine.pattern_inference.infer_field_patterns`. Optional —
            without it the regex tier is skipped (everything routes to LLM).
        gold_docs: list of gold doc texts feeding pattern inference. Required if
            ``field_examples`` is provided.
        enable_semantic_diff: skip the post-normalization LLM diff when ``False``.
        max_concurrent: bound on parallel doc processing (default 4).
        local_only: when ``True``, raise :class:`RefusedRemoteCallError` if any
            LLM-bearing input is supplied. Use for regulated deployments where
            no document content may leave the host.

    Returns:
        :class:`BatchReport` with per-doc outcomes and aggregate counts.
    """
    if local_only and llm is not None:
        from engine.security.local_only import RefusedRemoteCallError

        raise RefusedRemoteCallError("normalize_batch with local_only=True received an LLM provider")

    output_dir.mkdir(parents=True, exist_ok=True)

    started = datetime.now(UTC).isoformat()
    log.info(
        "batch.start",
        template=str(template_path),
        source_dir=str(source_dir),
        output_dir=str(output_dir),
    )

    # 1. Schema inference
    template_text = extract(template_path).text
    schemas = detect_placeholders(template_text)
    if llm is not None:
        schemas = await enrich_with_llm(schemas, llm)

    # 2. Pattern inference (optional)
    inferred_patterns: dict[str, InferredPattern] = {}
    if field_examples and gold_docs:
        inferred_patterns = infer_field_patterns(
            gold_docs=gold_docs,
            field_examples=field_examples,
        )

    # 3. Discover source docs
    source_paths = sorted(p for p in source_dir.iterdir() if p.suffix.lower() in _SUPPORTED_EXTS)
    log.info("batch.discovered", count=len(source_paths))

    # 4. Parallel processing
    sem = asyncio.Semaphore(max_concurrent)

    async def _bounded(p: Path) -> BatchItemResult:
        async with sem:
            return await _process_one(
                p,
                template_path,
                schemas,
                inferred_patterns,
                output_dir,
                llm=llm,
                enable_semantic_diff=enable_semantic_diff,
                audit=audit,
            )

    items = await asyncio.gather(*(_bounded(p) for p in source_paths))

    # 5. LLM call count (rough estimate)
    llm_calls = 0
    if llm is not None:
        llm_calls += len(schemas)  # schema enrichment
        for item in items:
            sources = {r.source for r in item.mapping.values()}
            if "llm" in sources or any(r.source == "missing" for r in item.mapping.values()):
                llm_calls += 1
            if enable_semantic_diff and item.output_path is not None:
                llm_calls += 1

    finished = datetime.now(UTC).isoformat()
    report = BatchReport(
        template_path=template_path,
        output_dir=output_dir,
        schemas=schemas,
        items=list(items),
        started_at=started,
        finished_at=finished,
        llm_call_count=llm_calls,
    )

    log.info("batch.done", **report.by_tier, llm_calls=llm_calls)
    return report
