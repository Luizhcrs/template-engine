from __future__ import annotations

from dataclasses import dataclass
from typing import TYPE_CHECKING

import pdfplumber
import structlog
from docx import Document

if TYPE_CHECKING:
    from pathlib import Path

log = structlog.get_logger(__name__)


@dataclass
class ExtractedDoc:
    """Output of document extraction, consumed by llm_mapper and validator.

    - text: flat concatenation (paragraphs + tables as " | " joined rows).
    - paragraphs: original paragraphs in order, non-empty only.
    - tables: 3-level nested list -- tables x rows x cells (all strings).
    - header_fields: heuristic key→value map. Currently only {"raw_header": str}
      for .docx or {} for .pdf. Consumers must not rely on specific keys.
    """

    text: str
    paragraphs: list[str]
    tables: list[list[list[str]]]
    header_fields: dict[str, str]


def extract(path: Path) -> ExtractedDoc:
    ext = path.suffix.lower()
    log.info("extractor.start", path=str(path), ext=ext)
    if ext == ".docx":
        result = _extract_docx(path)
    elif ext == ".pdf":
        result = _extract_pdf(path)
    else:
        log.warning("extractor.unsupported_format", path=str(path), ext=ext)
        raise ValueError(f"Formato não suportado: {ext}")
    log.info(
        "extractor.ok",
        path=str(path),
        paragraphs=len(result.paragraphs),
        tables=len(result.tables),
        chars=len(result.text),
    )
    return result


def _extract_docx(path: Path) -> ExtractedDoc:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = [[[c.text for c in row.cells] for row in t.rows] for t in doc.tables]

    header_text_parts: list[str] = []
    for section in doc.sections:
        header = section.header
        for p in header.paragraphs:
            if p.text.strip():
                header_text_parts.append(p.text)

    text_blocks = paragraphs[:]
    for t in tables:
        for row in t:
            text_blocks.append(" | ".join(row))
    text = "\n".join(text_blocks)

    return ExtractedDoc(
        text=text,
        paragraphs=paragraphs,
        tables=tables,
        header_fields={"raw_header": "\n".join(header_text_parts)},
    )


def _extract_pdf(path: Path) -> ExtractedDoc:
    paragraphs: list[str] = []
    tables: list[list[list[str]]] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            for line in t.split("\n"):
                if line.strip():
                    paragraphs.append(line)
            for tbl in page.extract_tables() or []:
                tables.append([[c or "" for c in row] for row in tbl])
    text = "\n".join(paragraphs)
    return ExtractedDoc(text=text, paragraphs=paragraphs, tables=tables, header_fields={})
